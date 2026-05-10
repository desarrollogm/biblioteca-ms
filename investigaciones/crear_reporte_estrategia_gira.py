import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_estrategia_presidenta():
    doc = Document()
    
    # Configuración de estilos base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Título (Arial 14, Bold)
    titulo_texto = 'ESTRATEGIA TÁCTICA: ENTREGA DE CARPETA EJECUTIVA - GIRA SONORA MAZO 2026'
    heading = doc.add_heading('', level=0)
    run = heading.add_run(titulo_texto)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Documento de Trabajo - CORE del Magisterio Sonorense').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('-' * 60)
    
    # 1. Análisis del Cartel
    h1 = doc.add_heading('1. ANÁLISIS DEL MENSAJE VISUAL', level=1)
    doc.add_paragraph('El cartel proyecta una "Movilización Institucional". No es de confrontación, es de gestión. El uso de la imagen propositiva de la Presidenta facilita que los equipos de seguridad no vean en la comisión una amenaza, sino una representación ciudadana organizada.')
    
    # 2. El "Elevator Pitch"
    h2 = doc.add_heading('2. EL PITCH DE CONTACTO (15 SEGUNDOS)', level=1)
    p_pitch = doc.add_paragraph()
    run_p = p_pitch.add_run('“¡Presidenta! Somos el Magisterio Sonorense. Traemos la propuesta de la PIP para el ISSSTE que ya está en el Senado. Es la solución técnica para las pensiones sin afectar las finanzas. Aquí está la carpeta.”')
    run_p.italic = True
    run_p.bold = True
    
    # 3. Puntos de Contacto
    h3 = doc.add_heading('3. UBICACIONES ESTRATÉGICAS (SUR-CENTRO)', level=1)
    doc.add_paragraph('• Ciudad Obregón (CBTA 197 - Providencia): Punto crítico de "Localía". Si el acceso es restringido, montar valla de honor en la entrada principal para atraer la atención de la comitiva.', style='List Bullet')
    doc.add_paragraph('• Navojoa (Hospital IMSS): Punto formal, ideal para entrega vía asistentes de primer anillo.', style='List Bullet')
    doc.add_paragraph('• Ures (Río Sonora): Punto comunitario, ideal para acercamiento físico directo por la naturaleza abierta del evento.', style='List Bullet')
    
    # 4. Recomendaciones "De Oro"
    h4 = doc.add_heading('4. RECOMENDACIONES TÁCTICAS', level=1)
    doc.add_paragraph('• Uniformidad: Usar camisas blancas/polo con logo del MS y gafete de la escuela.', style='List Bullet')
    doc.add_paragraph('• Deslinde Visual: Mantener distancia de la CNTE. Si ellos gritan, el MS guarda silencio solemne con la carpeta en alto.', style='List Bullet')
    doc.add_paragraph('• Evidencia: Documentar la entrega (foto/video) es vital para la legitimidad ante la base.', style='List Bullet')
    
    doc.add_paragraph('\n' + '-' * 60)
    final = doc.add_paragraph('El éxito es la foto de la entrega y el compromiso de una Mesa Técnica.')
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Estrategia_Gira_Presidenta_V2_Mayo2026.docx'
    doc.save(output_path)
    print(f"Estrategia guardada en: {output_path}")

if __name__ == "__main__":
    crear_estrategia_presidenta()
