import os
from docx import Document
from docx.shared import Pt

# 1. Append to Markdown
md_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\inteligencia_coahuila_mayo_2026.md'

wa_content = """
***

# Reporte de Inteligencia: Monitoreo de Grupos MS WhatsApp - Post 1ro de Mayo 2026

**Objetivo:** Análisis del estado anímico, táctico y organizativo de la base magisterial en los principales grupos de WhatsApp del Magisterio Sonorense y aliados tras la movilización del 1ro de Mayo.

## 1. Análisis General del Ecosistema WhatsApp
La estructura de comunicación del MS a través de WhatsApp es altamente descentralizada pero temáticamente unificada. Se observa un fuerte contraste entre los grupos "núcleo" (donde se discute estrategia política sindical) y los grupos regionales (donde domina el apoyo logístico, emocional y la difusión directa de posicionamientos). El rechazo a la actual dirigencia del SNTE es transversal y unánime en todas las regiones analizadas.

## 2. Desglose Estratégico por Grupos

**Pensión Intergeneracional Protegida (PIP):**
*   **Tonalidad:** Altamente combativo y propagandístico.
*   **Foco:** Sirve como el principal motor de consignas ("Sonora no se rinde", "SNTE traidor"). Rechazo total al "charrismo".
*   **Actores Clave:** Sarahí Gastélum y Profe Kristian David Copado (Cajeme) fungen como distribuidores principales de material audiovisual y motivacional de la marcha.

**MS (Magisterio Sonorense - Núcleo Estratégico):**
*   **Tonalidad:** Analítica, política y competitiva.
*   **Foco:** Es el "War Room" digital. Se discute intensamente la pugna por la representación sindical en Sonora (MS TE vs. CNTE). Aquí se filtran reportes de inteligencia y se debate la legitimidad de las facciones.
*   **Actores Clave:** Profe Jesus A Diaz y Profe Juan Pablo, quienes lideran la narrativa estratégica del movimiento.

**Región Hermosillo MS:**
*   **Tonalidad:** Solidaria y resiliente.
*   **Foco:** Documentación de campo ("en pie de lucha"). Se comparte contenido logístico y multimedia en tiempo real de la marcha. Sirve también como red de apoyo ante ausencias justificadas (ej. problemas de salud de miembros como Judith).

**Magisterio Sonorense ❤️🖤:**
*   **Tonalidad:** Empática y comunitaria.
*   **Foco:** La "Hermandad" del magisterio. El grupo trasciende lo político para enfocarse en el soporte emocional (ej. condolencias al Maestro Rosario). Esta cohesión emocional blindifica al grupo contra divisionismos externos. Liderado orgánicamente por coordinadores regionales.

**LEY DEL ISSSTE REGION YAQUI:**
*   **Tonalidad:** Organizada y determinada.
*   **Foco:** Coordinación inter-regional (conexión con Guaymas y Nogales). Su táctica es sincronizar la distribución de contenido ("vienen más acciones") para generar tendencias.
*   **Actores Clave:** El usuario "~Ya casi dos años!" actúa como pivote de difusión y mantenimiento de la consciencia a largo plazo.

**No a la Ley del ISSSTE Nogales y Navojoa y Álamos - Magisterio Sonorense:**
*   **Tonalidad:** Acción directa, focalizada localmente.
*   **Foco:** Presión política territorial. Es el canal de distribución de los posicionamientos locales (ej. video del posicionamiento en Nogales frente a autoridades municipales).
*   **Actores Clave:** Lupita Bórquez y Profe Gonzalo (Coordinador Nogales) dirigen la resistencia táctica en el norte y sur del estado.

**LEY ISSSTE 2025 YAQUI 2:**
*   **Tonalidad:** Solemne, motivacional, visión de futuro.
*   **Foco:** Mantener el objetivo a largo plazo (2025). El discurso se centra en "trascender la comodidad". Consideran que cada acción, por pequeña que sea en números, es gigante en consciencia política.
*   **Actores Clave:** Delegada Sindical Paloma lidera el marco filosófico de la lucha en este espacio.

## 3. Conclusión y Recomendación Operativa para el Core del MS
El control narrativo en WhatsApp es sólido. La estructura de grupos funciona como un sistema de poleas: lo que se acuerda en el grupo "MS" se irradia logísticamente a "Región Yaqui" y emocionalmente a "Magisterio Sonorense ❤️🖤".

**Acción Inmediata:** Capitalizar la alta moral y cohesión observada. Aprovechando el ecosistema sincronizado que maneja la región Yaqui, se sugiere inyectar el *Reporte de Contraste de Posicionamientos (1ro de Mayo)* en estos canales. Los liderazgos (Paloma, Gonzalo, Jesus A Diaz) deben utilizar los grupos regionales para anclar la superioridad técnica de las propuestas del MS frente a narrativas competidoras o desgaste natural post-marcha.
"""

with open(md_path, 'a', encoding='utf-8') as f:
    f.write(wa_content)

# 2. Create the Word Document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Heading
heading = doc.add_heading('Reporte de Inteligencia: Monitoreo de Grupos MS WhatsApp - Post 1ro de Mayo 2026', 0)
heading.style.font.name = 'Arial'
heading.style.font.size = Pt(14)
heading.style.font.bold = True

p1 = doc.add_paragraph()
r1 = p1.add_run('Objetivo: ')
r1.bold = True
p1.add_run('Análisis del estado anímico, táctico y organizativo de la base magisterial en los principales grupos de WhatsApp del Magisterio Sonorense y aliados tras la movilización del 1ro de Mayo.')

doc.add_paragraph()
h1 = doc.add_paragraph()
r_h1 = h1.add_run('1. Análisis General del Ecosistema WhatsApp')
r_h1.font.size = Pt(14)
r_h1.bold = True

doc.add_paragraph('La estructura de comunicación del MS a través de WhatsApp es altamente descentralizada pero temáticamente unificada. Se observa un fuerte contraste entre los grupos "núcleo" (donde se discute estrategia política sindical) y los grupos regionales (donde domina el apoyo logístico, emocional y la difusión directa de posicionamientos). El rechazo a la actual dirigencia del SNTE es transversal y unánime en todas las regiones analizadas.')

doc.add_paragraph()
h2 = doc.add_paragraph()
r_h2 = h2.add_run('2. Desglose Estratégico por Grupos')
r_h2.font.size = Pt(14)
r_h2.bold = True

def add_group(name, tone, focus, actors):
    p = doc.add_paragraph()
    p.add_run(f'{name}\n').bold = True
    p.add_run('Tonalidad: ').bold = True
    p.add_run(f'{tone}\n')
    p.add_run('Foco: ').bold = True
    p.add_run(f'{focus}\n')
    p.add_run('Actores Clave: ').bold = True
    p.add_run(f'{actors}\n')

add_group('Pensión Intergeneracional Protegida (PIP):', 'Altamente combativo y propagandístico.', 'Sirve como el principal motor de consignas ("Sonora no se rinde", "SNTE traidor"). Rechazo total al "charrismo".', 'Sarahí Gastélum y Profe Kristian David Copado (Cajeme) fungen como distribuidores principales de material audiovisual y motivacional de la marcha.')
add_group('MS (Magisterio Sonorense - Núcleo Estratégico):', 'Analítica, política y competitiva.', 'Es el "War Room" digital. Se discute intensamente la pugna por la representación sindical en Sonora (MS TE vs. CNTE). Aquí se filtran reportes de inteligencia y se debate la legitimidad de las facciones.', 'Profe Jesus A Diaz y Profe Juan Pablo, quienes lideran la narrativa estratégica del movimiento.')
add_group('Región Hermosillo MS:', 'Solidaria y resiliente.', 'Documentación de campo ("en pie de lucha"). Se comparte contenido logístico y multimedia en tiempo real de la marcha. Sirve también como red de apoyo ante ausencias justificadas (ej. problemas de salud de miembros como Judith).', 'Sin especificar, pero con fuerte sentido comunitario.')
add_group('Magisterio Sonorense ❤️🖤:', 'Empática y comunitaria.', 'La "Hermandad" del magisterio. El grupo trasciende lo político para enfocarse en el soporte emocional (ej. condolencias al Maestro Rosario). Esta cohesión emocional blindifica al grupo contra divisionismos externos. Liderado orgánicamente por coordinadores regionales.', 'Coordinadores regionales activos.')
add_group('LEY DEL ISSSTE REGION YAQUI:', 'Organizada y determinada.', 'Coordinación inter-regional (conexión con Guaymas y Nogales). Su táctica es sincronizar la distribución de contenido ("vienen más acciones") para generar tendencias.', 'El usuario "~Ya casi dos años!" actúa como pivote de difusión y mantenimiento de la consciencia a largo plazo.')
add_group('No a la Ley del ISSSTE Nogales y Navojoa y Álamos - Magisterio Sonorense:', 'Acción directa, focalizada localmente.', 'Presión política territorial. Es el canal de distribución de los posicionamientos locales (ej. video del posicionamiento en Nogales frente a autoridades municipales).', 'Lupita Bórquez y Profe Gonzalo (Coordinador Nogales) dirigen la resistencia táctica en el norte y sur del estado.')
add_group('LEY ISSSTE 2025 YAQUI 2:', 'Solemne, motivacional, visión de futuro.', 'Mantener el objetivo a largo plazo (2025). El discurso se centra en "trascender la comodidad". Consideran que cada acción, por pequeña que sea en números, es gigante en consciencia política.', 'Delegada Sindical Paloma lidera el marco filosófico de la lucha en este espacio.')

doc.add_paragraph()
h3 = doc.add_paragraph()
r_h3 = h3.add_run('3. Conclusión y Recomendación Operativa para el Core del MS')
r_h3.font.size = Pt(14)
r_h3.bold = True

doc.add_paragraph('El control narrativo en WhatsApp es sólido. La estructura de grupos funciona como un sistema de poleas: lo que se acuerda en el grupo "MS" se irradia logísticamente a "Región Yaqui" y emocionalmente a "Magisterio Sonorense ❤️🖤".')

p_rec = doc.add_paragraph()
p_rec.add_run('Acción Inmediata: ').bold = True
p_rec.add_run('Capitalizar la alta moral y cohesión observada. Aprovechando el ecosistema sincronizado que maneja la región Yaqui, se sugiere inyectar el Reporte de Contraste de Posicionamientos (1ro de Mayo) en estos canales. Los liderazgos (Paloma, Gonzalo, Jesus A Diaz) deben utilizar los grupos regionales para anclar la superioridad técnica de las propuestas del MS frente a narrativas competidoras o desgaste natural post-marcha.')

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Inteligencia_WhatsApp_MS_2026.docx'
doc.save(output_path)
print(f"WA Report saved to {output_path}")
