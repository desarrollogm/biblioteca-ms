import os
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Heading
heading = doc.add_heading('Reporte de Inteligencia: Contraste de Posicionamientos - 1ro de Mayo 2026', 0)
heading.style.font.name = 'Arial'
heading.style.font.size = Pt(14)
heading.style.font.bold = True

p = doc.add_paragraph()
r = p.add_run('Objetivo: ')
r.bold = True
p.add_run('Análisis discursivo y táctico del pronunciamiento de la CNTE (Nogales) vs. el Magisterio Sonorense (Mtra. Paloma) durante la marcha del 1ro de Mayo.')

doc.add_paragraph()

h1 = doc.add_paragraph()
r1 = h1.add_run('1. Análisis del Posicionamiento CNTE Nogales')
r1.font.size = Pt(14)
r1.bold = True

p2 = doc.add_paragraph()
p2.add_run('Vocera: ').bold = True
p2.add_run('Maestra de Nogales (Posicionamiento leído).\n')
p2.add_run('Tono: ').bold = True
p2.add_run('Ideológico, confrontativo global, maximalista (exige el todo o nada).\n')
p2.add_run('Enfoque Principal: ').bold = True
p2.add_run('Anticapitalismo, lucha de clases, abrogación total de leyes.\n')

p3 = doc.add_paragraph()
p3.add_run('Análisis de Fondo ("Hasta la cocina"):').bold = True

doc.add_paragraph('- Dispersión del Mensaje: La oradora de la CNTE inicia atacando el sistema de pensiones (lo cual es válido), pero rápidamente diluye el impacto al mezclarlo con demandas internacionales y ajenas al magisterio local. Pide romper relaciones con Israel, menciona a Palestina, habla de la desaparición de defensores de la tierra y los 43 de Ayotzinapa. Esta es su mayor debilidad táctica: al disparar a todos lados, no aciertan en el blanco principal que le duele al maestro sonorense hoy (sus pensiones reales).', style='List Bullet')
doc.add_paragraph('- Retórica de "Guerra": Utiliza términos como "gobierno neoliberal", "capitalismo que genera la muerte", "genocida", "oligarquías". Es un discurso de los años 70s que resuena con un sector muy radicalizado, pero aliena a la gran mayoría de la base magisterial que busca soluciones técnicas y legales, no una revolución armada.', style='List Bullet')
doc.add_paragraph('- Cero Propuesta Técnica: Exigen la abrogación de la Ley del ISSSTE 2007 y de la Ley Educativa 2019, así como 100% de aumento salarial y 40 horas laborales. Sin embargo, no explican CÓMO lo van a lograr, ni presentan corridas financieras, ni propuestas legislativas (como la PIP o la UMA-Fronteriza del MS). Su táctica es exigir y marchar, esperando que el gobierno ceda por presión, lo cual históricamente no ha devuelto el régimen solidario.', style='List Bullet')
doc.add_paragraph('- Lectura Plana: La oradora lee un documento prefabricado, mirando constantemente el papel, lo que resta pasión y conexión visual con la gente que marcha o escucha.', style='List Bullet')

doc.add_paragraph()

h2 = doc.add_paragraph()
r2 = h2.add_run('2. Análisis del Posicionamiento Magisterio Sonorense (Mtra. Paloma)')
r2.font.size = Pt(14)
r2.bold = True

p4 = doc.add_paragraph()
p4.add_run('Vocera: ').bold = True
p4.add_run('Mtra. Paloma.\n')
p4.add_run('Tono: ').bold = True
p4.add_run('Firme, propositivo, localizado, institucionalmente crítico, enfocado en derechos laborales tangibles.\n')
p4.add_run('Enfoque Principal: ').bold = True
p4.add_run('Rescate de pensiones, caja de ahorro, quinquenios, sistema de salud y rendición de cuentas sindical.\n')

p5 = doc.add_paragraph()
p5.add_run('Análisis de Fondo ("Hasta la cocina"):').bold = True

doc.add_paragraph('- Foco Láser en el Problema Local y Real: Paloma no habla de Medio Oriente; habla de Cajeme, Navojoa, Hermosillo. Habla del Décimo Transitorio, del saqueo a las cajas de ahorro (FAS), de la falta de rendición de cuentas de la Sección 28. Este discurso conecta inmediatamente con el bolsillo y la realidad del maestro sonorense.', style='List Bullet')
doc.add_paragraph('- Conocimiento de la Ley (Técnica Jurídica): El MS habla de "zonificación del 100%", "pago de quinquenios conforme a la ley vigente", "cuentas individuales que lastiman el futuro del retiro". Tienen mapeado exactamente qué artículos y qué fondos están fallando. No exigen "100% de aumento" al aire, exigen el pago correcto de lo que ya está legislado (Quinquenios Concepto 07).', style='List Bullet')
doc.add_paragraph('- Ataque Directo a la Cúpula Sindical ("El Doble Discurso"): Paloma es contundente contra la dirigencia del SNTE (Sección 28). Denuncia la farsa de que el SNTE está exigiendo cambios ahora, cuando "desde hace años están entregados a los diferentes gobiernos". El MS ataca la representación, exigiendo que las cuotas sindicales se transparenten ("transparencia y rendición de cuentas por parte del FAS").', style='List Bullet')
doc.add_paragraph('- Conexión con la Base (Lenguaje Corporal y Voz): Paloma improvisa o discursa con conocimiento del tema (no está leyendo un guion plano). Su voz refleja verdadera indignación y urgencia, interactuando con los contingentes (mira cómo la gente de las camisetas rojas está estructurada y en bloque). Es un discurso de liderazgo, no de resistencia abstracta.', style='List Bullet')

doc.add_paragraph()

h3 = doc.add_paragraph()
r3 = h3.add_run('3. Conclusión y Recomendación Operativa')
r3.font.size = Pt(14)
r3.bold = True

p6 = doc.add_paragraph()
p6.add_run('Conclusión "hasta la cocina": ').bold = True
p6.add_run('La CNTE hizo nuestro trabajo más fácil. Al salir con un discurso tan despegado de la realidad inmediata del docente sonorense (hablando de conflictos bélicos internacionales en lugar de cómo mejorar la jubilación local), el Magisterio Sonorense tiene la vía libre para presentarse ante las bases indecisas y los grupos de WhatsApp (como el de Nogales) como la ÚNICA fuerza seria, técnica y con posibilidades reales de recuperar las pensiones.')

doc.add_paragraph()

p7 = doc.add_paragraph()
p7.add_run('Acción Recomendada: ').bold = True
p7.add_run('Compartir el video de la Mtra. Paloma, acompañado de un texto corto que resalte: "Compañeros, la diferencia es clara. Mientras unos marchan para pedir la paz en el mundo diluyendo nuestra lucha magisterial, el Magisterio Sonorense exige HOY el pago de nuestros quinquenios, la auditoría al FAS y una pensión digna. El MS propone y lucha por lo nuestro. Aquí no hay cuentos, hay técnica y defensa real."')

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Inteligencia_1roMayo_2026.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
