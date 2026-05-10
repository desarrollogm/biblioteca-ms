import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, font_name='Arial', font_size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold

def create_doc():
    doc = Document()
    
    # Configuración de márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Logo
    logo_path = r'C:\Users\USUARIO\Downloads\logo MS.jpeg'
    if os.path.exists(logo_path):
        # Insertamos el logo centrado
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(2.0))
    
    # Título Principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('\nCONCEPTO E IDEOLOGÍA DEL MAGISTERIO SONORENSE')
    set_font(run, font_size=14, bold=True)
    
    # Autoría
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = author_para.add_run('Análisis de Movimiento para: EL JAD')
    set_font(run, font_size=11, bold=True)

    # Sección 1: El Concepto
    p = doc.add_paragraph()
    run = p.add_run('\n1. EL CONCEPTO: LA VANGUARDIA TÉCNICA')
    set_font(run, font_size=12, bold=True)
    
    p = doc.add_paragraph()
    run = p.add_run('El Magisterio Sonorense (MS) no es un sindicato tradicional; es un Cuerpo Técnico de Resistencia Institucional. Su esencia radica en la Soberanía Técnica: no se limita a exigir mejoras, sino que diseña las soluciones legislativas y actuariales necesarias (como la PIP) para que el Estado no tenga argumentos para la negativa.')
    set_font(run, font_size=12)

    # Sección 2: La Ideología
    p = doc.add_paragraph()
    run = p.add_run('\n2. LA IDEOLOGÍA: EL ADN DEL MOVIMIENTO')
    set_font(run, font_size=12, bold=True)

    ideas = [
        ('Racionalidad Técnica:', 'La verdad matemática y jurídica es más potente que el grito. Un maestro informado es un negociador infalible.'),
        ('Regionalismo y Dignidad:', 'Defensa de las condiciones de Sonora (UMA-Fronteriza, Rezonificación). El reconocimiento del costo de vida del norte.'),
        ('Justicia Generacional:', 'Lucha por rescatar el futuro de los jubilados y garantizar pensiones dignas para las nuevas generaciones.'),
        ('Autodeterminación:', 'Rechazo a la mediación de burocracias sindicales que negocian derechos por prebendas.')
    ]

    for titulo, desc in ideas:
        p = doc.add_paragraph(style='List Bullet')
        run_t = p.add_run(titulo)
        set_font(run_t, font_size=12, bold=True)
        run_d = p.add_run(f' {desc}')
        set_font(run_d, font_size=12)

    # Sección 3: Análisis "Hasta la Cocina"
    p = doc.add_paragraph()
    run = p.add_run('\n3. ESTRATEGIA DE INTELIGENCIA')
    set_font(run, font_size=12, bold=True)
    
    p = doc.add_paragraph()
    run = p.add_run('El MS opera bajo una lógica de Infiltración de Narrativa, logrando que términos técnicos sean hoy parte de la agenda oficial, neutralizando a los sindicatos oficiales mediante la superioridad documental y el manejo de datos.')
    set_font(run, font_size=12)

    # Cierre
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n"Porque la razón es nuestra fuerza, y la técnica nuestra herramienta."')
    set_font(run, font_size=12, bold=True)

    # Guardar
    output_path = r'c:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Identidad_y_Concepto_MS.docx'
    doc.save(output_path)
    print(f'Documento generado en: {output_path}')

if __name__ == '__main__':
    create_doc()
