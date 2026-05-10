from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilo base Arial 12
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

# ── Helpers ───────────────────────────────────────────────────────────────────

def titulo_principal(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

def subtitulo_portada(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

def seccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

def subseccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True

def parrafo(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12)

def parrafo_bold(doc, etiqueta, texto):
    p = doc.add_paragraph()
    r1 = p.add_run(etiqueta); r1.font.name = 'Arial'; r1.font.size = Pt(12); r1.bold = True
    r2 = p.add_run(texto);    r2.font.name = 'Arial'; r2.font.size = Pt(12)

def parrafo_nota(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def bala(doc, texto):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(texto); r.font.name = 'Arial'; r.font.size = Pt(12)

def espacio(doc):
    doc.add_paragraph()

def add_table(doc, encabezados, filas):
    table = doc.add_table(rows=len(filas)+1, cols=len(encabezados))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(encabezados):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True; run.font.name = 'Arial'; run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3A6E')
        tcPr.append(shd)
    for i, fila in enumerate(filas):
        fill = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
        for j, texto in enumerate(fila):
            cell = table.rows[i+1].cells[j]
            cell.text = texto
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Arial'; run.font.size = Pt(11); run.font.bold = False
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    return table

# ==============================================================================
# PORTADA
# ==============================================================================
titulo_principal(doc, 'REPORTE DE INTELIGENCIA ESTRATEGICA')
titulo_principal(doc, 'SECCION 54 DEL SNTE & ISSSTESON')
espacio(doc)
subtitulo_portada(doc, 'Anatomia del sindicato estatal de Sonora y su sistema de pensiones')
espacio(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Para uso exclusivo del CORE del Magisterio Sonorense')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
espacio(doc)
parrafo_bold(doc, 'Clasificacion: ', 'Inteligencia estrategica — uso interno MS')
parrafo_bold(doc, 'Fecha de analisis: ', '4 de mayo de 2026')
parrafo_bold(doc, 'Fuentes: ', 'snte54.com.mx, sonora.gob.mx (ISSSTESON), radiosonora.com.mx, oem.com.mx, elfinanciero.com.mx, staus.mx')

doc.add_page_break()

# ==============================================================================
# PARTE I — FICHA DE IDENTIDAD
# ==============================================================================
seccion(doc, 'PARTE I — FICHA DE IDENTIDAD: SECCION 54 DEL SNTE')
espacio(doc)
add_table(doc,
    ['Campo', 'Dato'],
    [
        ['Nombre oficial', 'Seccion 54 del Sindicato Nacional de Trabajadores de la Educacion'],
        ['Tipo de trabajadores', 'Maestros de plaza ESTATAL (pagados por el Gobierno del Estado de Sonora)'],
        ['Institucion de seguridad social', 'ISSSTESON (estatal) — NO el ISSSTE federal'],
        ['Marco legal pensional', 'Ley 38 del Estado de Sonora'],
        ['Secretario General (2022-2026)', 'Mtro. Jesus Javier Ceballos Corral'],
        ['Afiliados totales', '+17,000 trabajadores de la educacion estatales (2024)'],
        ['Alineamiento politico', 'SNTE nacional — Alfonso Cepeda Salas'],
        ['Fondo complementario propio', 'CMAP (Caja Magisterial de Ahorros y Prestamos)'],
        ['Sitio web', 'snte54.com.mx / snte.org.mx/seccion54'],
    ]
)
espacio(doc)
parrafo_nota(doc,
    'DATO CRITICO PARA EL MS: La Seccion 54 y la Seccion 28 son dos universos distintos dentro del mismo magisterio sonorense. '
    'Los estatales cotizan al ISSSTESON, los federalizados cotizan al ISSSTE. El conflicto UMA-ISSSTE que lidera el MS '
    'NO les afecta directamente a los maestros de la S54. Su batalla es estatal, no federal.')

doc.add_page_break()

# ==============================================================================
# PARTE II — LIDERAZGO
# ==============================================================================
seccion(doc, 'PARTE II — PERFIL DEL LIDERAZGO: JESUS JAVIER CEBALLOS CORRAL')
espacio(doc)
subseccion(doc, 'Secretario General S54 (periodo 2022-2026)')
espacio(doc)
add_table(doc,
    ['Indicador', 'Evaluacion'],
    [
        ['Alineamiento', 'SNTE nacional (Alfonso Cepeda) — linea corporativa institucional'],
        ['Postura sobre Ley ISSSTE', 'Sin pronunciamiento publico documentado — observador distante'],
        ['Postura sobre el MS', 'Sin confrontacion abierta — zona gris tactica'],
        ['Estilo de gestion', 'Institucional, no confrontacional — perfil opuesto a CNTE'],
        ['Palancas de presion', 'Relacion con Gobernador del Estado y SEP estatal (SEC Sonora)'],
    ]
)
espacio(doc)
parrafo_nota(doc,
    'No existe evidencia de comunicados oficiales de Ceballos Corral tomando postura sobre el conflicto federalizado '
    '(Ley del ISSSTE / UMAs). Esto representa una ventana tactica: la S54 no es un adversario activo del MS.')

doc.add_page_break()

# ==============================================================================
# PARTE III — ISSSTESON: LA CRISIS
# ==============================================================================
seccion(doc, 'PARTE III — ANATOMIA DEL ISSSTESON: LA CRISIS QUE NADIE NOMBRA')
espacio(doc)
subseccion(doc, '3.1 Que es el ISSSTESON')
parrafo(doc,
    'El Instituto de Seguridad y Servicios Sociales de los Trabajadores del Estado de Sonora es el equivalente estatal del ISSSTE: '
    'administra pensiones, servicios medicos y prestaciones de los trabajadores del Gobierno de Sonora — incluyendo los maestros de '
    'la Seccion 54, empleados de municipios incorporados, Poder Judicial y Poder Legislativo estatal. Se rige por la Ley 38 del Estado de Sonora.')
espacio(doc)
subseccion(doc, '3.2 La Crisis Financiera: Numeros Crudos')
espacio(doc)
add_table(doc,
    ['Elemento', 'Detalle'],
    [
        ['Deficit presupuestario documentado', '~4,000 millones de pesos (reducido desde cifras aun mayores en años anteriores)'],
        ['Causa 1', 'Cuotas y aportaciones insuficientes frente al gasto total'],
        ['Causa 2', 'Rezago de municipios y organismos en pago de cuotas'],
        ['Causa 3', 'Crecimiento constante de pensionados sin crecimiento proporcional de cotizantes'],
        ['Situacion actuarial', 'Reservas NEGATIVAS — estudios oficiales reconocen inviabilidad a largo plazo'],
        ['Rescate externo', 'Gobierno Federal ha inyectado recursos para cubrir aguinaldos y pensiones'],
    ]
)
espacio(doc)
parrafo_nota(doc,
    'BOMBA DE TIEMPO: Los estudios actuariales del propio ISSSTESON reconocen que el sistema es insostenible sin reformas profundas. '
    'El Gobierno del Estado lo sabe. La S54 lo sabe. Ninguno lo dice abiertamente a la base. '
    'Esta es una bomba de tiempo que el MS puede activar estrategicamente.')

doc.add_page_break()

# ==============================================================================
# PARTE IV — LAS TRES GENERACIONES
# ==============================================================================
seccion(doc, 'PARTE IV — LA TRAMPA DE LAS TRES GENERACIONES (Ley 38, reforma 2005)')
espacio(doc)
parrafo(doc, 'La reforma de 2005 a la Ley 38 del ISSSTESON creo una estructura generacional que reproduce la misma logica que la Ley del ISSSTE 2007:')
espacio(doc)
add_table(doc,
    ['Generacion', 'Criterio de ingreso', 'Base de calculo', 'Tope', 'Requisitos'],
    [
        ['G1 — Derechos Adquiridos', 'Pre-2005, 30+ años cotizados al 30/jun/2005', 'Ultimo salario cotizado (100%)', 'Sin tope duro', '30 H / 28 M años de servicio'],
        ['G2 — Actual', 'Ingresaron antes de julio 2005, sin derechos adquiridos', 'Promedio ultimos 3 años', 'Sin tope duro', 'Escala gradual segun año de ingreso'],
        ['G3 — Futura (LA MAS AFECTADA)', 'Ingresaron DESPUES de julio 2005', 'Promedio ultimos 10 años (sueldo regulador)', '20 salarios minimos mensuales', '35 H / 33 M años de servicio'],
    ]
)
espacio(doc)
parrafo_nota(doc,
    'PARALELISMO EXPLOSIVO: Los maestros estatales de la Generacion Futura tienen un tope de 20 salarios minimos en ISSSTESON. '
    'Los maestros federalizados tienen un tope de 10 UMAs en ISSSTE. Dos sistemas distintos. El mismo agravio estructural. '
    'Todo maestro contratado despues de 2005 — estatal o federal — enfrenta una jubilacion topada y reducida.')

doc.add_page_break()

# ==============================================================================
# PARTE V — EL CMAP
# ==============================================================================
seccion(doc, 'PARTE V — EL CMAP: EL ACTIVO QUE DELATA LA INSUFICIENCIA DEL SISTEMA')
espacio(doc)
parrafo(doc, 'La Seccion 54 gestiona la CMAP (Caja Magisterial de Ahorros y Prestamos) como fondo complementario ante las deficiencias del ISSSTESON:')
espacio(doc)
add_table(doc,
    ['Servicio CMAP', 'Descripcion'],
    [
        ['Bono de retiro', 'Pago unico al momento de jubilarse — complemento de la pension ISSSTESON'],
        ['Prestamos a corto plazo', 'Liquidez inmediata para afiliados activos'],
        ['Prestamos prendarios', 'Sobre bienes muebles'],
        ['Prestamos hipotecarios', 'Para vivienda'],
    ]
)
espacio(doc)
parrafo_nota(doc,
    'DATO CLAVE: Que la Seccion 54 haya necesitado crear su propia caja de ahorro complementaria es el reconocimiento implicito '
    'de que el ISSSTESON no alcanza para garantizar pensiones dignas por si solo. '
    'Esto es municion argumentativa directa para el MS: "Hasta el sindicato estatal sabe que su sistema de pensiones es insuficiente."')

doc.add_page_break()

# ==============================================================================
# PARTE VI — COMPARATIVO
# ==============================================================================
seccion(doc, 'PARTE VI — MAPA COMPARATIVO: S54 (ESTATALES) vs. S28/MS (FEDERALIZADOS)')
espacio(doc)
add_table(doc,
    ['Dimension', 'Estatales — Seccion 54', 'Federalizados — Seccion 28 / MS'],
    [
        ['Patron laboral', 'Gobierno del Estado de Sonora', 'Federacion (SEP/FONE)'],
        ['Institucion pensional', 'ISSSTESON', 'ISSSTE'],
        ['Marco legal', 'Ley 38 del Estado de Sonora', 'Ley del ISSSTE 2007'],
        ['Tope pensional (generacion joven)', '20 salarios minimos mensuales', '10 UMAs'],
        ['Fondo complementario', 'CMAP (propio del sindicato)', 'Ninguno — el MS busca crearlo via PIP'],
        ['Crisis del sistema', 'Deficit de 4,000 MDP + reservas negativas', 'Agotamiento de reservas actuariales ISSSTE'],
        ['Afiliados en Sonora', '~17,000 estatales', '~40,000+ federalizados'],
        ['Postura ante el MS', 'Observador neutro — sin confrontacion', 'Base natural del movimiento'],
    ]
)

doc.add_page_break()

# ==============================================================================
# PARTE VII — ANALISIS ESTRATEGICO
# ==============================================================================
seccion(doc, 'PARTE VII — ANALISIS ESTRATEGICO: FORTALEZAS Y DEBILIDADES')
espacio(doc)
subseccion(doc, 'Fortalezas de la S54 como actor politico en Sonora')
espacio(doc)
add_table(doc,
    ['Fortaleza', 'Descripcion'],
    [
        ['Masa de afiliados', '17,000 trabajadores con voto y presencia territorial'],
        ['Acceso al Gobernador', 'Canal institucional establecido con el Gobierno del Estado'],
        ['CMAP como recurso propio', 'Independencia financiera parcial del sistema ISSSTESON'],
        ['Perfil no confrontacional', 'No genera conflictos publicos — mayor credibilidad institucional'],
        ['Presencia estatal', 'Cubre todo Sonora — incluyendo zonas donde el MS tiene menor penetracion'],
    ]
)
espacio(doc)
subseccion(doc, 'Debilidades criticas de la S54')
espacio(doc)
add_table(doc,
    ['Debilidad', 'Descripcion', 'Oportunidad para el MS'],
    [
        ['Alineada al SNTE nacional', 'Dependiente de la linea entreguista de Alfonso Cepeda', 'El MS puede ofrecer la alternativa independiente'],
        ['Sin propuesta propia de reforma', 'No ha generado ningun documento tecnico para reformar el ISSSTESON', 'El MS tiene el modelo tecnico (PIP) adaptable al ambito estatal'],
        ['ISSSTESON en quiebra silenciosa', 'No ha salido a exigir reforma del sistema publicamente', 'El MS puede forzar el debate publico antes que ellos'],
        ['Generacion Futura desprotegida', 'Sus maestros jovenes tambien estan topados — no los ha organizado', 'Universo de jovenes estatales que el MS puede activar'],
    ]
)

doc.add_page_break()

# ==============================================================================
# PARTE VIII — VECTORES TACTICOS
# ==============================================================================
seccion(doc, 'PARTE VIII — VECTORES TACTICOS PARA EL MS')
espacio(doc)

subseccion(doc, 'Vector 1: La "Doble Trampa Generacional" — Argumento de unificacion')
parrafo(doc, 'Para auditorios mixtos (estatales + federalizados):')
p = doc.add_paragraph()
r = p.add_run(
    '"Tanto los maestros federalizados (10 UMAs en ISSSTE) como los estatales mas jovenes (20 salarios minimos en ISSSTESON) '
    'estan TOPADOS por sus propios sistemas. El problema no es solo federal — es un modelo de pensiones disenado para pagar menos. '
    'La diferencia es que el Magisterio Sonorense ya tiene la solucion tecnica documentada y presentada ante el Congreso."')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
espacio(doc)

subseccion(doc, 'Vector 2: El ISSSTESON en Quiebra — Presion al Gobernador')
parrafo(doc, 'El deficit de 4,000 millones de pesos del ISSSTESON es una vulnerabilidad politica del Gobierno Estatal. Si el MS lo expone:')
bala(doc, 'Presiona al Gobernador a tomar postura en el debate pensional nacional.')
bala(doc, 'Crea un argumento de urgencia para los 17,000 maestros estatales.')
bala(doc, 'Obliga a Ceballos Corral (S54) a definirse: con el Gobernador o con los maestros?')
espacio(doc)

subseccion(doc, 'Vector 3: La CMAP como Evidencia de Insuficiencia — Argumento de contexto')
p = doc.add_paragraph()
r = p.add_run(
    '"Cuando el propio sindicato estatal tuvo que crear su propia caja de ahorros porque el ISSSTESON no alcanza, '
    'eso nos dice todo sobre el estado del sistema de pensiones en Sonora. '
    'Estatales y federalizados enfrentamos el mismo problema: sistemas disenados para que no alcance."')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
espacio(doc)

subseccion(doc, 'Vector 4: La S54 como Aliado Latente — No es la CNTE')
parrafo(doc, 'A diferencia de la CNTE (adversario tactico), la S54 es un aliado pasivo que no se ha activado. Posibles acciones:')
bala(doc, 'Proponer un Foro Tecnico Conjunto sobre pensiones en Sonora (estatales + federalizados).')
bala(doc, 'Enviar comunicacion formal a Ceballos Corral invitandolo a un posicionamiento conjunto sobre pensiones dignas en Sonora.')
bala(doc, 'Crear la narrativa: "Magisterio Sonorense unido — estatales y federalizados — merece pensiones dignas".')

doc.add_page_break()

# ==============================================================================
# PARTE IX — EVALUACION DE RIESGO
# ==============================================================================
seccion(doc, 'PARTE IX — EVALUACION DE RIESGO')
espacio(doc)
add_table(doc,
    ['Riesgo', 'Nivel', 'Descripcion'],
    [
        ['S54 permanece neutral y MS no gana ese universo', 'MEDIO', 'Sin acercamiento activo, los 17,000 estatales quedan fuera del movimiento'],
        ['Gobierno usa a la S54 para dividir al magisterio sonorense', 'ALTO', 'El Gobernador puede negociar con la S54 para aislar politicamente al MS federal'],
        ['ISSSTESON quiebra sin que el MS lo capitalice', 'MEDIO', 'Si quiebra primero, el Gobierno controla la narrativa del rescate'],
        ['Jovenes estatales (G3) se organizan independientemente', 'BAJO', 'Poco probable sin catalizador externo — el MS puede ser ese catalizador'],
    ]
)

doc.add_page_break()

# ==============================================================================
# CONCLUSION
# ==============================================================================
seccion(doc, 'CONCLUSION EJECUTIVA')
espacio(doc)
parrafo(doc,
    'La Seccion 54 y el ISSSTESON no son actores hostiles al Magisterio Sonorense — son el espejo estatal del mismo problema federal. '
    'Los 17,000 maestros estatales de Sonora enfrentan un sistema de pensiones en crisis estructural, con una generacion joven topada a '
    '20 salarios minimos, y un sindicato que creo su propia caja de ahorro porque el ISSSTESON no alcanza.')
espacio(doc)
parrafo(doc,
    'El MS tiene la oportunidad historica de construir la narrativa mas poderosa del magisterio sonorense: '
    '"No importa si eres estatal o federalizado — en Sonora, ningun maestro tiene pension digna. '
    'El Magisterio Sonorense es el unico movimiento con soluciones tecnicas para TODOS."')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run(
    'TACTICA RECOMENDADA: Tender el puente hacia la S54 antes de que el Gobierno del Estado lo use para dividir. '
    'El MS gana mas con 17,000 aliados estatales que con 17,000 neutrales.')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True; r.italic = True

espacio(doc)
doc.add_paragraph('_' * 65)
parrafo(doc, 'Reporte elaborado para uso estrategico exclusivo del CORE del Magisterio Sonorense')
parrafo(doc, 'Fecha: 04/05/2026 | Fuentes: snte54.com.mx, sonora.gob.mx (ISSSTESON), radiosonora.com.mx, oem.com.mx, elfinanciero.com.mx, staus.mx')

# ==============================================================================
# GUARDAR
# ==============================================================================
output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Intel_S54_ISSSTESON.docx'
doc.save(output_path)
print(f'LISTO. Guardado en: {output_path}')
