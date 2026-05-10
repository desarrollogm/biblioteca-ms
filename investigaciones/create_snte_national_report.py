import os
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Heading
heading = doc.add_heading('Reporte de Inteligencia: Pronunciamiento Nacional SNTE - 1ro de Mayo 2026', 0)
heading.style.font.name = 'Arial'
heading.style.font.size = Pt(14)
heading.style.font.bold = True

p1 = doc.add_paragraph()
r1 = p1.add_run('Objetivo: ')
r1.bold = True
p1.add_run('Análisis estratégico del discurso del líder nacional del SNTE (Alfonso Cepeda) frente a la Presidenta Claudia Sheinbaum, y su impacto directo en la propuesta de la Pensión Intergeneracional Protegida (PIP) del MS.')

doc.add_paragraph()
h1 = doc.add_paragraph()
r_h1 = h1.add_run('1. Desglose del Discurso ("Hasta la cocina")')
r_h1.font.size = Pt(14)
r_h1.bold = True

doc.add_paragraph('- Contexto: El evento es oficial (La Mañanera del Pueblo del 1ro de Mayo). El SNTE funge como anfitrión de la Presidenta Claudia Sheinbaum.', style='List Bullet')
doc.add_paragraph('- La Declaración Clave: El líder del SNTE declara explícitamente el objetivo de "Construir un nuevo sistema de pensiones para las y los trabajadores al servicio del estado con viabilidad jurídica y presupuestaria de largo plazo".', style='List Bullet')
doc.add_paragraph('- Alineación Política: El SNTE se subordina y se alinea completamente como aliado incondicional del gobierno federal ("cuenta y seguirá contando con las maestras y maestros").', style='List Bullet')

doc.add_paragraph()
h2 = doc.add_paragraph()
r_h2 = h2.add_run('2. ¿Nos perjudica este video en nuestra propuesta (PIP)?')
r_h2.font.size = Pt(14)
r_h2.bold = True

p2 = doc.add_paragraph()
p2.add_run('Sí y No. Es un arma de doble filo y todo depende de quién gane la narrativa:')

p_risk = doc.add_paragraph()
p_risk.add_run('EL RIESGO (Cómo nos perjudica): ').bold = True
p_risk.add_run('El SNTE nacional y el Gobierno Federal ya pactaron que habrá una reforma ("un nuevo sistema"). El peligro es que el SNTE negociará esta reforma a puerta cerrada y presentará una solución "parche" (quizás un ligero aumento en las UMAs o un esquema híbrido débil) y se colgarán la medalla. Si logran esto, la PIP del Magisterio Sonorense quedará marginada como una "propuesta disidente" frente a una "reforma ya lograda por la vía institucional".')

p_opp = doc.add_paragraph()
p_opp.add_run('LA OPORTUNIDAD (Cómo lo usamos a favor): ').bold = True
p_opp.add_run('El líder nacional acaba de darle la razón al Magisterio Sonorense a nivel nacional. Al mencionar que se necesita un nuevo sistema con "viabilidad presupuestaria de largo plazo", está describiendo exactamente la esencia técnica de la PIP. Ellos tienen la voluntad política, pero nosotros tenemos el cálculo matemático real. ')

doc.add_paragraph()
h3 = doc.add_paragraph()
r_h3 = h3.add_run('3. Acción Táctica Inmediata (Contención y Contrataque)')
r_h3.font.size = Pt(14)
r_h3.bold = True

doc.add_paragraph('El Magisterio Sonorense NO debe atacar este video diciendo que es mentira. Debe "subirse a la ola" y apropiarse del mensaje con la siguiente narrativa en los grupos de WhatsApp:', style='List Bullet')

p_quote = doc.add_paragraph()
p_quote.add_run('Mensaje Sugerido para la Base:').bold = True
p_quote.add_run(' "Compañeros, el líder nacional del SNTE acaba de confirmar frente a la Presidenta lo que el MS lleva meses diciendo: Urge un NUEVO sistema de pensiones con viabilidad presupuestaria. El SNTE lo dice de dientes para afuera, pero el Magisterio Sonorense ya lo tiene estructurado matemáticamente: La Pensión Intergeneracional Protegida (PIP). No permitiremos que el SNTE negocie migajas a puerta cerrada. La PIP es el camino para ese nuevo sistema que hoy reconocen."')

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Inteligencia_Video_SNTE_Nacional.docx'
doc.save(output_path)
print(f"Report saved to {output_path}")
