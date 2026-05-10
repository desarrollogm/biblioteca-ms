from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, encabezados, filas, hdr_color='1F3A6E'):
    table = doc.add_table(rows=len(filas)+1, cols=len(encabezados))
    table.style = 'Table Grid'
    for i, h in enumerate(encabezados):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True; run.font.name = 'Arial'; run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, hdr_color)
    for i, fila in enumerate(filas):
        fill = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
        for j, texto in enumerate(fila):
            cell = table.rows[i+1].cells[j]
            cell.text = texto
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Arial'; run.font.size = Pt(11)
            shade_cell(cell, fill)
    return table

def seccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

def subseccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True

def parrafo(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12)

def parrafo_bold(doc, etiqueta, texto):
    p = doc.add_paragraph()
    r1 = p.add_run(etiqueta); r1.font.name = 'Arial'; r1.font.size = Pt(12); r1.bold = True
    r2 = p.add_run(texto);    r2.font.name = 'Arial'; r2.font.size = Pt(12)

def parrafo_nota(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def bala(doc, texto, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix); r1.font.name = 'Arial'; r1.font.size = Pt(12); r1.bold = True
        r2 = p.add_run(texto);       r2.font.name = 'Arial'; r2.font.size = Pt(12)
    else:
        r = p.add_run(texto); r.font.name = 'Arial'; r.font.size = Pt(12)

def espacio(doc):
    doc.add_paragraph()

# ==============================================================================
# PORTADA
# ==============================================================================
from docx.enum.text import WD_ALIGN_PARAGRAPH

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ESTRATEGIA OPERATIVA DE ENTREGA')
r.font.name = 'Arial'; r.font.size = Pt(18); r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Gira Presidencial Sonora — 8 al 10 de Mayo de 2026')
r.font.name = 'Arial'; r.font.size = Pt(14); r.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

espacio(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Como entregar el Resumen Ejecutivo a la Presidenta Claudia Sheinbaum Pardo')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True

espacio(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Para uso exclusivo del CORE del Magisterio Sonorense')
r.font.name = 'Arial'; r.font.size = Pt(11); r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

espacio(doc)
parrafo_bold(doc, 'Clasificacion: ', 'Estrategia operativa — uso interno MS')
parrafo_bold(doc, 'Fecha: ', '4 de mayo de 2026')
parrafo_bold(doc, 'Ventana de ejecucion: ', 'Viernes 8 — Domingo 10 de mayo de 2026')

doc.add_page_break()

# ==============================================================================
# PARTE I — CONTEXTO Y OPORTUNIDAD
# ==============================================================================
seccion(doc, 'PARTE I — EL CONTEXTO: POR QUE AHORA ES EL MOMENTO')
espacio(doc)
parrafo(doc,
    'La Presidenta Claudia Sheinbaum realizara una gira de trabajo a Sonora del 8 al 10 de mayo de 2026. '
    'Es su 5a visita al estado. La agenda contempla inauguraciones hospitalarias en Navojoa, '
    'la colocacion de la primera piedra del Hospital de Ures y el seguimiento al Plan de Justicia del Rio Sonora.')
espacio(doc)
parrafo(doc,
    'Esta gira representa la ventana tactica mas importante para el Magisterio Sonorense porque:'
)
bala(doc, 'La Presidenta estara fisicamente en Sonora — sin intermediarios entre CDMX y la base.')
bala(doc, 'El contexto de los actos (salud, justicia, infraestructura) crea un ambiente receptivo a propuestas sociales.')
bala(doc, 'El recorrido Sur-Norte pasa por zonas con alta concentracion de maestros federalizados.')
bala(doc, 'El timing es perfecto: dias despues de la reunion CNTE-Presidenta del 11 de mayo, el MS puede posicionarse como la alternativa tecnica.')
espacio(doc)
parrafo_nota(doc,
    'LECTURA CLAVE: Sheinbaum declaro que la gira NO esta relacionada con coyuntura politica. '
    'Eso es una VENTAJA para el MS: el documento tecnico que entregamos NO es coyuntural — es una propuesta '
    'estructural que encaja perfectamente con el perfil "solucion tecnica, no presion politica" que el MS tiene construido.')

doc.add_page_break()

# ==============================================================================
# PARTE II — LA AGENDA PRESIDENCIAL
# ==============================================================================
seccion(doc, 'PARTE II — AGENDA PRESIDENCIAL: LO CONFIRMADO')
espacio(doc)
add_table(doc,
    ['Dia', 'Lugar probable', 'Evento', 'Nivel de acceso para el MS'],
    [
        ['Viernes 8 mayo', 'Navojoa', 'Inauguracion Hospital IMSS Navojoa', 'MEDIO — zona de medios accesible'],
        ['Sabado 9 mayo', 'Ures', 'Primera piedra Hospital Regional + Informe Rio Sonora', 'ALTO — municipio pequeno, acto abierto'],
        ['Domingo 10 mayo', 'Hermosillo', 'Posible acto de cierre o retorno CDMX', 'BAJO — protocolo de salida muy cerrado'],
    ]
)
espacio(doc)
parrafo_nota(doc,
    'NOTA: El itinerario hora a hora no ha sido publicado. '
    'Se conocera 24-48 horas antes. El equipo de inteligencia MS debe monitorear: '
    'El Imparcial (elimparcial.com), Proyecto Puente (proyectopuente.com.mx) y Tribuna de Sonora para actualizaciones.')

doc.add_page_break()

# ==============================================================================
# PARTE III — PLAN DE 3 ESCENARIOS
# ==============================================================================
seccion(doc, 'PARTE III — PLAN DE OPERACION: LOS 3 ESCENARIOS')
espacio(doc)

# ESCENARIO A
p = doc.add_paragraph()
r = p.add_run('ESCENARIO A — VIA LEGISLATIVA (Prioridad 1 — Ejecutar HOY)')
r.font.name = 'Arial'; r.font.size = Pt(13); r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x7A, 0x3A)  # verde

espacio(doc)
add_table(doc,
    ['Paso', 'Accion', 'Plazo'],
    [
        ['1', 'Identificar al Senador o Diputado Federal de Morena-Sonora mas afin a la agenda social educativa', 'HOY — 4 mayo'],
        ['2', 'Contactarlo con la peticion: "Queremos que usted lleve el Resumen Ejecutivo de la PIP directamente a la Presidenta durante su gira en Sonora"', 'HOY o manana 5 mayo'],
        ['3', 'Preparar 3 copias del Resumen Ejecutivo selladas: una para el legislador, una para el equipo de la Presidenta, una para SEP', 'Antes del 7 mayo'],
        ['4', 'Ofrecerle al legislador reconocimiento publico como intermediario de la causa', 'En el acto de entrega'],
    ]
)
espacio(doc)
parrafo_nota(doc, 'POR QUE FUNCIONA: Un legislador federal tiene acceso directo a la comitiva presidencial. '
    'El costo politico para el legislador es CERO — entrega un documento tecnico propositivo, no un pliego de exigencias. '
    'El beneficio para el legislador es alto: ser visible como enlace de los maestros de Sonora.')

espacio(doc)

# ESCENARIO B
p = doc.add_paragraph()
r = p.add_run('ESCENARIO B — PRESENCIA DIRECTA EN URES (Prioridad 2 — Alto Impacto)')
r.font.name = 'Arial'; r.font.size = Pt(13); r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)  # naranja

espacio(doc)
subseccion(doc, 'Perfil de la comision (5 personas maximo):')
add_table(doc,
    ['Rol', 'Perfil ideal', 'Funcion'],
    [
        ['Vocero/Vocera principal', 'Maestra o maestro con 20+ anos de servicio, articulado/a', 'Hace el contacto verbal con la Presidenta o su equipo'],
        ['Portador del documento', 'Cualquier miembro del equipo', 'Lleva el sobre sellado con el Resumen Ejecutivo'],
        ['Documentalista', 'Persona con telefono con buena camara', 'Graba el momento del acercamiento para redes'],
        ['Respaldo 1 y 2', '2 maestros adicionales como acompanantes', 'Presencia visual, sin hablar'],
    ]
)
espacio(doc)

subseccion(doc, 'Protocolo de acercamiento en Ures:')
add_table(doc,
    ['Tiempo', 'Accion'],
    [
        ['3 horas antes del acto', 'Llegar al lugar del evento. Identificar la zona de medios y el acceso principal'],
        ['2 horas antes', 'Posicionarse en la zona donde los medios instalan camaras — es el punto de mayor contacto con la comitiva'],
        ['Durante el acto', 'NO interrumpir. Observar el protocolo. Identificar quien es el jefe de asistentes de la Presidenta'],
        ['Al finalizar el acto', 'Momento de oportunidad: la comitiva se desplaza al vehiculo. Son 30-60 segundos de acceso'],
        ['El guion exacto', '"Presidenta, somos maestros federalizados de Sonora. Tenemos una propuesta tecnica para el ISSSTE ya presentada en el Congreso. Nos permite entregarle este documento?" — NADA MAS.'],
    ]
)
espacio(doc)
parrafo_nota(doc,
    'SI NO HAY ACCESO DIRECTO A LA PRESIDENTA: Entregar el sobre a cualquier miembro visible del Estado Mayor Presidencial (EMP) '
    'con la siguiente frase: "Por favor haga llegar este documento tecnico a la Presidenta. Es una propuesta '
    'del Magisterio Sonorense sobre pensiones ISSSTE." El EMP tiene protocolo para recibir y canalizar documentos ciudadanos.')

espacio(doc)

# ESCENARIO C
p = doc.add_paragraph()
r = p.add_run('ESCENARIO C — CONFERENCIA DE PRENSA EN HERMOSILLO (Respaldo garantizado)')
r.font.name = 'Arial'; r.font.size = Pt(13); r.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

espacio(doc)
parrafo(doc,
    'Independientemente de si se logra la entrega directa (A o B), el Escenario C debe ejecutarse siempre '
    'como amplificador mediatico:')
espacio(doc)
add_table(doc,
    ['Paso', 'Accion', 'Canal'],
    [
        ['1', 'Convocar a medios el mismo dia de la gira (8-10 mayo) en Hermosillo', 'WhatsApp a periodistas de El Imparcial, Proyecto Puente, Tribuna de Sonora'],
        ['2', 'Comunicado oficial: "El Magisterio Sonorense entrega formalmente propuesta tecnica a la Presidenta durante su visita a Sonora"', 'Boletín de prensa en PDF'],
        ['3', 'Distribuir el Resumen Ejecutivo a los periodistas que cubren la gira presidencial', 'Email + WhatsApp directo a reporteros acreditados'],
        ['4', 'Publicar en redes sociales con etiquetas @Claudiashein @SEP_mx @Mario_Delgado_', 'Twitter/X, Facebook, Instagram del MS'],
    ]
)

doc.add_page_break()

# ==============================================================================
# PARTE IV — CRONOGRAMA
# ==============================================================================
seccion(doc, 'PARTE IV — CRONOGRAMA DE ACCION: PROXIMAS 96 HORAS')
espacio(doc)
add_table(doc,
    ['Cuando', 'Que hacer', 'Quien', 'Prioridad'],
    [
        ['HOY — 4 mayo (noche)', 'Identificar legislador federal Morena-Sonora para Escenario A', 'Liderazgo MS', 'CRITICA'],
        ['HOY — 4 mayo (noche)', 'Imprimir 5 copias del Resumen Ejecutivo (documento separado), selladas en sobre formal', 'Equipo tecnico MS', 'CRITICA'],
        ['5 mayo — manana', 'Llamar al legislador federal y presentar la solicitud', 'Liderazgo MS', 'URGENTE'],
        ['5 mayo — manana', 'Enviar Resumen Ejecutivo a Palacio Nacional por correo certificado', 'Liderazgo MS', 'URGENTE'],
        ['5 mayo — tarde', 'Seleccionar y confirmar la comision de 5 maestros para Ures', 'Coordinadores zona centro', 'URGENTE'],
        ['6-7 mayo', 'Monitorear agenda presidencial en medios locales hora a hora', 'Equipo inteligencia MS', 'IMPORTANTE'],
        ['7 mayo — noche', 'Briefing final: ensayar el guion de 30 segundos, confirmar roles', 'Vocero + comision', 'IMPORTANTE'],
        ['8 mayo — madrugada', 'Salida de la comision hacia Ures (llegar 3 horas antes del acto)', 'Comision MS', 'EJECUCION'],
        ['8-10 mayo', 'Documentar todo para redes + notificar a medios + Escenario C', 'Toda la comision', 'EJECUCION'],
    ]
)

doc.add_page_break()

# ==============================================================================
# PARTE V — GUION COMPLETO
# ==============================================================================
seccion(doc, 'PARTE V — GUION COMPLETO PARA EL VOCERO')
espacio(doc)

subseccion(doc, 'Si hay acceso directo a la Presidenta (30 segundos):')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r = p.add_run(
    '"Presidenta Sheinbaum, buenos dias. Somos maestros federalizados del Magisterio Sonorense. '
    'Tenemos una propuesta tecnica para reformar el sistema de pensiones del ISSSTE que ya presentamos '
    'en la Camara de Diputados y en el Senado. Se trata de homologar el techo de las UMAs al precedente '
    'que ya tiene el IMSS. Somos 40,000 maestros en Sonora y venimos a proponer, no a exigir. '
    'Le pedimos nos permita entregarle este documento para su revision."')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

espacio(doc)

subseccion(doc, 'Si solo hay acceso al equipo de la Presidenta (15 segundos):')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r = p.add_run(
    '"Somos el Magisterio Sonorense. Tenemos una propuesta tecnica para el ISSSTE que ya esta en el Congreso. '
    'Le pedimos haga llegar este documento a la Presidenta. Es una propuesta, no una demanda."')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

espacio(doc)

subseccion(doc, 'Si los medios preguntan (para camara):')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
r = p.add_run(
    '"El Magisterio Sonorense viene a Sonora con la Presidenta porque tenemos una propuesta concreta '
    'y viable para mejorar las pensiones de los 40,000 maestros federalizados del estado. '
    'No venimos a bloquear ni a presionar — venimos a proponer. '
    'Nuestra propuesta, la PIP, ya esta en el Congreso. Solo pedimos que el gobierno federal nos escuche."')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

espacio(doc)

subseccion(doc, 'Lo que NO se debe decir (bajo ninguna circunstancia):')
bala(doc, 'No mencionar a la CNTE ni a otras organizaciones.')
bala(doc, 'No usar la palabra "exigir", "demandar" o "huelga".')
bala(doc, 'No confrontar a nadie del equipo presidencial si no reciben el documento.')
bala(doc, 'No hacer declaraciones politicas sobre partidos o funcionarios.')
bala(doc, 'No hablar de la reunion CNTE-Presidenta del 11 de mayo como referente.')

doc.add_page_break()

# ==============================================================================
# PARTE VI — EVALUACION DE RIESGO
# ==============================================================================
seccion(doc, 'PARTE VI — EVALUACION DE RIESGO Y PLANES DE CONTINGENCIA')
espacio(doc)
add_table(doc,
    ['Riesgo', 'Nivel', 'Plan de contingencia'],
    [
        ['EMP bloquea acceso total al evento', 'ALTO', 'Ejecutar Escenario C (prensa) + dejar el sobre en recepcion del evento con nota oficial'],
        ['La Presidenta no pasa por la zona de medios', 'MEDIO', 'Entregar al jefe de asistentes + comunicado a prensa desde afuera del evento'],
        ['La CNTE llega primero y satura el espacio', 'ALTO', 'El MS llega 3 horas antes, ocupa la zona de medios y tiene presencia visual diferenciada (portafolio formal, no pancarta)'],
        ['El legislador no puede gestionar el Escenario A', 'MEDIO', 'Buscar segunda opcion: regidor, diputado local, o secretario de gobierno estatal'],
        ['La gira se cancela o modifica el itinerario', 'BAJO', 'El Resumen Ejecutivo ya fue enviado a Palacio Nacional — el esfuerzo no se pierde'],
        ['Medios no cubren la entrega', 'MEDIO', 'El MS publica el video de la entrega en redes propias independientemente de los medios'],
    ]
)

espacio(doc)

# Conclusion
seccion(doc, 'CONCLUSION: LA REGLA DE ORO')
espacio(doc)
parrafo(doc,
    'La Seccion 7 de Chiapas logro una reunion con la Presidenta bloqueando un aeropuerto. '
    'Eso les funciono porque tienen masa critica y una historia de mesas previas. '
    'El MS tiene algo mas poderoso: una propuesta tecnica que el gobierno QUIERE escuchar '
    'porque no le cuesta politicamente — le da una solucion a un problema real.')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run(
    'REGLA DE ORO: El MS gana si el documento llega a manos de la Presidenta o de su equipo '
    'Y si los medios lo registran. Basta con eso para que el MS sea reconocido como interlocutor '
    'tecnico de primer nivel — sin un solo bloqueo, sin un solo paro.')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True; r.italic = True

espacio(doc)
doc.add_paragraph('_' * 65)
parrafo(doc, 'Estrategia elaborada para uso exclusivo del CORE del Magisterio Sonorense')
parrafo(doc, 'Fecha: 04/05/2026 | Ejecucion: 8-10 de mayo de 2026 | Sonora, Mexico')

# GUARDAR
out2 = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Estrategia_Entrega_Presidenta_MS_Mayo2026.docx'
doc.save(out2)
print(f'ESTRATEGIA guardada: {out2}')
