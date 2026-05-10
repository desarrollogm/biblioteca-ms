import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Page 1 - Portada
title = doc.add_heading('RESUMEN EJECUTIVO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('PROPUESTA DE PENSIÓN INTERGENERACIONAL PROTEGIDA (PIP)', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('\n' * 2)
p = doc.add_paragraph()
p.add_run('PARA: ').bold = True
p.add_run('Dra. Claudia Sheinbaum Pardo, Presidenta Constitucional de los Estados Unidos Mexicanos.\n')
p.add_run('DE: ').bold = True
p.add_run('Magisterio Sonorense (MS).\n')
p.add_run('FECHA: ').bold = True
p.add_run('8 de mayo de 2026.\n')
p.add_run('ASUNTO: ').bold = True
p.add_run('Propuesta técnica para la dignificación del retiro de los trabajadores del Estado (ISSSTE).')

doc.add_page_break()

# Page 2 - Diagnóstico
doc.add_heading('PÁGINA 1: EL DIAGNÓSTICO DEL PROBLEMA', level=2)
doc.add_paragraph('El sistema actual de Cuentas Individuales (Ley ISSSTE 2007) ha fallado en garantizar un retiro digno para el magisterio.')
doc.add_paragraph('• Tope Salarial Injusto: Mientras en el IMSS el tope de jubilación es de 25 UMAs (Acuerdo 064/2020), en el ISSSTE permanece topado a 10 UMAs.', style='List Bullet')
doc.add_paragraph('• Pérdida de Poder Adquisitivo: El reemplazo salarial actual es insuficiente, condenando a miles de maestros a la pobreza en su vejez.', style='List Bullet')
doc.add_paragraph('• Falta de Viabilidad: Las AFOREs priorizan el rendimiento privado sobre la seguridad social del trabajador.', style='List Bullet')

doc.add_heading('PÁGINA 2: LA SOLUCIÓN — LA PIP', level=2)
doc.add_paragraph('La Pensión Intergeneracional Protegida (PIP) es un modelo técnico-financiero diseñado para rescatar el sistema solidario sin comprometer las finanzas públicas.')
doc.add_paragraph('1. Carácter Solidario e Intergeneracional: Estado y trabajador se respaldan mutuamente.', style='List Number')
doc.add_paragraph('2. Fondo Blindado (FPIP): Creación del Fondo de Pensión Intergeneracional Protegida, con blindaje constitucional.', style='List Number')
doc.add_paragraph('3. Fortalecimiento de PENSIONISSSTE: Migración voluntaria de cuentas individuales desde AFOREs privadas.', style='List Number')
doc.add_paragraph('4. Respaldo Legislativo: La propuesta cuenta con el sustento técnico del Instituto Belisario Domínguez del Senado de la República.', style='List Number')

doc.add_page_break()

# Page 3 - Mecanismo
doc.add_heading('PÁGINA 3: MECANISMO FINANCIERO Y REQUISITOS', level=2)
doc.add_paragraph('• Aportación del Trabajador: 2% del salario base quincenal.', style='List Bullet')
doc.add_paragraph('• Aportación del Estado: $3.25 por cada peso aportado por el trabajador.', style='List Bullet')
doc.add_paragraph('• Aportación de Pensionados: 2% de su salario mensual.', style='List Bullet')
doc.add_paragraph('• Requisitos de Retiro: 28 años de servicio (mujeres), 30 años (hombres), con edad mínima de 55 años.', style='List Bullet')
doc.add_paragraph('• Justicia Social: Permite el reingreso al Décimo Transitorio para afectados de 2007.', style='List Bullet')

doc.add_heading('PÁGINA 4: LA PETICIÓN CONCRETA', level=2)
p = doc.add_paragraph()
r = p.add_run('Solicitamos respetuosamente la instalación inmediata de MESAS TÉCNICAS DE ANÁLISIS con la participación de:')
r.bold = True

doc.add_paragraph('1. Autoridades de primer nivel del ISSSTE.', style='List Bullet')
doc.add_paragraph('2. Representantes de la Secretaría de Educación Pública (SEP).', style='List Bullet')
doc.add_paragraph('3. Comisión Técnica del Magisterio Sonorense.', style='List Bullet')

doc.add_paragraph('\n' * 2)
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.add_run('ATENTAMENTE\n').bold = True
p_final.add_run('Magisterio Sonorense\n').bold = True
p_final.add_run('Por la defensa de nuestros derechos laborales. Jubilación digna para tod@s.')

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Resumen_Ejecutivo_MS_Presidenta_V2_Mayo2026.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
