import os
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Heading
heading = doc.add_heading('Reporte de Inteligencia: Convocatoria SNTE Oficial (Foto 2) - 1ro de Mayo 2026', 0)
heading.style.font.name = 'Arial'
heading.style.font.size = Pt(14)
heading.style.font.bold = True

p1 = doc.add_paragraph()
r1 = p1.add_run('Objetivo: ')
r1.bold = True
p1.add_run('Análisis de aforo total en la fotografía grupal de la cúpula oficial (Sección 28/54) durante la concentración del 1ro de Mayo en la Región Yaqui.')

doc.add_paragraph()
h1 = doc.add_paragraph()
r_h1 = h1.add_run('1. Conteo de Compañeros (Aforo Total en la Toma Abierta)')
r_h1.font.size = Pt(14)
r_h1.bold = True

p2 = doc.add_paragraph()
p2.add_run('Analizando la fotografía panorámica (subida por el Profe Raggio a las 12:44 p.m.), que representa la "foto del recuerdo" o concentración final del contingente:')

doc.add_paragraph('- La toma agrupa a todo el contingente extendido a lo ancho de la intersección vial.', style='List Bullet')
doc.add_paragraph('- Calculando por densidad de los bloques (izquierda, centro detrás de lonas, y derecha), se estima un aforo total de entre 90 y 110 personas en su punto máximo de concentración.', style='List Bullet')
doc.add_paragraph('- Efectivamente hay más gente que en la foto del contingente marchando (donde eran ~50), lo que indica que para la foto grupal juntaron a todos los rezagados y comitivas.', style='List Bullet')

doc.add_paragraph()
h2 = doc.add_paragraph()
r_h2 = h2.add_run('2. Distinción entre Sección 28 y Sección 54')
r_h2.font.size = Pt(14)
r_h2.bold = True

doc.add_paragraph('- Se mantiene la táctica de uniformidad total: No es posible distinguir orgánicamente entre secciones. Todos visten idéntico (camisa blanca y gorra naranja).', style='List Bullet')
doc.add_paragraph('- Las lonas visibles en el fondo mantienen la tipografía y los colores (rojo) característicos de la representación nacional y de la Sección 28.', style='List Bullet')
doc.add_paragraph('- Si hay presencia de la Sección 54, está completamente mimetizada y subsumida bajo el bloque de la 28. No llevan identidad propia.', style='List Bullet')

doc.add_paragraph()
h3 = doc.add_paragraph()
r_h3 = h3.add_run('3. Conclusión de Inteligencia Táctica')
r_h3.font.size = Pt(14)
r_h3.bold = True

p3 = doc.add_paragraph()
p3.add_run('Aunque el número sube a ~100 personas al juntarlos a todos para la foto estática, la conclusión estratégica sigue siendo la misma: Es un fracaso de convocatoria para el aparato oficial.')

p4 = doc.add_paragraph()
p4.add_run('El SNTE oficial tiene la obligación (y los recursos, comisionados y presupuesto) para movilizar a miles en una región como Cajeme. Lograr juntar apenas a 100 personas para la foto principal del Día del Trabajo demuestra el colapso de su arrastre en las bases. ')
p4.add_run('El contraste con el Magisterio Sonorense (MS) es abismal, confirmando que la base rechaza a la dirigencia oficial.')

p5 = doc.add_paragraph()
p5.add_run('\n[NOTA: JAD, DEBES PEGAR LA FOTO DEL PROFE RAGGIO MANUALMENTE AQUÍ ABAJO ANTES DE MANDAR EL DOCUMENTO]').bold = True

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Inteligencia_Convocatoria_SNTE_Foto2.docx'
doc.save(output_path)
print(f"SNTE Report 2 saved to {output_path}")
