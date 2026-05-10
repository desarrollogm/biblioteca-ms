from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==============================================================================
# DOCUMENTO 1: RESUMEN EJECUTIVO (para entregar a la Presidenta — 4 paginas)
# ==============================================================================

doc = Document()

# Margenes ajustados para documento ejecutivo
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(doc, encabezados, filas, hdr_color='1F3A6E'):
    table = doc.add_table(rows=len(filas)+1, cols=len(encabezados))
    table.style = 'Table Grid'
    for i, h in enumerate(encabezados):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True; run.font.name = 'Arial'; run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, hdr_color)
    for i, fila in enumerate(filas):
        fill = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
        for j, texto in enumerate(fila):
            cell = table.rows[i+1].cells[j]
            cell.text = texto
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Arial'; run.font.size = Pt(11)
            shade_cell(cell, fill)
    return table

def espacio(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)

# ── PAGINA 1: PORTADA ────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('Hermosillo, Sonora — 9 de mayo de 2026')
r.font.name = 'Arial'; r.font.size = Pt(10); r.font.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

espacio(doc, 4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('Para:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
r2 = p.add_run('  Presidenta Claudia Sheinbaum Pardo')
r2.font.name = 'Arial'; r2.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('De:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
r2 = p.add_run('  Magisterio Sonorense — Maestros Federalizados de Sonora')
r2.font.name = 'Arial'; r2.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('Asunto:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
r2 = p.add_run('  Propuesta Tecnica para la Reforma del Sistema de Pensiones ISSSTE')
r2.font.name = 'Arial'; r2.font.size = Pt(12)

espacio(doc, 8)

# Linea divisoria
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('_' * 70)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

espacio(doc, 8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('RESUMEN EJECUTIVO')
r.font.name = 'Arial'; r.font.size = Pt(20); r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Propuesta de Reforma al Techo Pensional del ISSSTE')
r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True
r.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Homologacion al Precedente IMSS — Acuerdo 064/2020')
r.font.name = 'Arial'; r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

espacio(doc, 8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Magisterio Sonorense')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Representacion de los Maestros Federalizados del Estado de Sonora')
r.font.name = 'Arial'; r.font.size = Pt(11); r.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

espacio(doc, 6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Mayo 2026')
r.font.name = 'Arial'; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ── PAGINA 2: EL PROBLEMA ────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('I. EL PROBLEMA: LA BRECHA PENSIONAL QUE AFECTA A LOS MAESTROS DE SONORA')
r.font.name = 'Arial'; r.font.size = Pt(14); r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

espacio(doc, 4)

p = doc.add_paragraph()
r = p.add_run('Los maestros federalizados de Sonora enfrentan una inequidad estructural en el calculo de sus pensiones: '
    'la Ley del ISSSTE de 2007 establece un techo maximo de 10 Unidades de Medida y Actualizacion (UMAs) para el '
    'calculo de pensiones de jubilacion. Este techo implica que sin importar los anos de servicio o el salario '
    'cotizado, la pension maxima que puede recibir un maestro federalizado esta artificialmente limitada.')
r.font.name = 'Arial'; r.font.size = Pt(12)

espacio(doc, 6)

# Tabla de datos duros
p = doc.add_paragraph()
r = p.add_run('Los tres datos que definen el problema:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True

espacio(doc, 3)

add_table(doc,
    ['', 'ISSSTE (Maestros Federalizados)', 'IMSS (Precedente ya existente)'],
    [
        ['Techo pensional', '10 UMAs = ~$27,400 pesos mensuales', '25 UMAs = ~$68,500 pesos mensuales'],
        ['Base legal', 'Ley ISSSTE 2007 — Articulo 10', 'Acuerdo IMSS 064/2020 — vigente'],
        ['Trabajadores afectados en Sonora', '~40,000 maestros federalizados', 'Ya resuelto para trabajadores IMSS'],
    ]
)

espacio(doc, 8)

# Caja de impacto
p = doc.add_paragraph()
r = p.add_run('IMPACTO CONCRETO: Un maestro con 30 anos de servicio y salario de $45,000 mensuales '
    'recibe una pension maxima de $27,400 pesos bajo el esquema actual. '
    'El mismo perfil bajo el esquema IMSS recibiria hasta $45,000 pesos. '
    'La diferencia: $17,600 pesos menos al mes por ser maestro y no trabajador del IMSS.')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

espacio(doc, 6)

p = doc.add_paragraph()
r = p.add_run('Esta inequidad no tiene justificacion tecnica ni constitucional. '
    'El propio gobierno federal ya resolvio el problema para los trabajadores del IMSS mediante el Acuerdo 064/2020. '
    'Los maestros federalizados, que son trabajadores al servicio del Estado con igual dignidad constitucional, '
    'no han recibido la misma atencion.')
r.font.name = 'Arial'; r.font.size = Pt(12)

doc.add_page_break()

# ── PAGINA 3: LA PROPUESTA ───────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('II. LA PROPUESTA: PENSION INTEGRADORA PROPORCIONAL (PIP)')
r.font.name = 'Arial'; r.font.size = Pt(14); r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

espacio(doc, 4)

p = doc.add_paragraph()
r = p.add_run('Que es la PIP:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
p2 = doc.add_paragraph()
r2 = p2.add_run('La Pension Integradora Proporcional (PIP) es un mecanismo de calculo que complementa '
    'la pension base del ISSSTE con los recursos acumulados en la cuenta individual del trabajador (Afore), '
    'garantizando que el monto total se acerque al ultimo salario percibido — sin topar artificialmente la pension '
    'a 10 UMAs.')
r2.font.name = 'Arial'; r2.font.size = Pt(12)

espacio(doc, 6)

p = doc.add_paragraph()
r = p.add_run('Los 5 pilares de la propuesta:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True

espacio(doc, 3)

add_table(doc,
    ['#', 'Pilar', 'Detalle'],
    [
        ['1', 'Homologacion del techo UMA', 'Elevar el techo de 10 a 25 UMAs, igualando el precedente IMSS Acuerdo 064/2020'],
        ['2', 'Integracion de la cuenta individual', 'La Afore del trabajador complementa la pension base de forma proporcional a los anos cotizados'],
        ['3', 'Viabilidad fiscal comprobada', 'La PIP no requiere nuevo gasto federal — usa recursos ya existentes en cuentas individuales'],
        ['4', 'Sin reforma constitucional', 'Se instrumenta mediante acuerdo administrativo, igual que el Acuerdo IMSS 064/2020'],
        ['5', 'UMA-Fronteriza para Sonora', 'Propuesta adicional: diferencial geografico para zonas fronterizas por mayor costo de vida'],
    ]
)

espacio(doc, 6)

p = doc.add_paragraph()
r = p.add_run('Estado legislativo actual:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
espacio(doc, 3)
add_table(doc,
    ['Instancia', 'Accion realizada', 'Fecha'],
    [
        ['Camara de Diputados', 'Presentacion formal de la propuesta PIP en San Lazaro', '2026'],
        ['Senado de la Republica', 'Presentacion de propuesta UMA-Fronteriza', '2026'],
        ['Antecedente legal', 'Jurisprudencia 2a./J. 164/2019 — SCJN (precedente IMSS)', 'Vigente'],
    ]
)

doc.add_page_break()

# ── PAGINA 4: LO QUE SE PIDE ─────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run('III. LO QUE SOLICITAMOS AL GOBIERNO FEDERAL')
r.font.name = 'Arial'; r.font.size = Pt(14); r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

espacio(doc, 4)

p = doc.add_paragraph()
r = p.add_run('Tres peticiones concretas y viables:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
espacio(doc, 3)

add_table(doc,
    ['#', 'Peticion', 'Mecanismo sugerido'],
    [
        ['1', 'Instalacion de una mesa tecnica permanente entre el Magisterio Sonorense, SEP e ISSSTE para revisar la viabilidad de homologar el techo de UMAs al precedente del IMSS',
         'Decreto presidencial o instruccion directa a SEP/ISSSTE'],
        ['2', 'Revision y apoyo a la iniciativa PIP ya presentada en el Congreso',
         'Aval de la SEP como dependencia coordinadora del sector educativo'],
        ['3', 'Reconocimiento del diferencial geografico de la zona fronteriza de Sonora mediante la UMA-Fronteriza',
         'Instruccion a ISSSTE para elaborar estudio actuarial de factibilidad'],
    ]
)

espacio(doc, 8)

# Declaracion formal
p = doc.add_paragraph()
r = p.add_run('Declaracion del Magisterio Sonorense:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
espacio(doc, 3)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.right_indent = Cm(1.0)
r = p.add_run(
    '"El Magisterio Sonorense no viene a exigir lo imposible. Venimos a presentar una propuesta tecnica '
    'que ya fue posible para los trabajadores del IMSS mediante el Acuerdo 064/2020. '
    'Pedimos para los maestros de Sonora la misma justicia que ya existe para otros trabajadores del Estado. '
    'Nuestra propuesta esta documentada, es fiscalmente viable y no requiere reformar la Constitucion. '
    'Solo requiere voluntad politica."')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

espacio(doc, 8)

# Contacto
p = doc.add_paragraph()
r = p.add_run('Contacto para seguimiento:')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
espacio(doc, 3)
add_table(doc,
    ['Campo', 'Dato'],
    [
        ['Organizacion', 'Magisterio Sonorense — Representacion de Maestros Federalizados'],
        ['Estado', 'Sonora, Mexico'],
        ['Documentacion en el Congreso', 'Propuesta PIP — Camara de Diputados y Senado, 2026'],
        ['Antecedente legal de sustento', 'Jurisprudencia SCJN 2a./J. 164/2019 + Acuerdo IMSS 064/2020'],
    ]
)

espacio(doc, 6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('_' * 60)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Magisterio Sonorense — Mayo 2026')
r.font.name = 'Arial'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"Por la dignidad pensional de todos los maestros de Sonora"')
r.font.name = 'Arial'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

# GUARDAR RESUMEN EJECUTIVO
out1 = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Resumen_Ejecutivo_MS_Presidenta_Mayo2026.docx'
doc.save(out1)
print(f'RESUMEN EJECUTIVO guardado: {out1}')
