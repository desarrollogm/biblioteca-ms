import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_doc_zona(titulo_doc, zonas_data, nombre_archivo, imagen_path=None):
    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    # Titulo Principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo_doc)
    r.font.name = 'Arial'
    r.font.size = Pt(16)
    r.font.bold = True
    
    doc.add_paragraph() # Espacio
    
    # Si hay imagen, intentar insertarla
    if imagen_path and os.path.exists(imagen_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            r = p.add_run()
            r.add_picture(imagen_path, width=Inches(4.5))
        except Exception as e:
            p.add_run(f"[Error insertando imagen: {e}]")
        doc.add_paragraph()
        
    for seccion in zonas_data:
        # Titulo Seccion (Arial 14)
        p = doc.add_paragraph()
        r = p.add_run(seccion['titulo'])
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        
        # Contenido (Arial 12)
        for parrafo in seccion['parrafos']:
            if parrafo.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                r = p.add_run(parrafo[2:])
            elif parrafo.startswith('**') and '**' in parrafo[2:]:
                # Manejo basico de negritas al inicio
                partes = parrafo.split('**')
                p = doc.add_paragraph()
                for i, txt in enumerate(partes):
                    if not txt: continue
                    r = p.add_run(txt)
                    r.font.name = 'Arial'
                    r.font.size = Pt(12)
                    if i % 2 != 0: # Impar significa que estaba entre **
                        r.font.bold = True
            else:
                p = doc.add_paragraph()
                r = p.add_run(parrafo)
                r.font.name = 'Arial'
                r.font.size = Pt(12)
        doc.add_paragraph() # Espacio final de seccion
        
    out_dir = r"C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones"
    doc.save(os.path.join(out_dir, nombre_archivo))


# ================= DATA ZONA 1 =================
z1_data = [
    {"titulo": "CAMPAÑA: UMA-FRONTERIZA", "parrafos": ["Para uso de Coordinadores: Nogales, Agua Prieta, San Luis Río Colorado, Sonoyta, Naco, Cananea, Altar."]},
    {"titulo": "1. TEXTOS PARA REDES SOCIALES (COPY)", "parrafos": [
        "**Opción A: Post largo para Facebook**",
        "MAESTRO DE LA FRONTERA — ESTO AFECTA TU RETIRO",
        "¿Sabías que un trabajador del IMSS puede jubilarse con hasta $89,155 pesos al mes (25 UMAs), mientras que tú, maestro de la frontera norte, estás topado en $35,662 (10 UMAs)?",
        "La misma ley. El mismo país. Pero tú vives en una ciudad donde todo cuesta en dólares y pagas renta en zona libre.",
        "El Magisterio Sonorense ya presentó en el Senado de la República la propuesta de UMA-Fronteriza: que el tope de pensión se ajuste al costo de vida real de nuestra frontera.",
        "**Opción B: WhatsApp / Instagram**",
        "Gasto en dólares. Pago renta en zona libre. Pero el ISSSTE me jubila como si viviera en el sur.",
        "El MS tiene mi propuesta: UMA-FRONTERIZA."
    ]},
    {"titulo": "2. EMBUDO DE WHATSAPP", "parrafos": [
        "**Día 1 - El Gancho**",
        "Hola compañero/a. ¿Te parece justo que el gobierno nos pague el salario mínimo con tarifa de frontera, pero a la hora de jubilarnos el ISSSTE nos tope la pensión en 10 UMAs igual que en cualquier otro estado? El Magisterio Sonorense tiene un documento legal en el Senado para cambiar esto.",
        "**Día 3 - El Dato Duro**",
        "Te paso el dato exacto: Una pensión topada del IMSS hoy es de $89,155 al mes. Nuestra pensión topada del ISSSTE es de $35,662. Son $53,493 pesos que tú y tu familia pierden CADA MES de por vida.",
        "**Día 7 - El Cierre**",
        "La próxima semana vamos a mandar la primera lista de maestros fronterizos que respaldan la propuesta de la UMA-Fronteriza. ¿Te anoto en la lista de apoyo?"
    ]},
    {"titulo": "3. GUION PARA ASAMBLEAS O REUNIONES", "parrafos": [
        "Compañeros, buenos días. No vengo a hablarles de política nacional. Vengo a hablar del precio del supermercado aquí en Nogales. Todos sabemos que vivir en esta frontera es más caro.",
        "El gobierno federal lo sabe perfectamente, por eso existe la Zona Libre de la Frontera Norte. Pero parece que a la hora de jubilarnos se les olvida. Llegamos al final de nuestra carrera y nos dicen: 'tu tope son 10 UMAs'.",
        "La CNTE les va a decir que hagan huelga para pedir el 100% de aumento. El SNTE les va a decir que se conformen. El Magisterio Sonorense vino a traerles una propuesta real, que ya está en la Comisión del Senado: La UMA-Fronteriza."
    ]}
]

# ================= DATA ZONA 2 =================
z2_data = [
    {"titulo": "CAMPAÑA: REZONIFICACIÓN METROPOLITANA", "parrafos": ["Para uso de Coordinadores: Hermosillo, Bahía de Kino, La Colorada."]},
    {"titulo": "1. TEXTOS PARA REDES SOCIALES (COPY)", "parrafos": [
        "**Opción A: Post largo para Facebook**",
        "HERMOSILLO CRECIÓ, TU SUELDO NO.",
        "¿Sabías que Hermosillo es hoy la ciudad más cara de Sonora (después de la frontera), pero la SEP te sigue pagando con un tabulador de zona económica de hace más de 15 años?",
        "Una renta promedio no baja de los $8,000 pesos mensuales. La gasolina, la luz y los servicios están altísimos. Y cuando te jubiles, el ISSSTE te va a topar la pensión en $35,662 (10 UMAs).",
        "El Magisterio Sonorense exige que tu sueldo y tu pensión reconozcan el costo de vida real de la capital sonorense mediante la Rezonificación y la PIP.",
        "**Opción B: WhatsApp / Instagram**",
        "Tu ciudad creció y se modernizó. Tu tabulador SEP se quedó en el pasado. Pensión tope: $35,662 vs Rentas de $10,000. El MS tiene tu propuesta: REZONIFICACIÓN METROPOLITANA + PIP."
    ]},
    {"titulo": "2. EMBUDO DE WHATSAPP", "parrafos": [
        "**Día 1 - El Gancho**",
        "Hola compañero/a. Una pregunta honesta: ¿Cuánto gastas tú en gasolina y servicios básicos aquí en Hermosillo comparado con hace 5 años? La ciudad está carísima, pero nuestros tabuladores siguen igual. El MS exige la Rezonificación.",
        "**Día 3 - El Dato Duro**",
        "Aquí están los números: Hoy una pensión topada del ISSSTE es de $35,662 al mes. ¿Cuánto te queda libre después de pagar vivienda y servicios en Hermosillo? El MS está peleando subir el tope a 25 UMAs.",
        "**Día 7 - El Cierre**",
        "La próxima semana los coordinadores del MS entregarán la lista de maestros de la capital que respaldan la Rezonificación. ¿Te anoto en la lista de Hermosillo? Solo responde con un 'Sí apoyo'."
    ]},
    {"titulo": "3. GUION PARA ASAMBLEAS O REUNIONES", "parrafos": [
        "Compañeros, todos aquí batallamos con el tráfico, pagamos los recibos de la luz más altos en verano y vemos cómo cada día construyen plazas más caras. Hermosillo ya es una metrópoli. Pero para la SEP, seguimos cobrando como pueblo.",
        "El sistema está diseñado para que absorbamos el costo de la inflación urbana sin compensación.",
        "El Magisterio Sonorense elaboró la propuesta de Rezonificación Metropolitana para exigir al gobierno federal la actualización de nuestros tabuladores a costo de vida real. Apoyen la propuesta técnica, no la grilla."
    ]}
]

# ================= DATA ZONA 3 =================
z3_data = [
    {"titulo": "CAMPAÑA: QUINQUENIOS AL SUELDO BASE", "parrafos": ["Para uso de Coordinadores: Guaymas, Empalme, San Carlos."]},
    {"titulo": "1. TEXTOS PARA REDES SOCIALES (COPY)", "parrafos": [
        "**Opción A: Post largo para Facebook**",
        "MAESTRO DE GUAYMAS Y EMPALME: REVISA TU TALÓN DE CHEQUE",
        "¿Te has dado cuenta de que tus Quinquenios (Concepto 07) no están impactando tu sueldo base como marca la ley?",
        "Llevas años cotizando, pero el sistema tiene desindexada esta prima de tu sueldo tabular. Ganas menos hoy y tu jubilación será menor.",
        "A esto súmale que el costo de vida en la zona portuaria e industrial es altísimo. El MS está exigiendo Rezonificación Portuaria y la corrección inmediata del Concepto 07.",
        "**Opción B: WhatsApp / Instagram**",
        "Años de servicio. Mismos quinquenios mal calculados. Revisa tu Concepto 07 en tu talón. El MS tiene la ruta legal para compactarlo a tu sueldo base."
    ]},
    {"titulo": "2. EMBUDO DE WHATSAPP", "parrafos": [
        "**Día 1 - El Gancho**",
        "Hola compañero/a. Pregunta técnica rápida: ¿alguna vez te has fijado cómo viene calculado el 'Concepto 07' (Quinquenios) en tu talón? El MS encontró que los están calculando mal.",
        "**Día 3 - El Dato Duro**",
        "Te explico rápido cómo nos pega: Al no estar en el sueldo base, no solo cobras menos en la quincena, sino que al jubilarte, te van a calcular la pensión sobre una base mucho menor.",
        "**Día 7 - El Cierre**",
        "La próxima semana mandaremos una lista de maestros que respaldan la exigencia de corregir los Quinquenios. ¿Te anoto en la lista de respaldo?"
    ]},
    {"titulo": "3. GUION PARA ASAMBLEAS O REUNIONES", "parrafos": [
        "Compañeros, todos sabemos que nuestro sueldo cada vez alcanza para menos. Saquen su teléfono y revisen su último talón del FONE. Busquen el 'Concepto 07'.",
        "Ese concepto son sus quinquenios, pero lo tienen desindexado del sueldo base. Nos están robando dinero quincena tras quincena.",
        "El Magisterio Sonorense tiene la investigación legal que demuestra esta falla y la ruta para compactar el Concepto 07 al sueldo base. Defiendan su antigüedad."
    ]}
]

# ================= DATA ZONA 4 =================
z4_data = [
    {"titulo": "CAMPAÑA: LA GENERACIÓN AFECTADA (PIP + CONCEPTO 07)", "parrafos": ["Para uso de Coordinadores: Ciudad Obregón, Navojoa, Huatabampo, Álamos."]},
    {"titulo": "1. TEXTOS PARA REDES SOCIALES (COPY)", "parrafos": [
        "**Opción A: Post largo para Facebook**",
        "A LA GENERACIÓN QUE CONSTRUYÓ LAS ESCUELAS DEL SUR",
        "Llevas 25, 30 años de servicio. Has dejado tu vista y tu salud en el aula. ¿Y qué te dice el ISSSTE al jubilarte? Tu pensión está topada en 10 UMAs ($35,662 pesos al mes).",
        "Mientras tanto, un trabajador del IMSS con los mismos años se lleva hasta 25 UMAs ($89,155).",
        "El Magisterio Sonorense ya fue al Senado a entregar la propuesta técnica para corregir tus quinquenios y para aprobar la PIP (Pensión Intergeneracional Protegida) que rompe este techo.",
        "**Opción B: WhatsApp / Instagram**",
        "30 años de servicio. Jubilación topada en 10 UMAs. Tus quinquenios mal calculados. El MS tiene la propuesta legal: PIP + Concepto 07. Defiende tu jubilación."
    ]},
    {"titulo": "2. EMBUDO DE WHATSAPP", "parrafos": [
        "**Día 1 - El Gancho**",
        "Hola maestro/a. Le escribo porque usted es de la generación que levantó las escuelas en el sur. Da coraje ver cómo el ISSSTE le topa la pensión en 10 UMAs. El MS ya metió un documento al Senado para cambiar esto.",
        "**Día 3 - El Dato Duro**",
        "Maestro/a, además nos hemos dado cuenta de que durante toda su carrera le han calculado mal los Quinquenios (Concepto 07). El MS ya tiene la ruta jurídica para exigir que se integre a su base.",
        "**Día 7 - El Cierre**",
        "La próxima semana mandamos la lista oficial de Obregón y Navojoa con los docentes que respaldan la PIP. Me daría orgullo que su nombre estuviera en esa lista. ¿Lo anoto con un 'Sí apoyo'?"
    ]},
    {"titulo": "3. GUION PARA ASAMBLEAS O REUNIONES", "parrafos": [
        "Compañeros, veo aquí a maestros que me dieron clases, con 25 o 30 años de servicio. Ustedes son la generación a la que más le están quitando.",
        "La Ley del ISSSTE de 2007 (Décimo Transitorio) les puso un techo de cristal: 10 UMAs. Y además sus quinquenios se los han calculado por debajo del sueldo base.",
        "La CNTE pide marchar. Nosotros fuimos al Senado. El Magisterio Sonorense construyó el plan de reemplazo: la PIP. Defendamos su retiro con inteligencia, no con desgaste."
    ]}
]

# ================= DATA ZONA 5 =================
z5_data = [
    {"titulo": "CAMPAÑA: JUSTICIA EN LA SIERRA (PLAZAS + USICAMM)", "parrafos": ["Para uso de Coordinadores: Sahuaripa, Yécora, Zona Serrana, Comunidades Yaqui, Mayo y Seri."]},
    {"titulo": "1. TEXTOS PARA REDES SOCIALES (COPY)", "parrafos": [
        "**Opción A: Post largo para Facebook**",
        "EL MAESTRO DE LA SIERRA NO ES UN MAESTRO DE SEGUNDA",
        "Trabajas donde el asfalto termina. ¿Y cómo te responde el sistema educativo? Con plazas sin regularizar y con un USICAMM que te evalúa desde un escritorio refrigerado en Hermosillo.",
        "Exámenes estandarizados que no entienden ni respetan la realidad bilingüe, bicultural y rural de Sonora.",
        "El Magisterio Sonorense exige en el Senado la Regularización Inmediata de Plazas y la obligatoriedad de un USICAMM con Bilateralidad y Pertinencia Cultural.",
        "**Opción B: WhatsApp / Instagram**",
        "Aulas de cartón. Caminos de terracería. Y el USICAMM te evalúa como si estuvieras en Polanco. Exigimos plazas regularizadas. El MS está luchando por la sierra."
    ]},
    {"titulo": "2. EMBUDO DE WHATSAPP", "parrafos": [
        "**Día 1 - El Gancho**",
        "Hola compañero/a. Sé lo difícil que es la labor en la sierra. En el MS estamos hartos de que a los maestros rurales e indígenas se les trate con plazas temporales y evaluaciones desfasadas.",
        "**Día 3 - El Dato Duro**",
        "El problema es estructural: La ley asume que todos trabajan en las mismas condiciones, y eso es falso. El MS ya puso este tema en el Senado exigiendo regularización y bilateralidad.",
        "**Día 7 - El Cierre**",
        "Vamos a entregar firmas específicas de maestros rurales e indígenas. Queremos que en CDMX vean que la sierra existe y exige sus derechos. ¿Me das luz verde para anotar tu nombre?"
    ]},
    {"titulo": "3. GUION PARA ASAMBLEAS O REUNIONES", "parrafos": [
        "Compañeros, gracias por recibirme. El Magisterio Sonorense sabe que ustedes cubren las ausencias del Estado en las comunidades.",
        "Y sin embargo, los tienen con plazas irregulares y les aplican reglas de USICAMM pensadas para la capital. Exámenes que ignoran que ustedes enseñan en alta marginación.",
        "Nosotros llevamos su voz directamente a las mesas legislativas. Necesitamos su aval para que los senadores entiendan que esta demanda viene directamente desde la sierra."
    ]}
]

base_path = r"C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones"
img1 = os.path.join(base_path, "publicidad_zona1_frontera.png")
img2 = os.path.join(base_path, "publicidad_zona2_hermosillo.png")
img3 = os.path.join(base_path, "publicidad_zona3_guaymas.png")
img4 = os.path.join(base_path, "publicidad_zona4_sur.png")
img5 = os.path.join(base_path, "publicidad_zona5_sierra.png")

crear_doc_zona("CUADERNO ZONA 1: FRONTERA NORTE", z1_data, "Publicidad_Zona1_Frontera.docx", img1)
crear_doc_zona("CUADERNO ZONA 2: CAPITAL HERMOSILLO", z2_data, "Publicidad_Zona2_Hermosillo.docx", img2)
crear_doc_zona("CUADERNO ZONA 3: COSTA GUAYMAS-EMPALME", z3_data, "Publicidad_Zona3_Guaymas.docx", img3)
crear_doc_zona("CUADERNO ZONA 4: SUR OBREGÓN-NAVOJOA", z4_data, "Publicidad_Zona4_Sur.docx", img4)
crear_doc_zona("CUADERNO ZONA 5: SIERRA E INDÍGENA", z5_data, "Publicidad_Zona5_Sierra.docx", img5)

print("Todas las carpetas de publicidad generadas en formato DOCX con Arial 12/14 y las imagenes integradas.")
