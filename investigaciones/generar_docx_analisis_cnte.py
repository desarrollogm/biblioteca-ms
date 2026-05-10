import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_reporte_cnte():
    doc = Document()
    
    # Configuración de estilos base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Título (Arial 14, Bold)
    titulo_texto = 'REPORTE DE INTELIGENCIA: LA ÉLITE DE LA CNTE EN EL PODER (2018-2024)'
    heading = doc.add_heading('', level=0)
    run = heading.add_run(titulo_texto)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Documento Reservado para el CORE del MS')
    run.italic = True
    run.font.size = Pt(11)
    
    doc.add_paragraph('-' * 50)
    
    # Sección 1: RESUMEN EJECUTIVO
    h1 = doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
    for run in h1.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    resumen = (
        'El presente reporte detalla la integración de la otrora "cúpula combativa" de la CNTE a las estructuras de '
        'poder del gobierno federal (4T). Se identifica un patrón de cooptación institucional donde los liderazgos de '
        'las secciones más radicales (22 Oaxaca, 18 Michoacán, 9 CDMX) han permutado la movilización de base por '
        'escaños legislativos y direcciones administrativas.\n\n'
        'Este análisis sirve como herramienta discursiva para demostrar que la CNTE no es una oposición real, sino una '
        'cantera de cuadros para el oficialismo, lo que explica la parálisis en la solución de demandas históricas como '
        'la abrogación total de la Ley del USICAMM.'
    )
    doc.add_paragraph(resumen)
    
    # Sección 2: PERFILES CLAVE
    h2 = doc.add_heading('2. PERFILES CLAVE: "DE LA MARCHA AL ESCAÑO"', level=1)
    for run in h2.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    perfiles = [
        ('A. Azael Santiago Chepi (Oaxaca - Sección 22)', 
         '• Origen: Secretario General de la Sección 22 (2008-2012).\n'
         '• Posición Actual: Diputado Federal por Morena (reelecto).\n'
         '• Punto de Quiebre: Durante su gestión en la Comisión de Educación, fue uno de los principales operadores para la Reforma Educativa de 2019. '
         'La base en Oaxaca lo acusa de suavizar las demandas de abrogación total para alinearse con Palacio Nacional.'),
        
        ('B. Adela Piña Bernal (CDMX - Sección 9)', 
         '• Origen: Representante sindical de la Sección 9.\n'
         '• Trayectoria: Diputada Federal (Presidenta de la Comisión de Educación).\n'
         '• El "Pecado" Político: Fue nombrada Titular de la USICAMM. Es la responsable directa de la ejecución de los procesos de admisión y '
         'promoción que hoy tienen asfixiado al magisterio. Es el ejemplo más claro de un ex-CNTE administrando el sistema que juraron destruir.'),
        
        ('C. Raúl Morón Orozco (Michoacán - Sección 18)', 
         '• Origen: Líder histórico de la Sección 18.\n'
         '• Trayectoria: Senador de la República, Alcalde de Morelia.\n'
         '• Análisis: Morón representa el ala negociadora que utilizó la fuerza de la CNTE Michoacán para catapultar su carrera política personal. '
         'Su lealtad hoy es 100% partidista, subordinando los intereses magisteriales a la disciplina de Morena.'),
        
        ('D. Irán Santiago Manuel (Oaxaca - Sección 22)', 
         '• Origen: Secretario de Pagos de la Sección 22.\n'
         '• Posición Actual: Diputado Federal Morena.\n'
         '• Análisis: Conocido como el "operador financiero" de la 22. Su llegada al Congreso aseguró que los recursos y plazas en Oaxaca '
         'siguieran bajo un esquema de control que beneficia a su grupo político.'),

        ('E. Jorge Ángel Sibaja Mendoza (Oaxaca - Sección 22)', 
         '• Origen: Docente de la Sección 22.\n'
         '• Posición Actual: Diputado Federal Morena.\n'
         '• Escándalo de "Expulsión": Recientemente expulsado de mesas de negociación por la propia Sección 22, quienes lo señalaron de '
         '"usar al movimiento como trampolín". Es el ejemplo perfecto del "infiltrado" oficialista que ya no es bienvenido en las bases.'),

        ('F. Arcelia López Hernández (Oaxaca - Sección 22)', 
         '• Origen: Vínculo familiar (sobrina de Eloy López, ex-dirigente de la S22).\n'
         '• Posición Actual: Ex-Diputada Local Morena.\n'
         '• Análisis: Representa el nepotismo sindical. Su ascenso no fue por lucha, sino por herencia de poder dentro de la cúpula de la Sección 22.'),

        ('G. Antonio de Jesús Madriz Estrada (Michoacán - Sección 18)', 
         '• Origen: Magisterio Democrático (S18).\n'
         '• Trayectoria: Ex-Diputado Local, Director en Bienestar.\n'
         '• Uso del Poder: Ha sido denunciado por trabajadores de Bienestar por presunta coacción política. Es el salto de la "lucha social" al control partidista.'),

        ('H. Dolores Padierna Luna (CDMX)', 
         '• Origen: Fundadora histórica de la CNTE en los 70/80.\n'
         '• Rol Actual: Enlace Institucional de la SEP con el Poder Legislativo.\n'
         '• La "Madrina" del Pacto: Demuestra que la CNTE tiene "raíces" profundas en el mismo sistema que dicen combatir.')
    ]
    
    for titulo_p, contenido_p in perfiles:
        h_p = doc.add_heading(titulo_p, level=2)
        for run in h_p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_paragraph(contenido_p)
        
    # Sección 3: LA CONTRADICCIÓN ESTRATÉGICA
    h3 = doc.add_heading('3. LA CONTRADICCIÓN ESTRATÉGICA', level=1)
    for run in h3.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph(
        'Mientras la CNTE llama hoy a paros y huelgas nacionales exigiendo la abrogación de la ley de USICAMM, sus propios cuadros (y familiares de sus líderes) son quienes:'
    )
    doc.add_paragraph('1. Votaron a favor de la Ley General del Sistema para la Carrera en 2019.', style='List Number')
    doc.add_paragraph('2. Administran la USICAMM (Adela Piña).', style='List Number')
    doc.add_paragraph('3. Heredan las diputaciones por nepotismo (Arcelia López).', style='List Number')
    doc.add_paragraph('4. Operan desde la SEP para contener al movimiento (Dolores Padierna).', style='List Number')
    
    # Sección 4: BALAS ARGUMENTATIVAS
    h4 = doc.add_heading('4. BALAS ARGUMENTATIVAS PARA EL MS', level=1)
    for run in h4.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    balas = [
        '"La CNTE grita en la calle lo que sus diputados callan en la tribuna."',
        '"¿Cómo va a abrogar la CNTE la USICAMM si su propia ex-dirigente (Adela Piña) es quien la dirige?"',
        '"El deslinde del MS de la CNTE no es solo táctico, es ético: nosotros no buscamos curules, buscamos justicia para el maestro de Sonora."'
    ]
    for bala in balas:
        doc.add_paragraph(bala, style='List Bullet')
        
    # Conclusión
    doc.add_paragraph('-' * 50)
    concl = doc.add_paragraph()
    run_c = concl.add_run('Conclusión de Inteligencia: La CNTE funciona como una válvula de escape social. Sus movilizaciones sirven para legitimar a líderes que terminarán como candidatos de Morena. El MS debe capitalizar esta "traición" para atraer a los maestros que se sienten defraudados por la cúpula canteísta.')
    run_c.bold = True
    
    output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Intel_Cupula_CNTE_4T_V2.docx'
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")

if __name__ == "__main__":
    crear_reporte_cnte()
