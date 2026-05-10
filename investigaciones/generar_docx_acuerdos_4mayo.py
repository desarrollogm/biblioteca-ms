import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('MINUTA DE ACUERDOS: REUNIÓN ESTATAL DE COORDINADORES', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Fecha de Reunión: ').bold = True
p.add_run('4 de mayo de 2026\n')
p.add_run('Asunto: ').bold = True
p.add_run('Estrategia Operativa Mayo 2026\n')
p.add_run('Carácter: ').bold = True
p.add_run('Confidencial / Uso Interno MS')

doc.add_heading('1. ESTRATEGIA PARA CRECER EN REGIONES', level=2)
doc.add_paragraph('A) Visitas a los C.T. (Centros de Trabajo): Llevar la ruta de trabajo e invitar a los compañeros a unirse y apoyar la causa de manera presencial.', style='List Bullet')
doc.add_paragraph('B) Difusión en WhatsApp: Mantener el flujo constante de información técnica y propuestas generales en los grupos regionales y estatales.', style='List Bullet')
doc.add_paragraph('C) Expansión Regional: Identificar y buscar activamente a compañeros en las zonas geográficas donde la presencia del MS aún es débil.', style='List Bullet')

doc.add_heading('2. VISITA DE LA PRESIDENTA CLAUDIA SHEINBAUM', level=2)
doc.add_paragraph('A) Convocatoria Regional: Movilizar a las bases en las regiones donde estará la Presidenta para mostrar presencia organizada.', style='List Bullet')
doc.add_paragraph('B) Entrega de Documento: Asegurar la entrega física del oficio petitorio y la propuesta técnica PIP.', style='List Bullet')
doc.add_paragraph('C) Deslinde de la CNTE: Cuidar estrictamente no formar parte de las manifestaciones o bloqueos de la CNTE para preservar la identidad institucional del MS.', style='List Bullet')
doc.add_paragraph('D) Autonomía Regional: Respetar la organización y dinámicas propias de cada región en el estado.', style='List Bullet')
doc.add_paragraph('E) Gestión de Agenda: Solicitar apoyo para la agenda oficial con el diputado Pujol y el Senador Aguilar.', style='List Bullet')

doc.add_heading('3. ACCIONES DÍA DEL MAESTRO (15 DE MAYO)', level=2)
doc.add_paragraph('A) Manifestación SEC: Concentración masiva el 14 de mayo frente a las instalaciones de la Secretaría de Educación y Cultura (SEC).', style='List Bullet')
doc.add_paragraph('B) Campaña de Flyers: Distribución de volantes informativos sobre la manifestación, con énfasis especial en la región de Hermosillo.', style='List Bullet')
doc.add_paragraph('C) Registro de Demandas: Lanzamiento inmediato de un formulario digital para registrar las demandas específicas de la base y entregarlas el día de la manifestación.', style='List Bullet')

doc.add_heading('4. SEGUIMIENTO DE ACCIONES INSTITUCIONALES', level=2)
doc.add_paragraph('A) Cabildeo Federal/Local: Continuar las gestiones con la Senadora Lorenia Valles, así como con diputados federales y locales de la entidad.', style='List Bullet')
doc.add_paragraph('B) Gestión Froylán: Entablar diálogo directo con Froylán respecto a la solicitud pendiente.', style='List Bullet')
doc.add_paragraph('C) Seguimiento César Salazar: Coordinar con César Salazar el avance y seguimiento puntual de las demandas presentadas.', style='List Bullet')

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Minuta_Acuerdos_Estatal_4Mayo2026.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
