import os
import sys
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Force UTF-8 encoding for stdout and stderr to handle special characters on Windows
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

def generar_documento_analisis():
    doc = Document()
    
    # Configuración de estilos base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Título Principal (Arial 14, Bold, Color Negro)
    titulo_texto = 'ANÁLISIS DE INTELIGENCIA SINDICAL Y COMUNICADO POLÍTICO: VICTORIA TÁCTICA EN CAJEME'
    heading = doc.add_heading('', level=0)
    run = heading.add_run(titulo_texto)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo de Metadatos
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(
        'Fecha de Análisis: 10 de mayo de 2026 (Día de las Madres)\n'
        'Contexto: Gira Presidencial de la Dra. Claudia Sheinbaum Pardo en Sonora\n'
        'Lugar de Operación: Cajeme, Sonora — Comisión Especial del Magisterio Sonorense (MS)\n'
        'Clasificación: DOCUMENTO DE TRABAJO E INTERNACIÓN — CORE MS'
    )
    run_meta.font.size = Pt(10)
    run_meta.italic = True
    
    doc.add_paragraph('-' * 70)
    
    # ==================== SECCIÓN 1 ====================
    h1 = doc.add_heading('PARTE I: ANÁLISIS FORENSE Y EVALUACIÓN TÁCTICA DE LA ENTREGA', level=1)
    for run in h1.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(128, 0, 0) # Dark Red for section headers
        
    p1 = doc.add_paragraph()
    p1.add_run('1. El Hecho Histórico y su Contexto Simbólico:\n').bold = True
    p1.add_run(
        'Hoy, domingo 10 de mayo de 2026, coincidiendo con la conmemoración del Día de las Madres, '
        'el Magisterio Sonorense (MS) concretó una de las operaciones políticas y tácticas más trascendentales '
        'de su historia reciente. En Cajeme, Sonora, durante el marco de la quinta gira oficial de la Presidenta '
        'Dra. Claudia Sheinbaum Pardo por el estado, una delegación de base encabezada valientemente por las compañeras '
        'Kristina y Paloma logró eludir los cercos de seguridad presidencial y de logística estatal para establecer '
        'contacto físico y directo con el Poder Ejecutivo Federal y el Gobernador Alfonso Durazo Montaño.'
    )
    
    p2 = doc.add_paragraph()
    p2.add_run('2. La Firma Presidencial (Evidencia Irrefutable de Recepción):\n').bold = True
    p2.add_run(
        'A diferencia de las manifestaciones tradicionales que concluyen en la entrega fría a oficiales de correspondencia, '
        'las compañeras del MS consiguieron que la propia Presidenta de la República firmara de su puño y letra '
        '—con su distintiva firma abreviada "claudia Sh"— la esquina superior derecha del Oficio Núm. MS-01/2026, fechado '
        'en Hermosillo a 8 de mayo de 2026. Esta rúbrica, estampada directamente sobre el membrete institucional '
        'del Magisterio Sonorense, constituye una prueba jurídica y política irrefutable de que la propuesta técnica '
        'de la Pensión Intergeneracional Protegida (PIP) ha llegado formalmente al primer nivel del Estado Mexicano. '
        'Nadie en el aparato gubernamental o sindical de Sonora puede simular ignorancia o negar la existencia de esta vía.'
    )
    
    p3 = doc.add_paragraph()
    p3.add_run('3. Análisis del Video y Audio de Contacto (ms_presidenta sheimbaum_100526.mp4):\n').bold = True
    p3.add_run(
        'El análisis acústico y de contenido de la grabación de audio-video revela una ejecución discursiva de altísima precisión, '
        'ajustándose perfectamente a la "Tercera Vía" y al protocolo estratégico preestablecido por el MS:\n'
    )
    
    p3_a = doc.add_paragraph(style='List Bullet')
    p3_a.add_run('Posicionamiento Técnico: ').bold = True
    p3_a.add_run(
        '“...es una propuesta que tenemos para ustedes... no lastima la economía de los trabajadores ni las finanzas públicas '
        'y queremos que nos apoyen para poder recibir una mesa técnica que se ponga a analizar esta propuesta, presidenta...” '
        'Este enunciado neutraliza cualquier narrativa defensiva de la Secretaría de Hacienda federal o de los asesores de la SEP, '
        'exponiendo que la PIP no representa un boquete presupuestal, sino una reingeniería de fondos basada en el ahorro '
        'solidario y el fortalecimiento de PENSIONISSSTE.'
    )
    
    p3_b = doc.add_paragraph(style='List Bullet')
    p3_b.add_run('El Deslinde Clave ("¡No somos la CNTE!"): ').bold = True
    p3_b.add_run(
        'En un momento de máxima tensión del audio, el vocero declara con total nitidez: “¡No somos la CNTE, somos la base '
        'que queremos ser escuchados, queremos que nos dé un espacio, presidenta!”. Este deslinde es la pieza maestra de la '
        'interacción. Al desmarcarse explícitamente de la CNTE, el MS se inmuniza contra la etiqueta de "radicalismo intransigente" '
        'o "grupo de choque político". El MS se posiciona ante Sheinbaum como lo que es: un movimiento técnico, constructivo, '
        'democrático y genuinamente representativo de Sonora.'
    )
    
    p3_c = doc.add_paragraph(style='List Bullet')
    p3_c.add_run('Presión Cruzada (Gobernador Alfonso Durazo): ').bold = True
    p3_c.add_run(
        'El registro sonoro muestra que el MS también acorraló respetuosamente al Gobernador del estado en el mismo acto: '
        '“Gobernador, gobernador, somos el magisterio sonorense... queremos las mesas técnicas... ahí está mi compañero Cristian... '
        'firme esto por favor, gracias”. Al presionar a Durazo en presencia de Sheinbaum, se crea una pinza política: el Gobernador '
        'queda obligado ante la Presidenta a dar seguimiento local a las demandas de su propio magisterio.'
    )
    
    p4 = doc.add_paragraph()
    p4.add_run('4. Contraste de Operaciones (El Fracaso del SNTE 28 vs. El Éxito del MS):\n').bold = True
    p4.add_run(
        'Esta victoria táctica en Cajeme cobra mayor relevancia al contrastarse con las acciones recientes de la cúpula oficialista de la Sección 28 del SNTE. '
        'El pasado 7 de mayo, en una clínica del ISSSTE en el sur de Sonora, la burocracia del SNTE intentó realizar una manifestación simulada '
        'con consignas plagiadas del MS ("Jubilaciones Dignas" y "Abrogación de la Ley 2007"). Fue una maniobra de contención oportunista previa '
        'a la llegada de Sheinbaum para "lavarse la cara" ante las bases.\n\n'
        'El contraste es abismal: mientras el SNTE gesticula de espaldas a los acontecimientos y demanda una imposible "abrogación absoluta" '
        'sin números ni viabilidad, el MS camina con una propuesta seria bajo el brazo —la PIP, con pre-estudio favorable del Senado— y logra '
        'que la propia Presidenta le firme de recibido y autorice de facto el análisis en mesas técnicas. La cúpula simula; la base avanza.'
    )
    
    doc.add_paragraph('-' * 70)
    
    # ==================== SECCIÓN 2 ====================
    h2 = doc.add_heading('PARTE II: COMUNICADO OFICIAL PARA LA BASE MAGISTERIAL', level=1)
    for run in h2.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0) # Dark Green for the communiqué section
        
    p_com_header = doc.add_paragraph()
    p_com_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_com_h = p_com_header.add_run(
        '¡CON PROPUESTA Y DIÁLOGO: LA BASE MAGISTERIAL LE FIRMA EL FUTURO A SONORA!\n'
        'A LAS COMPAÑERAS Y COMPAÑEROS DE LAS BASES EN SONORA\n'
        'AL PUEBLO EN GENERAL'
    )
    run_com_h.font.size = Pt(12)
    run_com_h.bold = True
    
    p_body = doc.add_paragraph()
    p_body.add_run(
        'Hoy, 10 de mayo de 2026, en Cajeme, la dignidad de la base magisterial ha dado un paso de proporciones históricas. '
        'Mientras los líderes de oficina y las cúpulas cooptadas del sindicato tradicional se encierran a maquinar '
        'simulaciones estériles para justificar su inacción, la verdadera base del Magisterio Sonorense ha salido a la calle '
        'a construir el porvenir.'
    )
    
    p_body2 = doc.add_paragraph()
    p_body2.add_run(
        'En un día profundamente significativo para nuestra patria, en pleno Día de las Madres, fueron nuestras compañeras, '
        'las maestras del sur del estado, las que alzaron la voz y el documento de la justicia. Kristina y Paloma, encarnando la '
        'fuerza moral e intelectual del Magisterio Sonorense, rompieron todas las barreras burocráticas y entregaron '
        'directamente en las manos de la Presidenta de la República, Dra. Claudia Sheinbaum Pardo, nuestro Oficio MS-01/2026 '
        'con la propuesta de la Pensión Intergeneracional Protegida (PIP).'
    )
    
    p_body3 = doc.add_paragraph()
    p_body3.add_run(
        'La evidencia es irrefutable: la Presidenta ha recibido y firmado de puño y letra nuestro pliego técnico. '
        '¡No fuimos a gritar al vacío, fuimos a poner en el escritorio del país la única solución viable y actuarial '
        'que devolverá la jubilación digna a las y los trabajadores del Estado sin lastimar las finanzas nacionales! '
        'De frente a la Dra. Sheinbaum y en su propio rostro, le dejamos claro con orgullo y determinación: '
    ).italic = True
    
    p_quote = doc.add_paragraph()
    p_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_q = p_quote.add_run(
        '“¡No somos la CNTE, no somos la cúpula entreguista del SNTE! ¡Somos la base magisterial '
        'sonorense y exigimos ser escuchados en mesas técnicas con diálogo y propuesta!”'
    )
    run_q.bold = True
    run_q.font.size = Pt(12)
    run_q.italic = True
    
    p_body4 = doc.add_paragraph()
    p_body4.add_run(
        'Asimismo, acorralamos pacífica pero firmemente al Gobernador Alfonso Durazo Montaño ante los ojos de la comitiva presidencial. '
        'Le exigimos comprometer las fechas para las Mesas Técnicas de Trabajo con las autoridades del ISSSTE y la SEP. '
        'Porque la justicia no se mendiga; se arranca con el peso del argumento técnico, la seriedad intelectual y el respaldo del '
        'Senado de la República.'
    )
    
    p_body5 = doc.add_paragraph()
    p_body5.add_run(
        'El Magisterio Sonorense sigue demostrando que la lucha social de este siglo no se hace con pancartas vacías ni con paros '
        'que afectan a la niñez; se hace con inteligencia aplicada, rigor actuarial y un deslinde absoluto de los que usan '
        'al magisterio como trampolín político personal. Nuestra causa es única e inquebrantable: la dignificación del retiro '
        'de cada maestro y trabajador que cotiza ante el ISSSTE.'
    )
    
    p_body6 = doc.add_paragraph()
    p_body6.add_run(
        '¡La firma está puesta! ¡El compromiso de las mesas técnicas está en la mesa! ¡La base de Sonora ha ganado la batalla de la legitimidad!\n\n'
        'Que lo sepan bien los simuladores: la marea de la base ya no tiene marcha atrás. Seguimos trabajando, seguimos proponiendo, '
        'y seguiremos de pie hasta que la jubilación digna a 25 UMAs sea una realidad de ley para todos y todas.'
    )
    p_body6.bold = True
    
    p_com_footer = doc.add_paragraph()
    p_com_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_com_footer.add_run(
        'FRATERNALMENTE\n'
        '“POR LA DEFENSA DE NUESTROS DERECHOS LABORALES”\n'
        '“JUBILACIÓN DIGNA PARA TOD@S”\n\n'
        'EL COMITÉ DE COORDINACIÓN REGIONAL DEL MAGISTERIO SONORENSE'
    ).bold = True
    
    output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Analisis_Entrega_Presidenta_Cajeme_10Mayo2026.docx'
    doc.save(output_path)
    print(f"Document saved successfully to {output_path}")

if __name__ == "__main__":
    generar_documento_analisis()
