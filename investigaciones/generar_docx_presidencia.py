import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Configurar estilo Normal a Arial 12
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Arial'
font.size = Pt(12)

def add_heading(doc, text, level):
    if level == 0:
        h = doc.add_heading(text, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        h = doc.add_heading(text, level=level)
        
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    return h

# Título
add_heading(doc, 'INFORME CONFIDENCIAL - MAGISTERIO SONORENSE (MS)', 0)

# Metadatos
p = doc.add_paragraph()
p.add_run('A: ').bold = True
p.add_run('Dra. Claudia Sheinbaum Pardo, Presidenta Constitucional de los Estados Unidos Mexicanos\n')
p.add_run('CC: ').bold = True
p.add_run('Secretaría Técnica de la Presidencia / Oficina de Atención Ciudadana\n')
p.add_run('FECHA: ').bold = True
p.add_run('08 de mayo de 2026\n')
p.add_run('LUGAR: ').bold = True
p.add_run('Hermosillo, Sonora (Entregado en mano durante evento Punta Chueca)')

doc.add_paragraph()
doc.add_paragraph('Asunto: Radiografía de la Simulación Magisterial en Sonora y Presentación del Modelo PIP.').runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph('Señora Presidenta:')
doc.add_paragraph('Sea usted bienvenida al Estado de Sonora.')

doc.add_paragraph('Aprovechando el inicio de su gira de fin de semana, el Magisterio Sonorense (MS) se acerca a su equipo técnico de manera respetuosa e institucional, no con bloqueos ni chantajes mediáticos, sino con datos, auditorías forenses y propuestas técnicas actuariales.')

doc.add_paragraph('Mientras que ciertas fracciones magisteriales (como la cúpula de la CNTE) aprovechan la coyuntura del Mundial 2026 para amenazar con paros nacionales, y mientras la dirigencia oficial del SNTE guarda un silencio cómplice frente a la precariedad de los maestros sonorenses, el MS ha desarrollado la arquitectura técnica para la solución real de la crisis pensionaria:')

p_lista1 = doc.add_paragraph(style='List Bullet')
p_lista1.add_run('La Ley de Pensión Intergeneracional Protegida (PIP): ').bold = True
p_lista1.add_run('Un modelo matemático y jurídico sustentable que garantiza el retiro digno sin quebrar las finanzas estatales.')

p_lista2 = doc.add_paragraph(style='List Bullet')
p_lista2.add_run('Dossier "Fábricas de Ignorancia": ').bold = True
p_lista2.add_run('Una auditoría sociopolítica completa. Anexo a este informe, hacemos entrega confidencial del Tomo XXI: La Cúpula de Cristal, donde se expone con evidencia irrefutable cómo líderes sindicales han cooptado posiciones en la SEP y en las cámaras para su beneficio personal, traicionando a las bases.')

doc.add_paragraph('El Magisterio Sonorense está listo para sentarse en mesas técnicas, no en mesas de gritos. Somos el equilibrio, el "adulto en la habitación" que su gobierno necesita para aterrizar soluciones reales en el norte del país sin caer en las viejas prácticas del chantaje.')

doc.add_paragraph('Solicitamos respetuosamente que su equipo revise los tomos digitales anexos y agende una mesa de revisión técnica con nuestra Coordinación de Inteligencia Estratégica.')

doc.add_paragraph()
doc.add_paragraph('Atentamente,')
doc.add_paragraph()
doc.add_paragraph('La Coordinación General del Magisterio Sonorense (MS)\nConciencia, Lucha y Dignidad').runs[0].bold = True

# Guardar documento
doc.save(r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Informe_Presidencia_Sheinbaum_MS.docx')
print("Documento DOCX generado exitosamente.")
