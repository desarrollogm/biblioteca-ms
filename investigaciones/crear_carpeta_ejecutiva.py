from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Estilo base Arial 12 ---
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

def titulo_principal(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.name = 'Arial'
    r.font.size = Pt(16)
    r.font.bold = True
    return p

def seccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'
    r.font.size = Pt(14)
    r.font.bold = True
    return p

def subseccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.font.bold = True
    return p

def parrafo(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    return p

def parrafo_bold(doc, etiqueta, texto):
    p = doc.add_paragraph()
    r1 = p.add_run(etiqueta)
    r1.font.name = 'Arial'
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run(texto)
    r2.font.name = 'Arial'
    r2.font.size = Pt(12)
    return p

def bala(doc, texto):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(texto)
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    return p

def espacio(doc):
    doc.add_paragraph()

# =============================================
# PORTADA / ENCABEZADO
# =============================================
titulo_principal(doc, 'CARPETA EJECUTIVA')
titulo_principal(doc, 'MAGISTERIO SONORENSE')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Propuesta Técnica para Mesa de Primer Nivel')
r.font.name = 'Arial'
r.font.size = Pt(13)
r.font.bold = True
espacio(doc)

parrafo_bold(doc, 'Para: ', 'C. Presidenta Claudia Sheinbaum Pardo / C. Secretario Mario Delgado Carrillo (SEP)')
parrafo_bold(doc, 'De: ', 'Magisterio Sonorense — Representación de Docentes Federalizados, Sonora')
parrafo_bold(doc, 'Fecha: ', 'Mayo de 2026')
parrafo_bold(doc, 'Clasificación: ', 'Propuesta Técnica Formal')

espacio(doc)
p = doc.add_paragraph()
r = p.add_run('"No llegamos con la mano extendida. Llegamos con una propuesta construida desde las bases."')
r.font.name = 'Arial'
r.font.size = Pt(12)
r.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# =============================================
# PÁGINA 1 — QUIÉNES SOMOS
# =============================================
seccion(doc, 'PÁGINA 1 — QUIÉNES SOMOS')
subseccion(doc, 'El Magisterio Sonorense: Tres datos que nos definen')
espacio(doc)

parrafo_bold(doc, '1. Somos la voz técnica del magisterio federalizado en Sonora. ', 
    'El Magisterio Sonorense representa a los docentes federalizados adscritos a la Sección 28 del SNTE en el estado de Sonora. Somos un movimiento autónomo, apartidista y técnico, construido desde las bases del magisterio, no desde las cúpulas sindicales.')

espacio(doc)
parrafo_bold(doc, '2. Tenemos propuestas formalmente presentadas ante el Congreso. ',
    'El 21 de abril de 2026, el Magisterio Sonorense entregó formalmente en la Cámara de Diputados y en el Senado de la República dos instrumentos técnicos y legislativos:')
bala(doc, 'La Pensión Intergeneracional Protegida (PIP)')
bala(doc, 'La Iniciativa de Homologación UMA-Fronteriza para Sonora')

espacio(doc)
parrafo_bold(doc, '3. No pedimos la abrogación de una ley — pedimos su corrección quirúrgica. ',
    'A diferencia de otras organizaciones, el MS no exige destruir el sistema vigente sin proponer un reemplazo. Presentamos una alternativa financieramente viable, jurídicamente sustentada y presupuestalmente calculada.')

doc.add_page_break()

# =============================================
# PÁGINA 2 — EL PROBLEMA
# =============================================
seccion(doc, 'PÁGINA 2 — EL PROBLEMA: LA BRECHA QUE NADIE HA QUERIDO VER')
subseccion(doc, 'El maestro federalizado: el trabajador del Estado más castigado por la ley')
espacio(doc)

parrafo(doc, 'Existe una desigualdad estructural que viola el Artículo 1° Constitucional (principio de no discriminación) y afecta directamente a decenas de miles de maestros en Sonora y en toda la República:')
espacio(doc)

# Tabla comparativa
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
encabezados = ['Concepto', 'Sector Privado (IMSS)', 'Maestro Federalizado (ISSSTE)']
for i, h in enumerate(encabezados):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(12)

filas = [
    ['Tope de pensión máxima', '25 UMAs = $89,155/mes', '10 UMAs = $35,662/mes'],
    ['Base legal', 'Ley del Seguro Social, Art. 28', 'Ley ISSSTE, Art. Décimo Transitorio'],
    ['¿Cobra más si cotizó más?', 'SÍ', 'NO. El excedente se pierde'],
    ['Respaldo jurisprudencial', 'IMSS retuvo las 25 UMAs por voluntad política', 'SCJN ratificó el tope de 10 UMAs en 2021'],
]
for i, fila in enumerate(filas):
    for j, celda in enumerate(fila):
        cell = table.rows[i+1].cells[j]
        cell.text = celda
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(12)

espacio(doc)
subseccion(doc, 'Valores UMA vigentes 2026 (INEGI, vigentes desde 1 de febrero de 2026):')
bala(doc, 'UMA Diario: $117.31')
bala(doc, 'UMA Mensual: $3,566.22')
bala(doc, 'UMA Anual: $42,794.64')

espacio(doc)
subseccion(doc, 'El efecto real en un maestro sonorense')
parrafo(doc, 'Un maestro con 30 años de servicio en Sonora, que en activo percibía $60,000 mensuales, al jubilarse bajo el esquema de Décimo Transitorio TOPA EN $35,662 (10 UMAs). Pierde irreversiblemente más del 40% de su poder adquisitivo de por vida.')
espacio(doc)
parrafo(doc, 'Un trabajador del IMSS en la misma situación puede acceder hasta a $89,155. La misma ley, el mismo país, dos ciudadanos de distinta categoría.')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run('Esta no es una demanda ideológica. Es una demanda de igualdad constitucional.')
r.font.name = 'Arial'
r.font.size = Pt(12)
r.bold = True
r.italic = True

doc.add_page_break()

# =============================================
# PÁGINA 3 — LA PIP
# =============================================
seccion(doc, 'PÁGINA 3 — LA PIP: PENSIÓN INTERGENERACIONAL PROTEGIDA')
subseccion(doc, '¿Qué es la PIP?')
parrafo(doc, 'La Pensión Intergeneracional Protegida es la propuesta técnica central del Magisterio Sonorense. Fue diseñada con apoyo de datos actuariales y presentada formalmente en el Congreso en abril de 2026.')
espacio(doc)

subseccion(doc, 'Principios del modelo PIP')
espacio(doc)

table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Table Grid'
enc2 = ['Principio', 'Descripción']
for i, h in enumerate(enc2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(12)
filas2 = [
    ['Solidaridad intergeneracional', 'Los trabajadores activos contribuyen a un fondo que garantiza las pensiones de los jubilados actuales, como en el sistema previo a 2007.'],
    ['Viabilidad actuarial', 'El MS cuenta con estudios actuariales de ISSSTESON, UNISON e ITSON que respaldan la sostenibilidad del fondo.'],
    ['Garantía de derechos adquiridos', 'Las aportaciones individuales en Afores no se confiscan — se migran al fondo colectivo con valor garantizado.'],
    ['Progresividad', 'La transición es gradual, con horizonte de 15-20 años, evitando choques al sistema.'],
]
for i, fila in enumerate(filas2):
    for j, celda in enumerate(fila):
        cell = table2.rows[i+1].cells[j]
        cell.text = celda
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(12)

espacio(doc)
subseccion(doc, '¿Por qué la PIP es legislativamente viable?')
bala(doc, 'No requiere crear un organismo nuevo — opera dentro de la estructura del ISSSTE vigente.')
bala(doc, 'No implica gasto inmediato extraordinario — se financia con la redirección progresiva de aportaciones existentes.')
bala(doc, 'Tiene precedente en México — el IMSS mantuvo durante décadas un esquema solidario funcional (Ley 1973).')
bala(doc, 'Es comparable con modelos internacionales — Alemania, Canadá y países nórdicos operan pensiones solidarias sostenibles.')

espacio(doc)
subseccion(doc, 'Lo que el MS solicita sobre la PIP')
parrafo(doc, 'Que la SEP y la Presidencia instalen una mesa técnica con especialistas del ISSSTE, la SEP y representantes del MS para revisar, ajustar y viabilizar legislativamente la PIP antes del 15 de junio de 2026.')

doc.add_page_break()

# =============================================
# PÁGINA 4 — UMA-FRONTERIZA
# =============================================
seccion(doc, 'PÁGINA 4 — LA UMA-FRONTERIZA: JUSTICIA ESPECÍFICA PARA SONORA')
subseccion(doc, 'El problema específico de Sonora')
parrafo(doc, 'Los docentes federalizados en Sonora, particularmente en zonas fronterizas (Nogales, Agua Prieta, San Luis Río Colorado, Sonoyta), enfrentan una doble injusticia:')
bala(doc, 'El tope de 10 UMAs que aplica a todo el país.')
bala(doc, 'El costo de vida en zonas fronterizas es estructuralmente más alto (hasta 30-40% superior al promedio nacional) dada su dinámica económica binacional con los Estados Unidos.')

espacio(doc)
subseccion(doc, 'La Propuesta UMA-Fronteriza')
parrafo(doc, 'El Magisterio Sonorense propone al Congreso una diferenciación regional de la UMA para efectos de cálculo de pensiones en municipios fronterizos, similar a los criterios de zona económica especial que ya existen en materia fiscal y laboral.')
espacio(doc)

table3 = doc.add_table(rows=4, cols=3)
table3.style = 'Table Grid'
enc3 = ['Municipio', 'UMA aplicable', 'Pensión máxima mensual']
for i, h in enumerate(enc3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(12)
filas3 = [
    ['Interior de Sonora (estándar)', '10 UMAs', '$35,662/mes (actual, sin cambio inmediato)'],
    ['Zona Fronteriza (meta corto plazo)', '15 UMAs', '$53,493/mes (+$17,831 adicionales)'],
    ['Meta mediano plazo (con PIP)', '25 UMAs', '$89,155/mes (equiparación con IMSS)'],
]
for i, fila in enumerate(filas3):
    for j, celda in enumerate(fila):
        cell = table3.rows[i+1].cells[j]
        cell.text = celda
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(12)

espacio(doc)
subseccion(doc, '¿Por qué esto es urgente?')
bala(doc, 'Los maestros en zonas fronterizas compiten laboralmente con el sector privado que paga en dólares.')
bala(doc, 'La fuga de docentes calificados hacia empleos binacionales es un fenómeno documentado en las escuelas sonorenses.')
bala(doc, 'El gobierno federal tiene interés estratégico en estabilizar el magisterio en la frontera norte.')

espacio(doc)
parrafo(doc, 'La Iniciativa de UMA-Fronteriza fue presentada en el Senado de la República por el Magisterio Sonorense en abril de 2026. Establece la reforma al Artículo Décimo Transitorio de la Ley del ISSSTE para incluir un criterio de ajuste geográfico en zonas fronterizas.')

doc.add_page_break()

# =============================================
# PÁGINA 5 — RESPALDO LEGAL
# =============================================
seccion(doc, 'PÁGINA 5 — EL RESPALDO LEGAL: POR QUÉ ESTO ES POSIBLE')
subseccion(doc, 'El precedente que el gobierno federal ya creó: el caso IMSS')
parrafo(doc, 'En 2020, la Suprema Corte de Justicia de la Nación emitió la Jurisprudencia 2a./J. 164/2019, dictaminando que la generación de transición del IMSS debía toparse en 10 salarios mínimos al optar por la Ley de 1973.')
espacio(doc)
parrafo(doc, 'Sin embargo, el IMSS no lo aplicó. El Director General Zoé Robledo y el Consejo Técnico del IMSS tomaron la decisión política de respetar las 25 UMAs para quienes habían cotizado bajo ese parámetro, argumentando la protección de derechos adquiridos.')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run('Conclusión jurídica clave: ')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
r2 = p.add_run('Si el Estado mexicano, a través del IMSS, puede decidir proteger derechos por encima de una jurisprudencia de la SCJN cuando es en beneficio del trabajador, el mismo Estado, a través del ISSSTE, puede y debe aplicar el mismo principio protector a los maestros federalizados.')
r2.font.name = 'Arial'; r2.font.size = Pt(12)

espacio(doc)
subseccion(doc, 'Argumentos constitucionales y legales')
espacio(doc)

table4 = doc.add_table(rows=6, cols=2)
table4.style = 'Table Grid'
enc4 = ['Fundamento', 'Contenido']
for i, h in enumerate(enc4):
    cell = table4.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.name = 'Arial'
    cell.paragraphs[0].runs[0].font.size = Pt(12)
filas4 = [
    ['Artículo 1° Constitucional', 'Prohibición de discriminación. No puede haber tope de 10 UMAs para maestros y 25 para trabajadores del IMSS sin violar este precepto.'],
    ['Principio Pro Persona', 'Ante dos interpretaciones posibles de la ley, el Estado debe elegir la que más beneficia al trabajador.'],
    ['Artículo 123 Constitucional', 'Garantía de seguridad social para todos los trabajadores en condiciones de dignidad e igualdad.'],
    ['Ley Federal del Trabajo', 'Principio de irrenunciabilidad de los derechos laborales — el excedente de cotización no puede ser confiscado.'],
    ['Precedente IMSS', 'El propio gobierno federal ya sentó el estándar de 25 UMAs como umbral digno para la generación de transición.'],
]
for i, fila in enumerate(filas4):
    for j, celda in enumerate(fila):
        cell = table4.rows[i+1].cells[j]
        cell.text = celda
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(12)

espacio(doc)
subseccion(doc, 'Ruta legislativa propuesta')
bala(doc, 'PASO 1: Mesa técnica MS + SEP + ISSSTE (junio 2026)')
bala(doc, 'PASO 2: Reforma al Artículo Décimo Transitorio de la Ley ISSSTE (elevar tope de 10 a 25 UMAs + diferencial fronterizo)')
bala(doc, 'PASO 3: Presentación en Comisión de Educación y Comisión de Seguridad Social del Congreso')
bala(doc, 'PASO 4: Aprobación y vigencia (ciclo legislativo 2026-2027)')

doc.add_page_break()

# =============================================
# PÁGINA 6 — 5 PUNTOS PETITORIOS
# =============================================
seccion(doc, 'PÁGINA 6 — LO QUE PEDIMOS: 5 PUNTOS ESPECÍFICOS Y VIABLES')
espacio(doc)

subseccion(doc, 'PUNTO 1 — Mesa Técnica Permanente con la SEP e ISSSTE')
parrafo_bold(doc, 'Qué pedimos: ', 'La instalación de una mesa técnica de trabajo conjunta entre representantes del Magisterio Sonorense, la SEP y el ISSSTE, con calendario de sesiones y minutas firmadas.')
parrafo_bold(doc, 'Para qué: ', 'Revisar, ajustar y viabilizar legislativamente la PIP y la Iniciativa UMA-Fronteriza.')
parrafo_bold(doc, 'Plazo propuesto: ', 'Primera sesión antes del 15 de junio de 2026.')
espacio(doc)

subseccion(doc, 'PUNTO 2 — Congelamiento del Deterioro de Pensiones Durante el Proceso')
parrafo_bold(doc, 'Qué pedimos: ', 'Que durante el tiempo que dure la mesa técnica, no se apliquen medidas que deterioren adicionalmente las pensiones del Décimo Transitorio (nuevas jurisprudencias restrictivas, cambios reglamentarios, etc.).')
parrafo_bold(doc, 'Para qué: ', 'Negociar en un piso estable, sin que la situación empeore mientras se busca la solución legislativa.')
espacio(doc)

subseccion(doc, 'PUNTO 3 — Revisión del Artículo Décimo Transitorio')
parrafo_bold(doc, 'Qué pedimos: ', 'Que la SEP y la Presidencia respalden ante el Congreso la revisión legislativa del Artículo Décimo Transitorio de la Ley ISSSTE para elevar progresivamente el tope de 10 a 25 UMAs, homologando el criterio del IMSS.')
parrafo_bold(doc, 'Para qué: ', 'Cerrar la brecha de discriminación entre trabajadores del sector público y privado en materia de pensiones.')
espacio(doc)

subseccion(doc, 'PUNTO 4 — Reconocimiento Formal del MS como Interlocutor Técnico')
parrafo_bold(doc, 'Qué pedimos: ', 'Que el gobierno federal reconozca formalmente al Magisterio Sonorense como actor técnico legítimo en las mesas de negociación de política educativa y pensionaria en Sonora.')
parrafo_bold(doc, 'Para qué: ', 'Garantizar que la voz de los maestros sonorenses esté representada en las decisiones que los afectan directamente, más allá de las cúpulas sindicales tradicionales.')
espacio(doc)

subseccion(doc, 'PUNTO 5 — Agenda para la UMA-Fronteriza')
parrafo_bold(doc, 'Qué pedimos: ', 'Que la Iniciativa de UMA-Fronteriza presentada en el Senado sea turnada formalmente a la Comisión de Seguridad Social para su dictaminación en el periodo legislativo de septiembre-diciembre de 2026.')
parrafo_bold(doc, 'Para qué: ', 'Que los docentes de Nogales, Agua Prieta, San Luis Río Colorado y otras ciudades fronterizas tengan una solución específica a su situación de costo de vida diferenciado.')

espacio(doc)
doc.add_paragraph('─' * 60)
espacio(doc)

seccion(doc, 'CIERRE — NUESTRA OFERTA AL GOBIERNO FEDERAL')
parrafo(doc, 'El Magisterio Sonorense no llega a esta reunión como adversario. Llegamos como aliados técnicos de una transformación que el propio gobierno ha prometido y que la base magisterial necesita ver materializada.')
espacio(doc)
subseccion(doc, 'Lo que el gobierno gana al sentarse con el MS:')
bala(doc, 'Una solución técnica lista para legislar — sin necesidad de empezar desde cero.')
bala(doc, 'Un movimiento que no bloquea aeropuertos ni para escuelas — que construye propuestas.')
bala(doc, 'Un precedente de modernización pensionaria que puede ser presentado como logro de la administración.')
bala(doc, 'Estabilidad en el magisterio de Sonora — zona estratégica de la frontera norte.')
bala(doc, 'Un contrapeso técnico a las demandas maximalistas de la CNTE — moderación con contenido real.')

espacio(doc)
doc.add_paragraph('─' * 60)
parrafo_bold(doc, 'Representación: ', 'Magisterio Sonorense — Docentes Federalizados, Sección 28, Sonora')
parrafo_bold(doc, 'Propuestas entregadas: ', 'Cámara de Diputados y Senado de la República (abril 2026)')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run('Carpeta preparada para entrega en reunión de primer nivel. Versión: Mayo 2026')
r.font.name = 'Arial'; r.font.size = Pt(11); r.italic = True

# --- GUARDAR ---
output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Carpeta_Ejecutiva_MS_Presidencia.docx'
doc.save(output_path)
print(f"LISTO. Documento guardado en: {output_path}")
