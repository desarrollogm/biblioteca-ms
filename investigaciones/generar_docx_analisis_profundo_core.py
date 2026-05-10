import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('INFORME DE INTELIGENCIA ESTRATÉGICA: ANÁLISIS DE REUNIÓN DE COORDINADORES (04/05/26)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Documento de Trabajo de Alta Dirección para el CORE del MS\n').italic = True
p.add_run('Fecha de Elaboración: ').bold = True
p.add_run('5 de mayo de 2026\n')
p.add_run('Fuentes: ').bold = True
p.add_run('Audio de reunión, Minuta de Acuerdos, Documentos de Anatomía del CORE y Ofensiva de Mayo.')

doc.add_heading('1. DIAGNÓSTICO DE LIDERAZGO (ANATOMÍA DEL CORE)', level=2)
doc.add_paragraph('El equipo ncleo del Magisterio Sonorense ha alcanzado una madurez poltica sin precedentes, pasando de la reaccin emocional a la proyeccin tcnica. Sin embargo, se identifican roles crticos que deben optimizarse para la "Guerra de Mayo":')

p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('La Dualidad Operativa (Cristian - Juan Pablo): ').bold = True
p1.add_run('Mientras Cristian funge como el motor logstico y rostro diplomtico, Juan Pablo aporta el "Logos" (la razn) mediante la PIP. Esta frmula es el activo ms valioso del movimiento ante el Senado.')

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run('El Filtro Ético y el Dato (Paloma - Jisus): ').bold = True
p2.add_run('La firmeza de Paloma evita la contaminacin del movimiento con oportunismos, mientras que Jisus construye la infraestructura tecnolgica (Formularios/Atlas de Demandas) que sustenta la lucha.')

doc.add_heading('2. ANÁLISIS DE LA "OFENSIVA HERMOSILLO" (14-15 MAYO)', level=2)
doc.add_paragraph('La centralizacin de acciones en la capital el jueves 14 de mayo es un acierto tctico. El movimiento abandona la dispersin regional que "no pesa" para golpear en el corazn de la SEC.')

doc.add_heading('2.1 El Arma Secreta: El "Atlas de Demandas"', level=3)
doc.add_paragraph('El formulario digital no es solo una encuesta; es la transformacin de la queja en inteligencia. El 14 de mayo, el MS no llegar a la SEC a "pedir", sino a "notificar" un estado de crisis documentado escuela por escuela. Esto obliga a las autoridades a dar respuestas tcnicas, no polticas.')

doc.add_heading('2.2 Contrarrestando el "Pan y Circo"', level=3)
doc.add_paragraph('El sindicato oficial est intentando "comprar" a la base con regalos y bailes. El MS debe atacar este punto con un mensaje de contraste: "Mientras ellos te regalan una sombrilla, nosotros luchamos por tu jubilacin". Dignidad vs. Distraccin.')

doc.add_heading('3. FACTORES DE RIESGO Y NEUTRALIZACIÓN', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Riesgo Identificado'
hdr_cells[1].text = 'Estrategia de Neutralización'

row1 = table.rows[1].cells
row1[0].text = 'Mimetismo con la CNTE'
row1[1].text = 'Distanciamiento fsico y esttico. Cuidar la "foto" para que Presidencia nos vea como tcnicos, no como choque.'

row2 = table.rows[2].cells
row2[0].text = 'Logstica del "Megapuentn"'
row2[1].text = 'Enfocar la energa en el da 14. Convertir el 15 en un da de "vigilia informativa" digital tras la respuesta gubernamental.'

row3 = table.rows[3].cells
row3[0].text = 'Sabotaje en Hermosillo'
row3[1].text = 'Reforzar la red de delegados en la capital. Chayo requiere apoyo del CORE para evitar la infiltracin sindical oficial.'

doc.add_heading('4. CONCLUSIONES Y ACCIÓN "HASTA LA COCINA"', level=2)
doc.add_paragraph('El CORE tiene el control de la narrativa. La visita de la Presidenta (8-10 mayo) es el "pre-detonante" y el Da del Maestro es la "explosin".')

doc.add_paragraph('• Acción Inmediata: Viralizar el formulario de demandas. El "Atlas" debe ser masivo para ser incontestable.', style='List Bullet')
doc.add_paragraph('• Lobbying de Pinza: Usar a Lorenia Valles para validar la PIP mientras la base presiona en la SEC.', style='List Bullet')
doc.add_paragraph('• Mensaje Clave: "Unidad desde las bases, no desde las cpulas".', style='List Bullet')

p_final = doc.add_paragraph()
p_final.add_run('\n"El éxito no dependerá de ser miles, sino de la disciplina operativa."').italic = True
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Analisis_Profundo_Reunion_CORE_4Mayo.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
