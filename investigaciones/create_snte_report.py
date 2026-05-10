import os
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# Heading
heading = doc.add_heading('Reporte de Inteligencia: Convocatoria SNTE Oficial - 1ro de Mayo 2026', 0)
heading.style.font.name = 'Arial'
heading.style.font.size = Pt(14)
heading.style.font.bold = True

p1 = doc.add_paragraph()
r1 = p1.add_run('Objetivo: ')
r1.bold = True
p1.add_run('Análisis de aforo y presencia de la cúpula oficial (Sección 28/54) durante el desfile del 1ro de Mayo en la Región Yaqui.')

doc.add_paragraph()
h1 = doc.add_paragraph()
r_h1 = h1.add_run('1. Conteo de Compañeros (Aforo Real)')
r_h1.font.size = Pt(14)
r_h1.bold = True

p2 = doc.add_paragraph()
p2.add_run('Haciendo un conteo fila por fila del bloque principal (que viene marchando con una formación estructurada en el paso a desnivel):')

doc.add_paragraph('- Hay 2 personas sosteniendo la lona al frente.', style='List Bullet')
doc.add_paragraph('- Detrás de ellos vienen aproximadamente 8 filas de entre 4 y 6 personas cada una.', style='List Bullet')
doc.add_paragraph('- Contando a los rezagados que se ven más atrás en la toma, el bloque total es de unas 45 a 55 personas máximo. Es un contingente muy reducido para tratarse de la maquinaria oficial del sindicato en una región tan importante como la Yaqui (Cajeme).', style='List Bullet')

doc.add_paragraph()
h2 = doc.add_paragraph()
r_h2 = h2.add_run('2. Distinción entre Sección 28 y Sección 54')
r_h2.font.size = Pt(14)
r_h2.bold = True

doc.add_paragraph('- Todos pertenecen al mismo bloque: Es imposible distinguir visualmente dos secciones diferentes aquí, porque todos están uniformados de manera idéntica (camisa blanca/polo y la misma gorra color naranja). No hay dos frentes, ni dos tipos de uniformes que sugieran una participación nutrida de ambas secciones.', style='List Bullet')
doc.add_paragraph('- La Lona los delata: Al acercar la imagen a la lona principal que llevan al frente, dice claramente "SINDICATO NACIONAL DE TRABAJADORES DE LA EDUCACIÓN" en letras rojas, y en las letras negras más pequeñas de abajo se distingue "SECCIÓN 28".', style='List Bullet')

doc.add_paragraph()
h3 = doc.add_paragraph()
r_h3 = h3.add_run('3. Conclusión de Inteligencia Táctica')
r_h3.font.size = Pt(14)
r_h3.bold = True

p3 = doc.add_paragraph()
p3.add_run('Es una muestra de debilidad de convocatoria extrema. Para ser el SNTE oficial (Sección 28) desfilando en la Región Yaqui, juntar apenas a ~50 personas uniformadas evidencia la crisis de legitimidad y el descontento de la base. No se ven contingentes separados de la Sección 54 apoyándolos, es un solo bloque pequeño y altamente controlado.')

p4 = doc.add_paragraph()
p4.add_run('\n[NOTA: INGENIERO, COMO LA FOTO ME LA PASASTE POR EL CHAT, DEBES PEGARLA MANUALMENTE AQUÍ ABAJO ANTES DE MANDAR EL DOCUMENTO]').bold = True

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Inteligencia_Convocatoria_SNTE.docx'
doc.save(output_path)
print(f"SNTE Report saved to {output_path}")
