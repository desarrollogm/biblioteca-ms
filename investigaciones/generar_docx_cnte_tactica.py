import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Configurar estilo Normal a Arial 12
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Arial'
font.size = Pt(12)

def add_heading(doc, text, level):
    if level == 0:
        h = doc.add_heading(text, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        h = doc.add_heading(text, level=level)
        
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    return h

# Título
add_heading(doc, 'INFORME DE INTELIGENCIA TÁCTICA', 0)

# Metadatos
p = doc.add_paragraph()
p.add_run('OBJETIVO: ').bold = True
p.add_run('Análisis del Comunicado CNTE (07 Mayo 2026) sobre el Calendario Escolar y el Mundial 2026.\n')
p.add_run('PARA: ').bold = True
p.add_run('Coordinación General MS\n')
p.add_run('FECHA: ').bold = True
p.add_run('08 de mayo de 2026')

doc.add_paragraph()

# Sección 1
add_heading(doc, '1. PUNTOS CLAVE DEL COMUNICADO (RADIOGRAFÍA DEL DOCUMENTO)', level=1)

p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('Rechazo a la SEP: ').bold = True
p1.add_run('La CNTE rechaza formalmente el adelanto del fin de cursos al 5 de junio de 2026 anunciado por Mario Delgado (SEP) y CONAEDU.')

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run('Descrédito del argumento oficial: ').bold = True
p2.add_run('Califican de "totalmente falso" el pretexto de las olas de calor para adelantar las vacaciones, argumentando que el gobierno jamás se ha preocupado por la infraestructura climática de las escuelas, dejándole esa carga a padres y maestros.')

p3 = doc.add_paragraph(style='List Bullet')
p3.add_run('La verdadera narrativa: ').bold = True
p3.add_run('Afirman que es una decisión "cupular y neoliberal" para vaciar las calles y escuelas, protegiendo el negocio y los intereses de la "oligarquía financiera" durante el Mundial de Fútbol 2026.')

p4 = doc.add_paragraph(style='List Bullet')
p4.add_run('La Amenaza Principal (El Chantaje): ').bold = True
p4.add_run('Reafirman su amenaza de estallar una Huelga Nacional justo en las fechas del Mundial para sabotearlo si no se cumplen sus demandas. Su consigna lo resume: "¡Si no hay solución, no rodará su balón!".')

doc.add_paragraph()

# Sección 2
add_heading(doc, '2. IMPLICACIONES ESTRATÉGICAS PARA EL MAGISTERIO SONORENSE (MS)', level=1)

doc.add_paragraph('A. El Suicidio Mediático de la CNTE').runs[0].bold = True
doc.add_paragraph('Al amenazar con boicotear el Mundial de Fútbol (un evento de máxima distracción y arraigo popular en México), la CNTE está cometiendo un error táctico monumental. Se van a echar a la opinión pública, a los padres de familia y a los medios de comunicación encima. El gobierno federal tendrá la excusa perfecta (el repudio popular) para reprimirlos o ignorarlos sin costo político.')

doc.add_paragraph()
doc.add_paragraph('B. La Oportunidad de Oro para el Contraste (El "Adulto en la Habitación")').runs[0].bold = True
doc.add_paragraph('Este es el momento perfecto para que el MS brille por contraste:')
p_sub1 = doc.add_paragraph(style='List Bullet 2')
p_sub1.add_run('CNTE: ').bold = True
p_sub1.add_run('Gritos, bloqueos, chantaje internacional y repudio social.')
p_sub2 = doc.add_paragraph(style='List Bullet 2')
p_sub2.add_run('SNTE Oficial: ').bold = True
p_sub2.add_run('Sumisión total, aceptando los recortes al calendario sin chistar.')
p_sub3 = doc.add_paragraph(style='List Bullet 2')
p_sub3.add_run('Magisterio Sonorense (MS): ').bold = True
p_sub3.add_run('Propuesta técnica, rigor actuarial (Ley PIP), auditorías forenses y defensa legítima en los Congresos.')

doc.add_paragraph()
doc.add_paragraph('C. Recomendación de Posicionamiento MS:').runs[0].bold = True
p_rec1 = doc.add_paragraph('1. No subirse al ring del Mundial: ', style='List Number')
p_rec1.runs[0].bold = True
p_rec1.add_run('El MS no debe secundar ni aplaudir el boicot al Mundial. Es una batalla perdida mediáticamente.')

p_rec2 = doc.add_paragraph('2. Aprovechar el argumento de la infraestructura: ', style='List Number')
p_rec2.runs[0].bold = True
p_rec2.add_run('El único punto rescatable del documento de la CNTE es la denuncia sobre la falta de infraestructura para el calor. El MS en Sonora (donde el calor es extremo) puede tomar ese argumento y volverlo técnico: "En Sonora no necesitamos que adelanten el ciclo por el calor, necesitamos inversión auditable en refrigeración y subestaciones, dinero que se pierde en la corrupción del SNTE".')

p_rec3 = doc.add_paragraph('3. Reforzar la Vía Institucional: ', style='List Number')
p_rec3.runs[0].bold = True
p_rec3.add_run('Emitir mensajes internos a la base magisterial sonorense mostrando que mientras la cúpula de la CNTE juega a la política del chantaje con el fútbol, el MS está empujando leyes de pensiones viables en el Senado.')

# Guardar documento
doc.save(r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Informe_Tactica_CNTE_Mayo2026.docx')
print("Documento DOCX generado exitosamente.")
