import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('ANÁLISIS ESTRATÉGICO DE DOCUMENTO PETITORIO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Documento: "Establecimiento de mesas Técnicas de análisis" (Entregable a la Presidenta)', level=1)

p = doc.add_paragraph()
p.add_run('Clasificación: ').bold = True
p.add_run('INTELIGENCIA ESTRATÉGICA — REVISIÓN PRE-ENTREGA\n')
p.add_run('Fecha de análisis: ').bold = True
p.add_run('5 de mayo de 2026\n')
p.add_run('Destinatario final: ').bold = True
p.add_run('Dra. Claudia Sheinbaum Pardo, Presidenta de México\n')
p.add_run('Remitente: ').bold = True
p.add_run('Kristian David López Copado (Magisterio Sonorense)')

doc.add_heading('1. DIAGNÓSTICO GENERAL DEL DOCUMENTO', level=2)
doc.add_paragraph('El documento presenta una estructura impecable para el perfil de la destinataria. Claudia Sheinbaum tiene un perfil científico y técnico; el texto evita la retórica sindical tradicional (confrontativa o de victimización) y adopta un tono propositivo, institucional y técnico. Esta es la mayor fortaleza del documento.')

doc.add_heading('2. ANÁLISIS DE POSICIONAMIENTO (DIFERENCIACIÓN)', level=2)
doc.add_paragraph('El texto ejecuta magistralmente la "Tercera Vía" del Magisterio Sonorense:')

p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('Desmarque del SNTE: ').bold = True
p1.add_run('El texto ataca sutilmente la propuesta del SNTE al señalar que "involucra el aprovechamiento de los recursos... pero no toca los elementos más perjudiciales". Expone al SNTE como administrador del problema, no solucionador.')

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run('Desmarque de la CNTE: ').bold = True
p2.add_run('En el cuarto párrafo, se declara explícitamente: "El Magisterio Sonorense es un movimiento independiente... ajeno a la Coordinadora Nacional de Trabajadores de la Educación". Esto es crucial para que Presidencia no encapsule al MS como un grupo de choque radical.')

doc.add_heading('3. ANÁLISIS DE LA PETICIÓN CENTRAL (CALL TO ACTION)', level=2)
p = doc.add_paragraph()
r = p.add_run('Fortaleza táctica: No se exige la abrogación inmediata de la ley mediante decreto.')
r.bold = True
doc.add_paragraph('Exigir abrogación inmediata suele cerrar puertas. En su lugar, el MS solicita "mesas técnicas de trabajo con autoridades de primer nivel del ISSSTE y la SEP".')
doc.add_paragraph('• Es una petición altamente concedible por el Ejecutivo.\n• Traslada la presión de "cambiar la ley" a "sentarnos a revisar los números".\n• Mantiene la puerta abierta al diálogo institucional.')

doc.add_heading('4. ANÁLISIS DEL ANEXO TÉCNICO (LA PIP)', level=2)
doc.add_paragraph('La segunda parte del documento desglosa la Pensión Intergeneracional Protegida (PIP).')

doc.add_heading('Puntos fuertes a resaltar:', level=3)
doc.add_paragraph('• Blindaje Constitucional: El FPIP como fondo protegido evita el fantasma de los "saqueos" gubernamentales del pasado.', style='List Bullet')
doc.add_paragraph('• Fortalecimiento de PENSIONISSSTE: Trasladar cuentas privadas (AFOREs) al sector público alinea la propuesta con la narrativa estatista del gobierno actual ("fortalecer las instituciones del Estado").', style='List Bullet')
doc.add_paragraph('• El comodín de oro (Instituto Belisario Domínguez): Mencionar que la propuesta está respaldada por un estudio del Senado le otorga un peso institucional masivo. Ya no es "una idea de maestros", es un modelo con sustento legislativo.', style='List Bullet')

doc.add_heading('5. RECOMENDACIONES DE MEJORA ANTES DE LA IMPRESIÓN FINAL', level=2)
p_alert = doc.add_paragraph()
r_alert = p_alert.add_run('⚠️ REVISIÓN CRÍTICA PRE-ENTREGA:')
r_alert.bold = True
r_alert.font.color.rgb = RGBColor(255, 0, 0)

doc.add_paragraph('1. Cuidado con la codificación (Ortografía técnica): El texto extraído digitalmente muestra errores de caracteres (ej. "Técnicas", "diálogo", "Pensión" aparecen con símbolos raros en la lectura digital). ES VITAL asegurarse de que el PDF o Word final impreso tenga los acentos y las "ñ" perfectamente legibles. Un documento presidencial no puede tener errores de fuente.', style='List Bullet')
doc.add_paragraph('2. El término "Ahorro Solidario": El documento menciona "$3.25 por cada peso del trabajador... (antes llamado ahorro solidario)". El Ahorro Solidario fue un mecanismo de la reforma de Felipe Calderón (2007). Para evitar rechazo ideológico, es acertado que digan "antes llamado", pero se sugiere enfatizar que es "Aportación Patronal del Estado".', style='List Bullet')
doc.add_paragraph('3. Formato físico: El documento debe entregarse en papel opalina o membretado de alta calidad, firmado en original (tinta azul preferentemente para evidenciar que no es copia), en un sobre manila o carpeta institucional limpia.', style='List Bullet')

doc.add_heading('6. CONCLUSIÓN', level=2)
p_end = doc.add_paragraph()
r_end = p_end.add_run('El documento es una pieza maestra de cabildeo técnico. Está perfectamente calibrado para la administración Sheinbaum: técnico, institucional, propositivo y con respaldo del Senado. Si se entrega exitosamente durante la gira en Sonora, posicionará automáticamente al MS por encima del SNTE y la CNTE en el debate nacional de pensiones.')
r_end.bold = True

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Analisis_Documento_Presidenta_Mayo2026.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
