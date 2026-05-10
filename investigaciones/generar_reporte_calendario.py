import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_intel_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('INTELIGENCIA ESTRATÉGICA MS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header Info
    p = doc.add_paragraph()
    p.add_run('REPORTE: ').bold = True
    p.add_run('Análisis del Calendario de la Ignorancia (Ciclo 2025-2026)\n')
    p.add_run('FECHA: ').bold = True
    p.add_run('06 de mayo de 2026\n')
    p.add_run('DESTINO: ').bold = True
    p.add_run('CORE Magisterio Sonorense\n')
    p.add_run('CLASIFICACIÓN: ').bold = True
    p.add_run('ESTRICTAMENTE CONFIDENCIAL / NIVEL 1\n')
    
    # Sections
    sections = [
        ("1. EL HALLAZGO: LA ENCUESTA DE MARIO DELGADO", 
         "Se ha detectado una consulta interna liderada por el Secretario Mario Delgado para validar el cierre del ciclo escolar el 30 de junio de 2026. Este ajuste responde a la incapacidad operativa de la SEP ante el Mundial FIFA y las olas de calor climáticas."),
        
        ("2. PUNTOS CRÍTICOS DEL RECORTE", 
         "* Cierre de clases adelantado al 30 de junio.\n"
         "* Pérdida neta de 15 días laborables y académicos.\n"
         "* Justificación oficial: Protección ante el calor. Justificación real: Logística del Mundial."),
        
        ("3. DEBILIDAD INSTITUCIONAL EN SONORA", 
         "El Secretario de Educación en Sonora, Froylán Gámez, se encuentra operando activamente su candidatura a la gubernatura de 2027 (destapado públicamente el 2 de mayo). Este enfoque electoral deja a la SEC sin una postura técnica real, supeditándose a las decisiones de la Federación sin defender la particularidad de la zona fronteriza o sonora."),
        
        ("4. ACCIÓN ESTRATÉGICA RECOMENDADA", 
         "El Magisterio Sonorense (MS) debe filtrar esta información a la base docente antes del anuncio oficial de mañana (7 de mayo). La narrativa debe ser la denuncia de la 'Fábrica de Ignorancia' que prefiere cerrar escuelas por un mundial que invertir en infraestructura digna y en el aumento a la UMA magistral.")
    ]
    
    for title_text, body_text in sections:
        h = doc.add_heading(title_text, level=1)
        doc.add_paragraph(body_text)
    
    # Save
    output_path = r'c:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Analisis_Mundialista_Calendario_Mario_Delgado_Mayo2026.docx'
    doc.save(output_path)
    print(f"Reporte generado exitosamente en: {output_path}")

if __name__ == "__main__":
    create_intel_report()
