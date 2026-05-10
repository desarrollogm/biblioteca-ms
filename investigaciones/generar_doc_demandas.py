import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_demands_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('CATÁLOGO DE DEMANDAS Y NECESIDADES ESTATALES 2026', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Organization Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run('MAGISTERIO SONORENSE\n')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(139, 0, 0) # Dark Red
    
    # Subheader
    info = doc.add_paragraph()
    info.add_run('REPORTE TÉCNICO DE NECESIDADES ESCOLARES\n').bold = True
    info.add_run('ESTADO: ').bold = True
    info.add_run('Sonora\n')
    info.add_run('FECHA: ').bold = True
    info.add_run('06 de mayo de 2026\n')
    info.add_run('ASUNTO: ').bold = True
    info.add_run('Ejes de Acción y Encuesta de Diagnóstico para la Gira Presidencial\n')
    
    # Introduction
    intro = doc.add_paragraph()
    intro.add_run('El presente documento detalla las carencias críticas y demandas estatales que el Magisterio Sonorense (MS) ha identificado como prioritarias. Estas demandas deben servir como base para la consulta a las bases y como sustento técnico para la entrega del pliego petitorio a la Presidencia de la República.').italic = True

    # Demands List
    demands = [
        ("1. INFRAESTRUCTURA DE SUPERVIVENCIA", 
         "* SUBESTACIONES ELÉCTRICAS: Modernización total para soportar carga de AC en zonas de calor extremo.\n"
         "* RED HIDRÁULICA: Garantizar agua potable y enfriadores para el 100% de la matrícula.\n"
         "* MANTENIMIENTO: Presupuesto directo a planteles para impermeabilización y clima."),
        
        ("2. SEGURIDAD E INTEGRIDAD ESCOLAR", 
         "* BLINDAJE PERIMETRAL: Construcción de bardas en zonas de alta incidencia delictiva.\n"
         "* PERSONAL DE APOYO (PAAE): Reposición de plazas de veladores e intendentes.\n"
         "* VIGILANCIA: Vinculación de cámaras escolares al C5i del Estado."),
        
        ("3. JUSTICIA LABORAL Y ADMINISTRATIVA (SEC)", 
         "* CONCEPTO 07: Compactación total de la prima quinquenal al sueldo base.\n"
         "* CARRERA ADMINISTRATIVA: Activación de escalafones para el personal de apoyo.\n"
         "* VIDA CARA: Rezonificación de tabuladores por inflación fronteriza."),
        
        ("4. SALUD Y SEGURIDAD SOCIAL (ISSSTESON)", 
         "* ABASTO TOTAL: Eliminación del sistema de vales; medicinas en farmacia.\n"
         "* CITAS DE ESPECIALIDAD: Reducción drástica de tiempos de espera.\n"
         "* EQUIPAMIENTO REGIONAL: Hospitales dignos en Navojoa, Guaymas y la Sierra.")
    ]
    
    for title_text, body_text in demands:
        h = doc.add_heading(title_text, level=1)
        # Apply dark red to headings
        for run in h.runs:
            run.font.color.rgb = RGBColor(139, 0, 0)
        doc.add_paragraph(body_text)

    # Footer
    footer = doc.add_paragraph('\n\n"Por un Magisterio de Propuesta y Profesionalismo"\n')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('CORE MAGISTERIO SONORENSE').bold = True

    # Save
    output_path = r'c:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Catalogo_Demandas_Estatales_Sonora_2026.docx'
    doc.save(output_path)
    print(f"Documento generado exitosamente en: {output_path}")

if __name__ == "__main__":
    create_demands_doc()
