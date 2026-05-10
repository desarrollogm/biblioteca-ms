import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_doc_medios(titulo_doc, subtitulo, secciones, nombre_archivo):
    doc = Document()
    
    # Estilo base
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    # Titulo Principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo_doc)
    r.font.name = 'Arial'
    r.font.size = Pt(16)
    r.font.bold = True
    
    # Subtitulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitulo)
    r.font.name = 'Arial'
    r.font.size = Pt(12)
    r.font.italic = True
    
    doc.add_paragraph() # Espacio
    
    for seccion in secciones:
        # Titulo Seccion (Arial 14)
        p = doc.add_paragraph()
        r = p.add_run(seccion['titulo'])
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        
        # Contenido (Arial 12)
        for parrafo in seccion['parrafos']:
            if parrafo.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                r = p.add_run(parrafo[2:])
                r.font.name = 'Arial'
                r.font.size = Pt(12)
            elif parrafo.startswith('**') and '**' in parrafo[2:]:
                # Manejo basico de negritas al inicio
                partes = parrafo.split('**')
                p = doc.add_paragraph()
                for i, txt in enumerate(partes):
                    if not txt: continue
                    r = p.add_run(txt)
                    r.font.name = 'Arial'
                    r.font.size = Pt(12)
                    if i % 2 != 0: # Impar significa que estaba entre **
                        r.font.bold = True
            else:
                p = doc.add_paragraph()
                r = p.add_run(parrafo)
                r.font.name = 'Arial'
                r.font.size = Pt(12)
        doc.add_paragraph() # Espacio final de seccion
        
    out_dir = r"C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Estrategia_Medios_15_Mayo"
    doc.save(os.path.join(out_dir, nombre_archivo))


# ================= DATA GENERAL KRISTIAN =================
gen_data = [
    {"titulo": "1. OBJETIVO DEL 15 DE MAYO", "parrafos": [
        "Posicionar al Magisterio Sonorense (MS) en radio, prensa y TV como la ÚNICA alternativa seria y técnica.",
        "**El Contraste a Vender:** Mientras el SNTE sortea carros (pan y circo) y la CNTE propone parar labores en pleno Mundial (desgaste social), el MS está en la Cámara de Diputados peleando jurídicamente la PIP y el fin del robo de las 10 UMAs."
    ]},
    {"titulo": "2. CÓMO PITCHEAR (VENDER) LA NOTA AL PERIODISTA", "parrafos": [
        "Profe Kristian, este es el mensaje de WhatsApp que le mandarás a los productores de noticieros (Proyecto Puente, Uniradio, etc.) el 13 o 14 de mayo:",
        "**Guion de WhatsApp:**",
        "'Hola [Nombre del periodista]. Este 15 de mayo los maestros no tenemos nada que celebrar. El ISSSTE nos está robando más de 50 mil pesos mensuales a nuestra generación por el tope de las 10 UMAs. Mientras otros sindicatos hacen rifas o huelgas, el Magisterio Sonorense ya metió la iniciativa de la Pensión Intergeneracional Protegida (PIP) en el Senado. Tenemos los expedientes y los números fríos. ¿Nos das 10 minutos en cabina para explicarle a Sonora la verdadera crisis magisterial?'"
    ]},
    {"titulo": "3. MEDIA TRAINING (MANEJO DE PREGUNTAS TRAMPA)", "parrafos": [
        "**Trampa 1:** '¿Ustedes son una copia de la CNTE o van a hacer huelga?'",
        "**Respuesta (Técnica del Puente):** 'Nosotros respetamos las manifestaciones, PERO el MS es una organización técnica. Nosotros no paramos escuelas, nosotros arrastramos el lápiz. Ya presentamos la iniciativa en el Senado con corrida financiera.'",
        "**Trampa 2:** '¿Por qué quieren dividir al SNTE Sección 28?'",
        "**Respuesta:** 'No dividimos, evolucionamos. El SNTE oficial se estancó en aplaudir y rifar carros. Nosotros representamos a la base que exige resultados legislativos, no aplausos al gobierno.'"
    ]},
    {"titulo": "4. SOUNDBITES (FRASES DE 10 SEGUNDOS PARA TV/RADIO)", "parrafos": [
        "- 'El 15 de mayo no queremos rifas, queremos que nos quiten el tope de las 10 UMAs.'",
        "- 'Un trabajador del IMSS se jubila con 89 mil pesos; al maestro que educó a Sonora lo topan en 35 mil.'",
        "- 'El MS no hace grilla, hace propuestas con viabilidad financiera.'"
    ]}
]

# ================= DATA CAJEME (KRISTIAN/PALOMA) =================
caj_data = [
    {"titulo": "1. EL ENFOQUE PARA OBREGÓN (CAJEME)", "parrafos": [
        "Cajeme es el bastión histórico de la 'Generación Afectada'. Los maestros que construyeron el sistema educativo del sur.",
        "En la entrevista local, el foco no es pelear, es dignificar a los maestros que están a punto de jubilarse y se toparán con el muro de cristal de las 10 UMAs."
    ]},
    {"titulo": "2. MENSAJE CLAVE EN CABINA", "parrafos": [
        "**Mtra. Paloma / Profe Kristian:**",
        "'En Cajeme tenemos miles de maestros que entregaron su salud por 30 años en las aulas. Y hoy, la Ley del ISSSTE de 2007 les da la espalda.'",
        "'Mientras la inflación en Cajeme sube, a nosotros nos dicen que lo máximo a lo que podemos aspirar son 35,662 pesos mensuales. La PIP (Pensión Intergeneracional Protegida) es el mecanismo solidario que nosotros ya propusimos en México para romper ese tope sin quebrar al estado.'"
    ]},
    {"titulo": "3. PREGUNTA LOCAL", "parrafos": [
        "Si preguntan en la radio de Obregón: '¿Y qué va a pasar este Día del Maestro aquí?'",
        "**Respuesta:** 'Vamos a seguir recolectando las firmas de la base. El MS no va a los festejos oficiales porque hoy la urgencia es defender nuestro patrimonio. Hacemos un llamado a todos los maestros del Valle del Yaqui a sumarse a la iniciativa técnica de la PIP.'"
    ]}
]

# ================= DATA NOGALES (GONZALO) =================
nog_data = [
    {"titulo": "1. EL ENFOQUE PARA NOGALES", "parrafos": [
        "Profe Gonzalo, tu misión en los medios de Nogales es visibilizar la brutalidad del costo de vida fronterizo contrastado con la pensión del ISSSTE.",
        "Tu bandera exclusiva es la **UMA-Fronteriza**."
    ]},
    {"titulo": "2. MENSAJE CLAVE EN CABINA", "parrafos": [
        "'En Nogales pagamos renta en dólares y la canasta básica está inflada por la frontera. El gobierno federal reconoce esto con la Zona Libre de la Frontera Norte, pero el ISSSTE nos ignora.'",
        "'Es absurdo que el tope de pensión sea el mismo en Nogales que en una comunidad del sur del país donde la vida cuesta la mitad. El MS ya puso en la mesa del Senado la creación de la UMA-Fronteriza para ajustar los tabuladores a nuestra realidad económica.'"
    ]},
    {"titulo": "3. EL ATAQUE INDIRECTO (A LA CNTE)", "parrafos": [
        "Si te preguntan por las marchas de otros grupos:",
        "'Respetamos a la coordinadora, pero marchar por marchar en Nogales no cambia la ley en México. Nosotros llevamos un documento técnico que demuestra la viabilidad de la UMA-Fronteriza. Eso es ser un sindicato inteligente.'"
    ]}
]

# ================= DATA HERMOSILLO (CHAYO) =================
hillo_data = [
    {"titulo": "1. EL ENFOQUE PARA HERMOSILLO", "parrafos": [
        "Profe Chayo, en la capital están los noticieros más pesados. Tu enfoque es la Rezonificación Metropolitana.",
        "Debes conectar con la clase media de Hermosillo, que entiende que la ciudad está carísima."
    ]},
    {"titulo": "2. MENSAJE CLAVE EN CABINA", "parrafos": [
        "'Hermosillo creció, se modernizó, pero el tabulador de la SEP se quedó en el pasado. Hoy una renta no baja de 10 mil pesos y los servicios de luz en verano nos ahorcan.'",
        "'El Magisterio Sonorense pide justicia laboral: Rezonificación Metropolitana. No estamos pidiendo lujos, pedimos que nuestro salario y futura pensión estén indexados a la inflación urbana de Hermosillo.'"
    ]},
    {"titulo": "3. SOUNDBITE PARA RADIO LOCAL", "parrafos": [
        "'El 15 de mayo el SNTE va a rifar un carro para 30 mil maestros. Nosotros queremos asegurar que los 30 mil maestros puedan pagar la gasolina de su carro todos los días. Esa es la diferencia del MS.'"
    ]}
]

# ================= DATA GUAYMAS-EMPALME (JESUS) =================
guaymas_data = [
    {"titulo": "1. EL ENFOQUE PARA GUAYMAS / EMPALME", "parrafos": [
        "Profe Jesús, la región portuaria necesita datos duros. Tu enfoque es exhibir el robo silencioso de los Quinquenios (Concepto 07).",
        "En la radio (ej. FM 105), diles a los radioescuchas que saquen su talón."
    ]},
    {"titulo": "2. MENSAJE CLAVE EN CABINA", "parrafos": [
        "'Le pido a todos los maestros de Guaymas y Empalme que nos escuchan, que busquen el Concepto 07 en su talón FONE.'",
        "'Nos están robando quincena tras quincena porque lo tienen desindexado del sueldo base. Esto no solo nos afecta hoy, sino que condena nuestra jubilación. El MS trae la ruta jurídica para compactarlo. No son promesas, son fallos legales.'"
    ]},
    {"titulo": "3. EL LLAMADO A LA ACCIÓN", "parrafos": [
        "Cierra la entrevista invitando a la organización:",
        "'El MS tiene las puertas abiertas para todo maestro federalizado de la región costa que quiera sumarse a la demanda de la compactación de los quinquenios. Este 15 de mayo, infórmate y defiende tu dinero.'"
    ]}
]

# ================= DATA NAVOJOA (JUAN PABLO) =================
nav_data = [
    {"titulo": "1. EL ENFOQUE PARA NAVOJOA", "parrafos": [
        "Profe Juan Pablo, el Valle del Mayo tiene una fuerte identidad sindical, pero a menudo se sienten abandonados por la cúpula de Hermosillo.",
        "Tu enfoque es la representación regional fuerte y la PIP."
    ]},
    {"titulo": "2. MENSAJE CLAVE EN CABINA", "parrafos": [
        "'En Navojoa, los maestros federalizados vemos cómo la cúpula sindical de la Sección 28 se queda en Hermosillo haciendo eventos. Nosotros estamos en las escuelas, sintiendo el golpe de la inflación.'",
        "'La PIP no es una idea al aire, es un documento que el MS construyó con expertos financieros. A los compañeros del sur les digo: no podemos dejarle nuestra jubilación a la suerte ni a las marchas de la CNTE, necesitamos fuerza técnica y legal.'"
    ]},
    {"titulo": "3. EL POSICIONAMIENTO", "parrafos": [
        "Haz que el MS se vea como el protector de la base del sur:",
        "'Este 15 de mayo, el verdadero reconocimiento al maestro de Navojoa sería garantizarle un retiro digno y tumbar la barrera de las 10 UMAs. Mientras eso no pase, el MS seguirá en pie de lucha legal.'"
    ]}
]

base_path = r"C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Estrategia_Medios_15_Mayo"

crear_doc_medios("MANUAL DE VOCERÍA Y MEDIOS: 15 DE MAYO", "Documento Estratégico General - Uso exclusivo: Profe Kristian", gen_data, "Manual_General_Medios_Kristian.docx")
crear_doc_medios("GUION DE ENTREVISTA: CAJEME (OBREGÓN)", "Voceros: Profe Kristian / Mtra. Paloma", caj_data, "Entrevista_Cajeme_Paloma_Kristian.docx")
crear_doc_medios("GUION DE ENTREVISTA: NOGALES", "Vocero: Profe Gonzalo", nog_data, "Entrevista_Nogales_Gonzalo.docx")
crear_doc_medios("GUION DE ENTREVISTA: HERMOSILLO", "Vocero: Profe Chayo", hillo_data, "Entrevista_Hermosillo_Chayo.docx")
crear_doc_medios("GUION DE ENTREVISTA: GUAYMAS/EMPALME", "Vocero: Profe Jesús", guaymas_data, "Entrevista_Guaymas_Jesus.docx")
crear_doc_medios("GUION DE ENTREVISTA: NAVOJOA", "Vocero: Profe Juan Pablo", nav_data, "Entrevista_Navojoa_JuanPablo.docx")

print("Se generaron correctamente los 6 documentos de Estrategia de Medios en Arial 12/14.")
