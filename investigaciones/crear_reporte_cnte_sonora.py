from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilo base Arial 12
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

def titulo_principal(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(16); r.font.bold = True

def seccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True

def subseccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True

def parrafo(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12)

def parrafo_bold(doc, etiqueta, texto):
    p = doc.add_paragraph()
    r1 = p.add_run(etiqueta); r1.font.name = 'Arial'; r1.font.size = Pt(12); r1.bold = True
    r2 = p.add_run(texto); r2.font.name = 'Arial'; r2.font.size = Pt(12)

def bala(doc, texto):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(texto); r.font.name = 'Arial'; r.font.size = Pt(12)

def espacio(doc):
    doc.add_paragraph()

def add_table(doc, encabezados, filas):
    table = doc.add_table(rows=len(filas)+1, cols=len(encabezados))
    table.style = 'Table Grid'
    for i, h in enumerate(encabezados):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True; run.font.name = 'Arial'; run.font.size = Pt(12)
    for i, fila in enumerate(filas):
        for j, texto in enumerate(fila):
            cell = table.rows[i+1].cells[j]
            cell.text = texto
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Arial'; run.font.size = Pt(12)
    return table

# ==================================================
# PORTADA
# ==================================================
titulo_principal(doc, 'REPORTE DE INTELIGENCIA')
titulo_principal(doc, 'CNTE SONORA')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Analisis del posicionamiento de la Mtra. Mercedes Flores y su organizacion')
r.font.name = 'Arial'; r.font.size = Pt(13); r.font.bold = True
espacio(doc)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Para uso exclusivo del CORE del Magisterio Sonorense')
r2.font.name = 'Arial'; r2.font.size = Pt(12); r2.italic = True
espacio(doc)
parrafo_bold(doc, 'Fuente primaria: ', 'Facebook CNTE sonora — ID: 61561007156309')
parrafo_bold(doc, 'Fecha de analisis: ', '3 de mayo de 2026')
parrafo_bold(doc, 'Clasificacion: ', 'Inteligencia estrategica — uso interno MS')

doc.add_page_break()

# ==================================================
# PARTE I — FICHA DE IDENTIDAD
# ==================================================
seccion(doc, 'PARTE I — FICHA DE IDENTIDAD DEL ACTIVO DIGITAL')
espacio(doc)
add_table(doc,
    ['Campo', 'Dato'],
    [
        ['Nombre de la pagina', 'CNTE sonora'],
        ['ID Facebook', '61561007156309'],
        ['Tipo de pagina', 'Organizacion sin animo de lucro'],
        ['Seguidores', '634'],
        ['Siguiendo', '28'],
        ['Sede declarada', 'Domicilio conocido, Ciudad Obregon, Sonora, C.P. 85000'],
        ['Base de operacion fisica', 'Local de la UGOCP, Calle Ramon Guzman, Ciudad Obregon'],
        ['Contacto digital', 'Messenger directo via pagina'],
        ['Actividad registrada', 'Octubre 2025 - Mayo 2026 (al menos)'],
    ]
)
espacio(doc)
parrafo(doc, 'NOTA TACTICA: Con solo 634 seguidores, la pagina tiene baja penetracion digital comparada con su presencia fisica en marchas. El MS puede superar este alcance digital rapidamente con contenido dirigido a la base de la Seccion 28 en Sonora.')

doc.add_page_break()

# ==================================================
# PARTE II — LIDERAZGO
# ==================================================
seccion(doc, 'PARTE II — MAPEO DE LIDERAZGO Y ACTORES CLAVE')
espacio(doc)
subseccion(doc, 'Figura central: Mtra. Mercedes Flores')
bala(doc, 'Rol: Lider local de CNTE Sonora.')
bala(doc, 'Presencia visual: Aparece en actos publicos con megafono, encabezando marchas y eventos de movilizacion.')
bala(doc, 'Estilo de liderazgo: Presencia territorial, discurso de confrontacion directa, no de propuesta tecnica.')
bala(doc, 'Debilidad identificada: No genera contenido propio tecnico. Su comunicacion es emocional y de consigna.')
espacio(doc)
subseccion(doc, 'Circulo operativo:')
espacio(doc)
add_table(doc,
    ['Actor', 'Rol observado'],
    [
        ['Rafael Trujano', 'Generador de contenido. Vincula CNTE con sindicatos universitarios (STAUS/STEUS), amplifica hacia la izquierda universitaria.'],
        ['Gustavo Felix', 'Integrante activo del nucleo de coordinacion.'],
        ['Manuel Ramon Mendoza', 'Apoyo y movilizacion de base.'],
        ['Kiabi Jocobi', 'Apoyo y movilizacion de base.'],
        ['Hector Julian Carrizosa Cota (QEPD)', 'Maestro fallecido (marzo 2026), usado como simbolo de memoria historica y elemento emocional de cohesion.'],
    ]
)
espacio(doc)
parrafo(doc, 'DATO CLAVE: La vinculacion de Rafael Trujano con STAUS/STEUS es un dato de inteligencia relevante: la CNTE Sonora esta intentando construir un frente amplio universitario-magisterial en Sonora. Si lo logra, amplia su base potencial de presion mas alla del magisterio puro.')

doc.add_page_break()

# ==================================================
# PARTE III — PLIEGO PETITORIO
# ==================================================
seccion(doc, 'PARTE III — PLIEGO PETITORIO Y POSICIONAMIENTO POLITICO')
espacio(doc)
add_table(doc,
    ['#', 'Demanda', 'Caracter'],
    [
        ['1', 'Abrogacion total de la Ley del ISSSTE 2007', 'Maximalista — sin propuesta de reemplazo'],
        ['2', 'Fuera las AFORES y las UMAs', 'Regresivo — piden salarios minimos como unidad de calculo'],
        ['3', 'Incremento salarial del 100%', 'Inviable presupuestalmente en el corto plazo'],
        ['4', 'Reinstalacion de Mesa CNTE-Presidenta', 'Tactica de presion directa — sin agenda tecnica'],
    ]
)
espacio(doc)
parrafo(doc, 'LECTURA TACTICA: Las 4 demandas son maximalistas y sin documentacion tecnica de respaldo. Esta es la mayor debilidad de la CNTE Sonora frente al MS: exigen pero no proponen modelos alternativos financieramente sustentables.')

doc.add_page_break()

# ==================================================
# PARTE IV — ANALISIS DE PUBLICACIONES
# ==================================================
seccion(doc, 'PARTE IV — ANALISIS DE PUBLICACIONES (CRONOLOGIA)')
espacio(doc)

subseccion(doc, 'Post del 1 de Mayo de 2026, 21:58h — Marcha Nogales')
parrafo_bold(doc, 'Contenido: ', '"CNTE NOGALES PRESENTE!"')
parrafo_bold(doc, 'Engagement: ', '23 reacciones (Me gusta + Me importa), 3 compartidos.')
subseccion(doc, 'Pancartas visibles en fotos:')
bala(doc, '"Claudia, mentiste con la ley del ISSSTE"')
bala(doc, '"Jubilacion para todos"')
bala(doc, '"No a las AFORES"')
parrafo(doc, 'Lectura tactica: Personalizan el ataque en Sheinbaum. Esto confirma que su estrategia de acceso a la mesa presidencial fue por costo politico, no por legitimidad propositiva. El insulto publico dificulta que el gobierno las reciba como interlocutores preferentes.')
espacio(doc)

subseccion(doc, 'Post del 1 de Mayo de 2026, 21:54h — Compartido Seccion 22 Oaxaca')
parrafo(doc, 'Contenido: Comparten comunicado de Red es Oaxaca sobre la Seccion 22 advirtiendo paro laboral si no hay respuestas.')
parrafo(doc, 'Lectura tactica: Confirma sincronizacion nacional con las secciones del sureste. La CNTE Sonora no opera sola — actua como nodo regional de una estrategia coordinada desde CDMX y Oaxaca. Esto les da fuerza nacional pero los hace dependientes de decisiones que se toman fuera de Sonora.')
espacio(doc)

subseccion(doc, 'Post del 30 de Abril de 2026 — Solidaridad con Huelga Universitaria')
parrafo(doc, 'Contenido: Rafael Trujano apoya huelga del STAUS y STEUS (Universidad de Sonora).')
parrafo(doc, 'Lectura tactica: Ampliacion de coalicion hacia el sector universitario. Si unen magisterio + universidades, aumentan su masa critica de presion en Sonora.')
espacio(doc)

subseccion(doc, 'Post del 19 de Octubre de 2025 — Encuentro con Frente Amplio')
parrafo(doc, 'Contenido visual: Fotos de Mtra. Mercedes y companeras en evento del Frente de Pueblos y Fraccionamientos (Atlas), con ropa tipica indigena y punos en alto.')
parrafo(doc, 'Lectura tactica: La CNTE Sonora construye alianzas con movimientos indigenas y sociales — no solo sindicales. Esto les da capital moral pero los aleja del discurso tecnico-laboral que el gobierno federal exige para sentarse a negociar.')
espacio(doc)

subseccion(doc, 'Carteles recurrentes observados en galeria:')
bala(doc, 'Paro Nacional de 48 horas (13 y 14 de noviembre de 2025)')
bala(doc, '"Abrogacion de la Ley del ISSSTE 2007" — pancarta permanente en marchas')
bala(doc, '"EXIGIMOS" + logos CNTE — estilo visual rojo/negro agresivo')
bala(doc, 'Aniversario 45 de la CNTE (diciembre 2025, Morelia)')
bala(doc, 'Referencia a los 43 de Ayotzinapa en fotos de archivo')

doc.add_page_break()

# ==================================================
# PARTE V — ANALISIS ESTRATEGICO
# ==================================================
seccion(doc, 'PARTE V — ANALISIS ESTRATEGICO DE COMUNICACION')
espacio(doc)
subseccion(doc, 'Fortalezas de la CNTE Sonora:')
espacio(doc)
add_table(doc,
    ['Fortaleza', 'Descripcion'],
    [
        ['Presencia territorial', 'Tienen celulas activas en Obregon, Hermosillo y Nogales.'],
        ['Identidad visual fuerte', 'Rojo/negro + consignas directas — facilmente reconocible.'],
        ['Coordinacion nacional', 'Sincronizados con Seccion 22 (Oaxaca) y el CEN CNTE.'],
        ['Alianzas amplias', 'Frente con universidades (STAUS), movimientos indigenas y sociales.'],
        ['Capital emocional', 'Usan la muerte del Mtro. Carrizosa como simbolo de cohesion.'],
        ['Zona fronteriza', 'Presencia en Nogales genera eco binacional que el gobierno teme.'],
    ]
)
espacio(doc)
subseccion(doc, 'Debilidades criticas de la CNTE Sonora:')
espacio(doc)
add_table(doc,
    ['Debilidad', 'Descripcion', 'Oportunidad para el MS'],
    [
        ['634 seguidores', 'Alcance digital minimo para una organizacion que pretende ser representativa.', 'El MS puede superar su presencia digital en semanas.'],
        ['Cero propuesta tecnica', 'No tienen PIP, no tienen modelos actuariales, no tienen corridas financieras.', 'El MS ya tiene documentos presentados en el Congreso.'],
        ['Demandas inviables', '100% de aumento salarial y abrogacion total son imposibles a corto plazo.', 'El MS propone cambios legislativos puntuales y viables.'],
        ['Imagen de confrontacion', '"Claudia mentiste" — personalizar el ataque genera antiparia en el gobierno.', 'El MS se presenta como aliado, no como adversario.'],
        ['Dependencia de movilizacion', 'Sin marcha, no existen mediaticamente.', 'El MS puede operar desde el Congreso sin bloqueos.'],
        ['Sin base frontera tecnica', 'No hablan de UMA-Fronteriza ni de diferenciales geograficos.', 'La UMA-Fronteriza es exclusiva del MS — argumento que ellos no tienen.'],
    ]
)

doc.add_page_break()

# ==================================================
# PARTE VI — CONTRAINTELIGENCIA
# ==================================================
seccion(doc, 'PARTE VI — MATRIZ DE CONTRAINTELIGENCIA')
parrafo(doc, 'Argumentos listos para debate en grupos de WhatsApp, redes sociales y asambleas:')
espacio(doc)

subseccion(doc, 'Ante: "Abrogacion total de la Ley del ISSSTE"')
parrafo(doc, 'Respuesta MS: "La abrogacion sin propuesta de reemplazo deja a 3 millones de trabajadores sin certeza. El MS ya tiene la PIP aprobable manana. Cual es la propuesta de la CNTE?"')
espacio(doc)

subseccion(doc, 'Ante: "Claudia, mentiste"')
parrafo(doc, 'Respuesta MS: "El confrontacionismo no construye acuerdos. El Magisterio Sonorense llego al Congreso con propuestas; la CNTE llega con pancartas de insulto. Cual estrategia creen que logra resultados reales para los maestros?"')
espacio(doc)

subseccion(doc, 'Ante: "No a las UMAs"')
parrafo(doc, 'Respuesta MS: "Exactamente por eso proponemos las 25 UMAs fronterizas — para que Sonora tenga el mismo techo que el IMSS. No pedimos eliminar las UMAs, pedimos igualdad constitucional. Nuestra propuesta ya esta en el Senado."')
espacio(doc)

subseccion(doc, 'Ante: "100% de aumento salarial"')
parrafo(doc, 'Respuesta MS: "El 100% de aumento no esta en ningun presupuesto federal. El MS pelea por lo que si se puede ganar: quinquenios al concepto correcto, pensiones dignas con la PIP y techo de 25 UMAs. Resultados concretos, no promesas."')

doc.add_page_break()

# ==================================================
# PARTE VII — EVALUACION DE RIESGO
# ==================================================
seccion(doc, 'PARTE VII — EVALUACION DE RIESGO PARA EL MS')
espacio(doc)
add_table(doc,
    ['Riesgo', 'Nivel', 'Descripcion'],
    [
        ['CNTE Sonora capture la agenda en Nogales', 'ALTO', 'Ya tienen presencia documentada en la frontera — zona clave del MS.'],
        ['Frente CNTE-STAUS gane masa critica', 'MEDIO', 'Si unen magisterio + universidades en Sonora, su peso politico crece.'],
        ['Reunion del 11 mayo Seccion 7 genere acuerdos que desplacen al MS', 'ALTO', 'Si la Presidenta cede a la CNTE, el MS queda sin argumento de urgencia.'],
        ['MS sea percibido como "moderado" y pierda base', 'MEDIO', 'La base mas radical puede preferir el discurso confrontacional.'],
        ['CNTE Sonora adopte propuesta tecnica del MS', 'BAJO', 'Su identidad ideologica no les permite abrazar propuestas dentro del sistema.'],
    ]
)

doc.add_page_break()

# ==================================================
# PARTE VIII — RECOMENDACIONES
# ==================================================
seccion(doc, 'PARTE VIII — RECOMENDACIONES ACCIONABLES PARA EL CORE MS')
espacio(doc)
subseccion(doc, 'URGENTES — Esta semana:')
espacio(doc)
parrafo_bold(doc, '1. Generar contenido en Nogales antes del 11 de mayo. ',
    'Comunicado o video posicionando la UMA-Fronteriza como la solucion especifica para las zonas fronterizas. Que la base de Nogales sepa que el MS tiene una propuesta para ellos.')
espacio(doc)
parrafo_bold(doc, '2. Aprovechar el momentum del 11 de mayo. ',
    'Emitir un comunicado el mismo dia que la Seccion 7 se reuna con Sheinbaum: "Mientras otros negocian bajo presion, el Magisterio Sonorense ya tiene una propuesta tecnica en el Senado. Pedimos que la mesa del 11 incluya la revision de la PIP y la UMA-Fronteriza."')
espacio(doc)
parrafo_bold(doc, '3. Superar los 634 seguidores de CNTE Sonora en redes. ',
    'Meta alcanzable en 2 semanas con contenido dirigido a las bases de la Seccion 28 en Sonora.')
espacio(doc)

subseccion(doc, 'ESTRATEGICAS — Proximo mes:')
espacio(doc)
parrafo_bold(doc, '4. No atacar a Mtra. Mercedes directamente. ',
    'Atacar la ausencia de propuesta tecnica, no a la persona. "Respetamos la lucha de nuestros companeros de la CNTE. Pero el maestro de Sonora necesita resultados, no marchas."')
espacio(doc)
parrafo_bold(doc, '5. Capitalizar la alianza CNTE-STAUS como argumento de amplificacion propia. ',
    'Si ellos se alian con universidades, el MS puede buscar respaldo de colegios de abogados, contadores y actuarios de Sonora que validen tecnicamente la PIP.')
espacio(doc)
parrafo_bold(doc, '6. Preparar contrapunto para el 15 de mayo (Dia del Maestro). ',
    'Acto propio MS con entrega simbolica de la Carpeta Ejecutiva a medios y legisladores en Sonora, el mismo dia que la CNTE marche.')

doc.add_page_break()

# ==================================================
# CONCLUSION
# ==================================================
seccion(doc, 'CONCLUSION EJECUTIVA')
espacio(doc)
parrafo(doc,
    'La CNTE Sonora bajo la coordinacion de la Mtra. Mercedes Flores opera con una estrategia de presion emocional y confrontacion directa que le ha dado visibilidad pero no resultados concretos. Su pliego petitorio es ideologico y maximalista, sin soporte tecnico ni financiero.')
espacio(doc)
parrafo(doc,
    'Su mayor activo es la coordinacion nacional con otras secciones y su presencia en Nogales. Su mayor debilidad es la ausencia total de propuesta tecnica y su limitada presencia digital (634 seguidores).')
espacio(doc)
parrafo(doc,
    'El MS tiene ventaja competitiva absoluta en el unico terreno que importa para el gobierno federal: la viabilidad tecnica y legislativa. La Carpeta Ejecutiva, la PIP y la UMA-Fronteriza son instrumentos que la CNTE Sonora nunca podra replicar sin traicionar su propia identidad ideologica.')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run('TACTICA RECOMENDADA: No competir con la CNTE en las calles — competir con ella en los despachos. El MS gana donde la CNTE no puede entrar: las mesas tecnicas del Congreso y la SEP.')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True; r.italic = True

espacio(doc)
doc.add_paragraph('_' * 60)
parrafo(doc, 'Reporte elaborado para uso estrategico exclusivo del CORE del Magisterio Sonorense')
parrafo(doc, 'Fecha: 03/05/2026 | Fuente: Facebook CNTE sonora (analisis directo), investigacion multifuente')

# GUARDAR
output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Intel_CNTE_Sonora_Mercedes.docx'
doc.save(output_path)
print(f"LISTO. Guardado en: {output_path}")
