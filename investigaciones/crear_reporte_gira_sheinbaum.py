from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def titulo_principal(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

def subtitulo_portada(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)

def seccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6E)

def subseccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True

def parrafo(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12)

def parrafo_bold(doc, etiqueta, texto):
    p = doc.add_paragraph()
    r1 = p.add_run(etiqueta); r1.font.name = 'Arial'; r1.font.size = Pt(12); r1.bold = True
    r2 = p.add_run(texto);    r2.font.name = 'Arial'; r2.font.size = Pt(12)

def parrafo_nota(doc, texto, color='444444'):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))

def bala(doc, texto):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(texto); r.font.name = 'Arial'; r.font.size = Pt(12)

def espacio(doc):
    doc.add_paragraph()

def add_table(doc, encabezados, filas, hdr_color='1F3A6E'):
    table = doc.add_table(rows=len(filas)+1, cols=len(encabezados))
    table.style = 'Table Grid'
    for i, h in enumerate(encabezados):
        cell = table.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True; run.font.name = 'Arial'; run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, hdr_color)
    for i, fila in enumerate(filas):
        fill = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
        for j, texto in enumerate(fila):
            cell = table.rows[i+1].cells[j]
            cell.text = texto
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Arial'; run.font.size = Pt(11)
            shade_cell(cell, fill)
    return table

# ==============================================================================
# PORTADA
# ==============================================================================
titulo_principal(doc, 'REPORTE DE INTELIGENCIA OPERATIVA URGENTE')
titulo_principal(doc, 'GIRA PRESIDENCIAL SONORA — MAYO 2026')
espacio(doc)
subtitulo_portada(doc, 'Agenda de la Presidenta Sheinbaum + Estrategia de Entrega de Documento')
espacio(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Para uso exclusivo del CORE del Magisterio Sonorense')
r.font.name = 'Arial'; r.font.size = Pt(12); r.italic = True
espacio(doc)
parrafo_bold(doc, 'Clasificacion: ', 'Inteligencia operativa urgente — uso interno MS')
parrafo_bold(doc, 'Fecha de analisis: ', '4 de mayo de 2026')
parrafo_bold(doc, 'Fuentes: ', 'Conferencia presidencial 04/05/2026, elimparcial.com, tribuna.com.mx, meganoticias.mx')

doc.add_page_break()

# ==============================================================================
# ALERTA CRITICA
# ==============================================================================
seccion(doc, 'PARTE I — DATO CRITICO: LA VENTANA SE ABRE EN 4 DIAS')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run('URGENTE: La Presidenta Claudia Sheinbaum estara en Sonora del VIERNES 8 AL DOMINGO 10 DE MAYO DE 2026. '
    'Es la 5a visita a Sonora en su gestion. La ventana para entregar el documento se cierra el domingo 10 a la noche. '
    'HAY QUE ACTIVAR EL PROTOCOLO HOY.')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_page_break()

# ==============================================================================
# PARTE II — FICHA DE LA GIRA
# ==============================================================================
seccion(doc, 'PARTE II — FICHA DE LA GIRA: LO QUE SE SABE')
espacio(doc)
add_table(doc,
    ['Campo', 'Dato'],
    [
        ['Fechas confirmadas', 'Viernes 8, Sabado 9 y Domingo 10 de mayo de 2026'],
        ['Numero de visita', '5a gira de trabajo a Sonora en la gestion Sheinbaum'],
        ['Gobernador anfitrio', 'Alfonso Durazo Montano'],
        ['Eje 1 de agenda', 'Seguimiento a Planes de Justicia para pueblos originarios de Sonora'],
        ['Eje 2 de agenda', 'Inauguracion de hospitales (IMSS Navojoa, 1a piedra Hospital Ures)'],
        ['Eje 3 de agenda', 'Revision de avances en obras carreteras y proyectos hidricos'],
        ['Eje 4 — DATO CLAVE', 'Informes a comunidades del Rio Sonora sobre acuerdos con Grupo Mexico'],
        ['Recorrido tentativo', 'De SUR a NORTE (Navojoa → Ures → Hermosillo, segun Durazo)'],
        ['Itinerario hora a hora', 'NO publicado aun — se conocera 24-48h antes'],
        ['Contexto politico', 'Sheinbaum descarto que la gira este relacionada con coyuntura politica'],
    ]
)
espacio(doc)
parrafo_nota(doc,
    'LECTURA TACTICA: Que Sheinbaum haya descartado la coyuntura politica como motivo NO significa que no pueda recibir el documento. '
    'Significa que el MS debe presentarse como actor tecnico-propositivo, NO como grupo de presion. '
    'El encuadre correcto es: "Venimos a entregar una propuesta, no a exigir."')

doc.add_page_break()

# ==============================================================================
# PARTE III — PUNTOS DE CONTACTO
# ==============================================================================
seccion(doc, 'PARTE III — MAPA DE PUNTOS DE CONTACTO POSIBLES')
espacio(doc)

subseccion(doc, 'Zona Sur: NAVOJOA (viernes 8 o sabado 9)')
parrafo(doc, 'Evento probable: Inauguracion o supervision del Hospital IMSS de Navojoa.')
espacio(doc)
add_table(doc,
    ['Factor', 'Evaluacion'],
    [
        ['Accesibilidad', 'ALTA — actos de inauguracion son abiertos al publico con invitados'],
        ['Publico esperado', 'Trabajadores de salud, autoridades municipales, comunidad'],
        ['Oportunidad MS', 'Presencia de maestros de la zona sur — terreno conocido (Cajeme, Navojoa)'],
        ['Punto de entrega', 'Llegada/salida del evento — zona de medios es el mejor momento'],
        ['Riesgo', 'EMP cierra el perimetro 2-3 horas antes del acto'],
    ]
)
espacio(doc)

subseccion(doc, 'Zona Centro: URES (sabado 9 o domingo 10) — PUNTO MAS RECOMENDADO')
parrafo(doc, 'Evento probable: Colocacion de primera piedra Hospital Regional de Ures + Informe Plan Rio Sonora.')
espacio(doc)
add_table(doc,
    ['Factor', 'Evaluacion'],
    [
        ['Accesibilidad', 'ALTA — actos comunitarios en municipios pequenos son mas abiertos'],
        ['Publico esperado', 'Comunidades del Rio Sonora, lideres indigenas, medios regionales'],
        ['Oportunidad MS', 'Acto con alta carga simbolica + comunidad local = ambiente mas accesible'],
        ['Punto de entrega', 'Zona de medios antes del acto — o acceso tras el presidium'],
        ['Riesgo', 'Menor afluencia puede significar mayor escrutinio a quienes se acercan'],
        ['VENTAJA TACTICA', 'Ures tiene ~7,000 habitantes — el MS puede ser de los pocos grupos organizados presentes'],
    ]
)
espacio(doc)

subseccion(doc, 'Hermosillo — Aeropuerto (punto de llegada/salida)')
add_table(doc,
    ['Factor', 'Evaluacion'],
    [
        ['Accesibilidad', 'BAJA — zona restringida con alto operativo de seguridad EMP + GN'],
        ['Oportunidad', 'Solo via intermediario (legislador o enlace del Gobernador)'],
        ['Recomendacion', 'NO intentar contacto directo en aeropuerto. Usar via institucional'],
    ]
)

doc.add_page_break()

# ==============================================================================
# PARTE IV — LOS 5 FACTORES TACTICOS
# ==============================================================================
seccion(doc, 'PARTE IV — LOS 5 FACTORES TACTICOS (Leccion de la Seccion 7 en Palenque)')
espacio(doc)
add_table(doc,
    ['Factor Seccion 7', 'Adaptacion para el MS en Sonora'],
    [
        ['1. Inteligencia de agenda', 'Ya la tenemos: Navojoa y Ures son los puntos calientes del 8-10 mayo'],
        ['2. Escala proporcional', 'No se necesitan 200 maestros — se necesitan 5-10 con identificacion y carpeta'],
        ['3. Escalera de intermediarios', 'El gobernador Durazo ES el enlace — un legislador Morena-Sonora puede pedir la reunion'],
        ['4. El documento fisico', 'La Carpeta Ejecutiva de 4 paginas — lista para entregar en mano'],
        ['5. Legitimidad preacumulada', 'La PIP ya esta en el Congreso — ese ES el credencial institucional del MS'],
    ]
)

doc.add_page_break()

# ==============================================================================
# PARTE V — PLAN DE OPERACION: 3 ESCENARIOS
# ==============================================================================
seccion(doc, 'PARTE V — PLAN DE OPERACION: 3 ESCENARIOS')
espacio(doc)

subseccion(doc, 'ESCENARIO A: Via Legislativa (Recomendado — Riesgo Bajo)')
parrafo(doc, 'Ejecutar HOY mismo:')
bala(doc, 'Contactar al Senador o Diputado Federal de Morena-Sonora mas afin a la agenda social. '
    'Pedirle que entregue personalmente la Carpeta Ejecutiva a la Presidenta o a su jefe de asesores durante la gira. '
    'Ofrecerle reconocimiento publico como canal de la causa.')
bala(doc, 'Enviar la Carpeta Ejecutiva a la Oficina de la Presidencia HOY via correo certificado: '
    'Palacio Nacional, Edificio 10, planta baja, Centro Historico, CDMX. '
    'Tambien por correo electronico a la Coordinacion de Enlace Ciudadano de la Presidencia.')
bala(doc, 'Notificar a medios regionales (El Imparcial, Proyecto Puente, Tribuna de Sonora) '
    'que el MS entregara un documento tecnico a la Presidenta durante su visita a Sonora.')
espacio(doc)

subseccion(doc, 'ESCENARIO B: Presencia Directa en Ures (Alto Impacto — Riesgo Medio)')
parrafo(doc, 'Ejecutar el 8-9 de mayo:')
bala(doc, 'Comision de 5 maestros representativos — preferentemente de la zona centro del estado. '
    '1 vocero/vocera con carpeta sellada y portafolio formal. 2-3 maestros con identificacion de escuela. '
    '1 persona para documentar (video/foto para redes).')
bala(doc, 'Posicionarse en zona de medios antes del acto (llegar 3 horas antes). '
    'Al terminar el acto, cuando la comitiva presidencial se desplaza al siguiente vehiculo, '
    'es el momento de mas de 30 segundos disponibles.')
bala(doc, 'El guion de los 30 segundos: "Presidenta, somos maestros federalizados de Sonora. '
    'Tenemos una propuesta tecnica para el ISSSTE — la PIP — que ya esta en el Congreso. '
    'Nos permite entregarle este documento?" — Sencillo. Sin slogan. Sin pancarta. Solo el documento.')
bala(doc, 'Si no hay acceso directo: Entregar a cualquier miembro del EMP con una nota de portada '
    'explicando que es para la Presidenta.')
espacio(doc)

subseccion(doc, 'ESCENARIO C: Conferencia de Prensa en Hermosillo (Riesgo Bajo — Impacto Moderado)')
bala(doc, 'Convocar a medios en Hermosillo el mismo dia de la gira. '
    'Comunicado: "Mientras la Presidenta visita Sonora, el Magisterio Sonorense le entrega formalmente '
    'una propuesta tecnica para reformar el ISSSTE."')
bala(doc, 'Entregar el documento a los medios y distribuirlo en digital a los periodistas que cubren la gira presidencial. '
    'Los reporteros que cubren la gira son el canal mas eficiente para que el documento llegue a los asesores.')
bala(doc, 'Publicar en redes el mismo momento con el hashtag de la visita presidencial.')

doc.add_page_break()

# ==============================================================================
# PARTE VI — EL DOCUMENTO QUE SE ENTREGA
# ==============================================================================
seccion(doc, 'PARTE VI — EL DOCUMENTO QUE SE ENTREGA (Especificaciones)')
espacio(doc)
parrafo_nota(doc, 'El documento que se entrega a la Presidenta debe ser un Resumen Ejecutivo de MAXIMO 4 PAGINAS. '
    'NO la carpeta larga. No hay tiempo ni condicion para que lea 30 paginas.')
espacio(doc)
add_table(doc,
    ['Pagina', 'Contenido'],
    [
        ['PAGINA 1 — PORTADA', 'Logo MS + "Para: Presidenta Claudia Sheinbaum Pardo" + Fecha y lugar de entrega'],
        ['PAGINA 2 — EL PROBLEMA', '3 datos duros: 10 UMAs = techo ISSSTE | 25 UMAs = lo que ya tiene IMSS | ~40,000 maestros federalizados en Sonora afectados'],
        ['PAGINA 3 — LA PROPUESTA', 'La PIP en 5 bullets: que es, por que es viable, por que no requiere reforma constitucional, por que el precedente IMSS la sustenta, que se le pide al Ejecutivo'],
        ['PAGINA 4 — LA PETICION', '"Solicitamos respetuosamente una mesa tecnica con SEP/ISSSTE para revisar la viabilidad de homologar el techo de UMAs al precedente ya existente en el IMSS." + Contacto'],
    ]
)

doc.add_page_break()

# ==============================================================================
# PARTE VII — EVALUACION DE RIESGO
# ==============================================================================
seccion(doc, 'PARTE VII — EVALUACION DE RIESGO OPERATIVO')
espacio(doc)
add_table(doc,
    ['Riesgo', 'Nivel', 'Mitigacion'],
    [
        ['EMP bloquea el acceso al evento', 'ALTO', 'Llegar 3 horas antes, posicionarse en zona de medios, no confrontar'],
        ['Presidenta no puede recibir directamente', 'MEDIO', 'Entregarlo a su jefe de asistentes o equipo de SEP que la acompana'],
        ['Medios no lo cubren', 'MEDIO', 'Lanzar comunicado a medios ANTES del acto — no despues'],
        ['CNTE intenta lo mismo y desplaza al MS', 'ALTO', 'El MS debe estar primero y con mayor profesionalismo visual'],
        ['Gobernador Durazo bloquea el acceso', 'MEDIO', 'Usar via legislativa federal (no estatal) como canal principal'],
    ]
)

doc.add_page_break()

# ==============================================================================
# PARTE VIII — CRONOGRAMA
# ==============================================================================
seccion(doc, 'PARTE VIII — CRONOGRAMA DE ACCION (LAS PROXIMAS 96 HORAS)')
espacio(doc)
add_table(doc,
    ['Hora/Fecha', 'Accion', 'Responsable'],
    [
        ['HOY — 4 mayo, noche', 'Finalizar y cerrar el Resumen Ejecutivo de 4 paginas', 'Equipo tecnico MS'],
        ['HOY — 4 mayo, noche', 'Identificar legislador federal Morena-Sonora para canal A', 'Liderazgo MS'],
        ['5 mayo — manana', 'Enviar carpeta por correo certificado a Palacio Nacional', 'Liderazgo MS'],
        ['5 mayo — manana', 'Contactar legislador federal y solicitar reunion urgente', 'Liderazgo MS'],
        ['5-6 mayo', 'Identificar 5 maestros para comision en Ures (Escenario B)', 'Coordinadores zona centro'],
        ['6-7 mayo', 'Monitorear confirmacion de agenda presidencial hora por hora', 'Equipo inteligencia MS'],
        ['7 mayo — noche', 'Briefing final con comision — guion de 30 segundos ensayado', 'Vocero principal'],
        ['8 mayo — madrugada', 'Desplazamiento comision al punto de mayor probabilidad (Ures o Navojoa)', 'Comision MS'],
        ['8-10 mayo', 'Ejecutar entrega + documentar para redes + notificar a medios', 'Toda la comision'],
    ]
)

doc.add_page_break()

# ==============================================================================
# CONCLUSION
# ==============================================================================
seccion(doc, 'CONCLUSION EJECUTIVA')
espacio(doc)
parrafo(doc,
    'La visita de la Presidenta Sheinbaum a Sonora del 8 al 10 de mayo es la ventana tactica mas importante '
    'que ha tenido el MS en su historia reciente. Los puntos de mayor probabilidad de contacto son Ures '
    '(acto comunitario en municipio pequeno, alta simbologia) y Navojoa (inauguracion hospitalaria).')
espacio(doc)
parrafo(doc,
    'El MS no necesita bloquear nada. Necesita llegar primero, con el documento correcto, y con los maestros correctos. '
    'Cinco maestros con portafolio y propuesta tecnica valen mas que 500 con pancartas.')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run(
    'TACTICA MAESTRA: Si el documento llega a manos de la Presidenta (o de su equipo) en Sonora, '
    'y el mismo dia sale en medios que el MS entrego una propuesta tecnica, el MS habra ganado visibilidad '
    'presidencial sin un solo bloqueo, sin un solo paro, y con la legitimidad tecnica intacta.')
r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True; r.italic = True

espacio(doc)
doc.add_paragraph('_' * 65)
parrafo(doc, 'Reporte elaborado para uso estrategico exclusivo del CORE del Magisterio Sonorense')
parrafo(doc, 'Fecha: 04/05/2026 | Fuentes: Conferencia presidencial 04/05/2026, elimparcial.com, tribuna.com.mx, meganoticias.mx')

# GUARDAR
output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Reporte_Intel_Gira_Sheinbaum_Sonora_Mayo2026.docx'
doc.save(output_path)
print(f'LISTO. Guardado en: {output_path}')
