import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('ANÁLISIS ESTRATÉGICO: ACUERDOS ESTATALES 4 DE MAYO', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('Documento de Trabajo para el CORE del MS\n').italic = True
p.add_run('Fecha de Análisis: ').bold = True
p.add_run('5 de mayo de 2026\n')
p.add_run('Clasificación: ').bold = True
p.add_run('Inteligencia Operativa / Uso Reservado')

doc.add_heading('1. VALORACIÓN GENERAL', level=2)
doc.add_paragraph('Los acuerdos alcanzados muestran una madurez operativa significativa. El Magisterio Sonorense (MS) está consolidando una "Estrategia de Pinza" que combina la presión de base con el cabildeo institucional de alto nivel.')

doc.add_heading('2. ANÁLISIS DE EJES TÁCTICOS', level=2)

doc.add_heading('2.1 La Gira Presidencial (8-10 de Mayo)', level=3)
doc.add_paragraph('El deslinde de la CNTE (Punto 2C) es la decisión estratégica más relevante. Preservar la identidad institucional del MS frente a la Presidenta Sheinbaum es vital para mantener la calidad de "interlocutor técnico". Si el MS es percibido como un grupo de choque, se cerrarán las puertas de Palacio Nacional.', style='List Bullet')
doc.add_paragraph('La gestión con el Dip. Pujol y el Sen. Aguilar debe enfocarse en que ellos funjan como escoltas institucionales de la comisión al momento de la entrega.', style='List Bullet')

doc.add_heading('2.2 Crecimiento y Consolidación Regional', level=3)
doc.add_paragraph('Las visitas presenciales a los Centros de Trabajo (C.T.) superan el impacto del WhatsApp, ya que generan compromiso real de base.', style='List Bullet')
doc.add_paragraph('Es necesario mapear las "zonas de silencio" (donde el apoyo es bajo) para evitar que otros actores sindicales ocupen ese vacío político.', style='List Bullet')

doc.add_heading('2.3 La Jornada del Día del Maestro (14-15 de Mayo)', level=3)
doc.add_paragraph('El registro de demandas mediante formulario digital es una "mina de oro" de datos. Permitirá transformar la manifestación en la SEC de un acto de protesta genérico a una entrega masiva de casos específicos con sustento real.', style='List Bullet')

doc.add_heading('3. ACTORES POLÍTICOS Y CABILDEO', level=2)
doc.add_paragraph('Sen. Lorenia Valles: Es la pieza clave. Su respaldo otorga legitimidad oficialista a la PIP.', style='List Bullet')
doc.add_paragraph('César Salazar y Froylán: Representan el canal técnico y operativo. Se debe mantener una comunicación fluida para evitar bloqueos administrativos.', style='List Bullet')

doc.add_heading('4. RECOMENDACIONES ESTRATÉGICAS ("HASTA LA COCINA")', level=2)

p1 = doc.add_paragraph()
r1 = p1.add_run('Prioridad Ures/Navojoa: ')
r1.bold = True
p1.add_run('Cuidar la "foto". Si hay presencia de la CNTE, la comisión del MS debe distanciarse físicamente para que el registro de prensa sea puramente institucional.')

p2 = doc.add_paragraph()
r2 = p2.add_run('Inteligencia de Datos: ')
r2.bold = True
p2.add_run('Procesar la información del formulario del 14 de mayo para generar un "Atlas de Demandas del Magisterio Sonorense". Esto desarmará dialécticamente a las autoridades de la SEC.')

p3 = doc.add_paragraph()
r3 = p3.add_run('Lobbying con Sen. Valles: ')
r3.bold = True
p3.add_run('El argumento debe ser la "Estabilidad Social". El MS se ofrece como el actor que garantiza orden en las escuelas a cambio de una revisión técnica de las pensiones (PIP).')

doc.add_heading('CONCLUSIÓN', level=2)
p_final = doc.add_paragraph()
p_final.add_run('El MS tiene el control de la narrativa. El éxito de mayo depende de la disciplina operativa (no mezclarse con la CNTE) y de la capacidad de convertir el sentimiento de la base en datos duros procesables.').bold = True

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Analisis_Estrategico_Acuerdos_4Mayo.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
