import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_video_analysis_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('INFORME DE ANÁLISIS: DECLARATORIA SEP (YOUTUBE)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Header
    p = doc.add_paragraph()
    p.add_run('REPORTE TÉCNICO: ').bold = True
    p.add_run('Análisis de Declaraciones de Mario Delgado\n')
    p.add_run('FECHA: ').bold = True
    p.add_run('06 de mayo de 2026\n')
    p.add_run('DESTINO: ').bold = True
    p.add_run('CORE Magisterio Sonorense\n')
    
    # Sections
    data = [
        ("1. OBJETIVO DEL VIDEO", 
         "El mensaje central es la confirmación de la reunión nacional del 7 de mayo para oficializar el recorte del calendario escolar. El Secretario Mario Delgado busca socializar la medida antes de su decreto oficial."),
        
        ("2. ACUERDOS IDENTIFICADOS", 
         "* Cierre de ciclo propuesto: 30 de junio de 2026.\n"
         "* Justificantes: Olas de calor extremo y logística del Mundial de Fútbol.\n"
         "* Método: 'Consenso' administrativo con los estados."),
        
        ("3. ANÁLISIS DE INTELIGENCIA MS", 
         "La SEP está claudicando ante la falta de infraestructura. Es una medida cosmética para evitar paros o descontento social durante el Mundial. En Sonora, esto deja al descubierto la falta de subestaciones eléctricas que el MS ha denunciado."),
        
        ("4. ACCIÓN RECOMENDADA", 
         "Publicar de inmediato la postura del MS denunciando la falta de inversión que obliga a estos cierres anticipados. Usar la frase: 'No queremos vacaciones por calor, queremos escuelas con luz y aire'.")
    ]
    
    for title_text, body_text in data:
        h = doc.add_heading(title_text, level=1)
        doc.add_paragraph(body_text)

    # Footer
    footer = doc.add_paragraph('\n"Por la Profesionalización del Magisterio"\n')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('CORE MS').bold = True

    output_path = r'c:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Analisis_Video_Mario_Delgado_SEP_2026.docx'
    doc.save(output_path)
    print(f"Documento generado en: {output_path}")

if __name__ == "__main__":
    create_video_analysis_doc()
