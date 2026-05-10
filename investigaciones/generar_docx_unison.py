import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Configurar estilo Normal a Arial 12
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Arial'
font.size = Pt(12)

def add_heading(doc, text, level):
    if level == 0:
        h = doc.add_heading(text, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        h = doc.add_heading(text, level=level)
        
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    return h

# Título
add_heading(doc, 'INFORME DE INTELIGENCIA ESTRATÉGICA', 0)

# Metadatos
p = doc.add_paragraph()
p.add_run('OBJETIVO: ').bold = True
p.add_run('Análisis del Comunicado Oficial del Gobernador de Sonora (07 Mayo 2026) sobre el Levantamiento del Paro en la UNISON.\n')
p.add_run('PARA: ').bold = True
p.add_run('Coordinación General Magisterio Sonorense (MS)\n')
p.add_run('FECHA: ').bold = True
p.add_run('08 de mayo de 2026')

doc.add_paragraph()

# Sección 1
add_heading(doc, '1. ANÁLISIS DEL COMUNICADO OFICIAL', level=1)
doc.add_paragraph('El documento es un comunicado oficial emitido por el Gobernador Alfonso Durazo Montaño el 7 de mayo de 2026, celebrando el fin del conflicto laboral (paro/huelga) en la Universidad de Sonora (UNISON).')

p1 = doc.add_paragraph(style='List Bullet')
p1.add_run('Tono Institucional y Conciliador: ').bold = True
p1.add_run('El Gobernador felicita la "voluntad mostrada" tanto por el Sindicato de Trabajadores Académicos (STAUS/STEUS) como por la administración universitaria. Utiliza palabras clave como "diálogo", "responsabilidad institucional" y "construcción de acuerdos".')

p2 = doc.add_paragraph(style='List Bullet')
p2.add_run('Narrativa del "Interés Superior": ').bold = True
p2.add_run('El texto subraya que la educación está por encima de las diferencias y que se privilegió a los miles de estudiantes sonorenses para que continúen sus actividades académicas y de investigación con normalidad.')

p3 = doc.add_paragraph(style='List Bullet')
p3.add_run('Postura del Gobierno del Estado: ').bold = True
p3.add_run('Cierra afirmando que mantendrá "las puertas abiertas" para construir soluciones. Se posiciona como un garante de la estabilidad y el bienestar, pero cuidando de mostrar que la resolución fue entre el Sindicato y la Administración, respetando la autonomía universitaria.')

doc.add_paragraph()

# Sección 2
add_heading(doc, '2. CONTEXTO E INVESTIGACIÓN (EL TRASFONDO)', level=1)
doc.add_paragraph('Tras revisar la cobertura de prensa y el entorno sociopolítico del 7 y 8 de mayo de 2026, el trasfondo de este comunicado revela lo siguiente:')

doc.add_paragraph('A. Presión por la Gira Presidencial').runs[0].bold = True
doc.add_paragraph('El levantamiento del paro no es una casualidad de calendario. Ocurre exactamente un día antes de la gira de la Presidenta Claudia Sheinbaum por Sonora (8 al 10 de mayo). El Gobierno del Estado necesitaba urgentemente "limpiar la casa" y desactivar cualquier foco rojo de protesta (como una huelga en la máxima casa de estudios) que pudiera empañar la visita o generar bloqueos en los eventos de la Presidenta en Hermosillo.')

doc.add_paragraph('B. Intervención Estatal Discreta pero Firme').runs[0].bold = True
doc.add_paragraph('Aunque el comunicado elogia el diálogo entre el Sindicato y la Rectoría, en la práctica, el Gobierno del Estado tuvo que intervenir inyectando recursos extraordinarios o forzando políticamente a la Rectoría a ceder a las demandas salariales/prestacionales del sindicato para garantizar la paz social previa a la llegada de Sheinbaum.')

doc.add_paragraph()

# Sección 3
add_heading(doc, '3. IMPLICACIONES Y OPORTUNIDADES PARA EL MAGISTERIO SONORENSE (MS)', level=1)

p_op1 = doc.add_paragraph(style='List Bullet')
p_op1.add_run('La Gira Presidencial como Acelerador de Soluciones: ').bold = True
p_op1.add_run('El levantamiento de la huelga en la UNISON demuestra que el Gobierno del Estado tiene los medios (financieros y políticos) para resolver conflictos laborales rápido cuando hay presión política de alto nivel (la visita de la Presidenta). El MS debe registrar este modus operandi. Cuando hay eventos de esta magnitud, la capacidad de negociación del gobierno local se flexibiliza enormemente por miedo a los "periodicazos" nacionales.')

p_op2 = doc.add_paragraph(style='List Bullet')
p_op2.add_run('El Contraste Educativo: ').bold = True
p_op2.add_run('El Gobernador dice que "la educación siempre debe estar por encima de las diferencias". El MS puede usar esta misma frase del Gobernador para exigirle congruencia respecto al magisterio estatal (Sección 28/54). Si la educación es prioridad, el gobierno no puede seguir solapando las omisiones de la dirigencia sindical charra que afectan a los maestros de educación básica.')

p_op3 = doc.add_paragraph(style='List Bullet')
p_op3.add_run('Posicionamiento Táctico: ').bold = True
p_op3.add_run('El MS no es un sindicato que cierre universidades, sino uno que audita fondos y propone leyes (Ley PIP). Podemos publicar un posicionamiento felicitando a los compañeros de la UNISON por su logro, pero recordando que en educación básica y media superior en Sonora, la "voluntad y el diálogo" del SNTE oficial solo sirve para entregar conquistas laborales, no para rescatarlas.')

# Guardar documento
doc.save(r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Informe_Paro_Unison_Mayo2026.docx')
print("Documento DOCX generado exitosamente.")
