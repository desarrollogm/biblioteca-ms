import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Title
title = doc.add_heading('ANÁLISIS FORENSE DE INTELIGENCIA SINDICAL', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Documento: "¡La falsa UNIDAD del SNTE!"', level=1)

p = doc.add_paragraph()
p.add_run('Clasificación: ').bold = True
p.add_run('INTELIGENCIA ESTRATÉGICA — USO INTERNO MS\n')
p.add_run('Fecha de análisis: ').bold = True
p.add_run('5 de mayo de 2026\n')
p.add_run('Elaborado por: ').bold = True
p.add_run('Área de Inteligencia — Magisterio Sonorense\n')
p.add_run('Fuente analizada: ').bold = True
p.add_run('Panfleto en circulación digital — hashtag #MagisterioNacional en pie de lucha')

doc.add_heading('1. IDENTIFICACIÓN DEL DOCUMENTO', level=2)
doc.add_heading('1.1 ¿De dónde viene este texto?', level=3)
doc.add_paragraph('El documento se difundió bajo el hashtag #MagisterioNacional en pie de lucha. Los colores rojo y negro son la identidad visual de la CNTE (Coordinadora Nacional de Trabajadores de la Educación). Por tanto:')

p = doc.add_paragraph()
r = p.add_run('⚠️ CONCLUSIÓN DE ORIGEN: Este documento NO fue emitido por el MS ni por un actor neutral. Es un panfleto producido por una facción alineada a la CNTE o a bases disidentes anti-SNTE. Su objetivo es demoler la credibilidad del SNTE ante el magisterio de base.')
r.bold = True

doc.add_heading('1.2 Encuadre del contexto', level=3)
doc.add_paragraph('El texto fue producido después del 1° de mayo de 2026, en referencia directa a las marchas nacionales convocadas por el SNTE bajo el slogan "UNIDOS SOMOS MÁS FUERTES", y en respuesta a la propuesta de pensiones presentada por el dirigente nacional Alfonso Cepeda Salas durante la Semana Nacional de la Seguridad Social.')

doc.add_heading('2. DIAGNÓSTICO TÁCTICO: ¿QUÉ INTENTA HACER ESTE DOCUMENTO?', level=2)
doc.add_paragraph('El texto tiene una sola misión táctica: demoler la legitimidad del SNTE ante su propia base magisterial, mediante cuatro ataques articulados:')

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Argumento'
hdr_cells[1].text = 'Contenido'
hdr_cells[2].text = 'Objetivo táctico'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

row1 = table.rows[1].cells
row1[0].text = 'La unidad es falsa'
row1[1].text = 'Las marchas del 1° de mayo son espectáculo, no unidad real'
row1[2].text = 'Deslegitimar la movilización del SNTE'

row2 = table.rows[2].cells
row2[0].text = 'Cepeda como traidor'
row2[1].text = 'Cepeda evitó exigir la abrogación de la Ley ISSSTE 2007'
row2[2].text = 'Señalar al líder como cómplice del status quo'

row3 = table.rows[3].cells
row3[0].text = '"Utopía" como capitulación'
row3[1].text = 'Cepeda llamó "utopía" a la abrogación de la ley'
row3[2].text = 'Evidenciar que el líder mató la aspiración central de la base'

row4 = table.rows[4].cells
row4[0].text = 'El Mundial como cortina'
row4[1].text = 'El contexto del Mundial inhibe movilizaciones reales'
row4[2].text = 'Insinuar coordinación gobierno-SNTE para desactivar presión'

doc.add_paragraph()

doc.add_heading('3. ANATOMÍA DE LA PROPUESTA CEPEDA SALAS', level=2)
doc.add_paragraph('Este es el hallazgo más relevante para el MS. El documento revela los contornos de la propuesta de Alfonso Cepeda Salas presentada ante la Semana Nacional de Seguridad Social:')

doc.add_heading('Lo que Cepeda SÍ hizo:', level=3)
doc.add_paragraph('✅ Presentó una "propuesta-idea" en materia de pensiones\n✅ Generó apariencia de acción y gestión ante el magisterio\n✅ Se posicionó en el espacio de la discusión previsional')

doc.add_heading('Lo que Cepeda NO hizo:', level=3)
doc.add_paragraph('❌ No exigió la abrogación o derogación de la Ley del ISSSTE 2007\n❌ No presentó un modelo financiero alternativo\n❌ No abordó la demanda central del magisterio de base\n❌ Llamó a la abrogación de la ley una "utopía" — descartándola públicamente')

p = doc.add_paragraph()
r = p.add_run('Diagnóstico de la propuesta Cepeda: Es una propuesta que administra el problema sin resolverlo. Mantiene el statu quo de 2007 con ajustes cosméticos, exactamente como la reforma que el magisterio rechaza. Al llamar "utopía" a la abrogación, Cepeda marca el techo de lo negociable desde la cúpula sindical — un techo que excluye las demandas de fondo.')
r.bold = True

doc.add_heading('4. MAPA DE FRACTURAS: EL ESCENARIO REAL DEL SINDICALISMO MAGISTERIAL', level=2)

doc.add_heading('4.1 La fractura SNTE — base propia', level=3)
doc.add_paragraph('El hecho de que este panfleto circule con fuerza revela que el SNTE está siendo atacado desde dentro. Su base no le cree. Las contradicciones son públicas:')
doc.add_paragraph('• El SNTE moviliza en calle pero evita la demanda central (Ley ISSSTE 2007)', style='List Bullet')
doc.add_paragraph('• El discurso de "unidad" choca con la realidad de secciones inconformes', style='List Bullet')
doc.add_paragraph('• Cepeda no puede presentarse ante la Presidenta como voz unificada del magisterio', style='List Bullet')

doc.add_heading('4.2 La fractura CNTE — propuesta vs. presión', level=3)
doc.add_paragraph('La CNTE produce este documento como crítica al SNTE, pero no ofrece alternativa técnica. El panfleto:')
doc.add_paragraph('• Destruye al SNTE con argumentos válidos', style='List Bullet')
doc.add_paragraph('• No menciona ningún modelo de pensión alternativo', style='List Bullet')
doc.add_paragraph('• No menciona la PIP (Pensión Intergeneracional Protegida)', style='List Bullet')

doc.add_heading('5. BALAS ARGUMENTATIVAS DERIVADAS (War Room MS)', level=2)
doc.add_paragraph('Bala 1 — Frente a la base del SNTE:').bold = True
doc.add_paragraph('"Hasta la propia base del SNTE denuncia que Cepeda Salas llamó \'utopía\' a la abrogación de la Ley ISSSTE. Mientras ellos se pelean entre sí, el Magisterio Sonorense ya entregó en el Senado la PIP: un modelo actuarial concreto y financiado. Nosotros no pedimos utopías. Nosotros legislamos."')

doc.add_paragraph('Bala 2 — Frente a la CNTE:').bold = True
doc.add_paragraph('"La CNTE tiene razón en que la \'unidad\' del SNTE es falsa. Pero la solución no es solo señalar al traidor — es presentar la alternativa. El MS ya la tiene: se llama PIP y ya está en el Congreso."')

doc.add_paragraph('Bala 3 — Para medios y autoridades:').bold = True
doc.add_paragraph('"El SNTE evita la demanda central. La CNTE demanda sin propuesta. El Magisterio Sonorense es el único actor que presenta al mismo tiempo la demanda Y el modelo de solución financiera."')

doc.add_heading('6. CONCLUSIÓN EJECUTIVA', level=2)
p = doc.add_paragraph()
r = p.add_run('El documento "¡La falsa UNIDAD del SNTE!" es un activo de inteligencia de alto valor para el Magisterio Sonorense. Confirma que el SNTE no representa a su base en la demanda central, que la CNTE tampoco tiene propuesta técnica, y que el MS es el único actor con propuesta técnica presentada formalmente ante el Poder Legislativo. El campo está vacío. La ventana es ahora — antes de la gira presidencial del 8 al 10 de mayo de 2026.')
r.bold = True

output_path = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Analisis_Tactica_SNTE_FalsaUnidad_Mayo2026.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
