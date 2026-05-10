from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)

def titulo_principal(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(16); r.font.bold = True

def seccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(14); r.font.bold = True

def subseccion(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True

def parrafo(doc, texto):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.font.name = 'Arial'; r.font.size = Pt(12)

def parrafo_bold(doc, etiqueta, texto):
    p = doc.add_paragraph()
    r1 = p.add_run(etiqueta); r1.font.name='Arial'; r1.font.size=Pt(12); r1.bold=True
    r2 = p.add_run(texto); r2.font.name='Arial'; r2.font.size=Pt(12)

def bala(doc, texto):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(texto); r.font.name='Arial'; r.font.size=Pt(12)

def espacio(doc):
    doc.add_paragraph()

def tabla(doc, encabezados, filas):
    t = doc.add_table(rows=len(filas)+1, cols=len(encabezados))
    t.style = 'Table Grid'
    for i, h in enumerate(encabezados):
        c = t.rows[0].cells[i]; c.text = h
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.name = 'Arial'
        c.paragraphs[0].runs[0].font.size = Pt(12)
    for i, fila in enumerate(filas):
        for j, txt in enumerate(fila):
            c = t.rows[i+1].cells[j]; c.text = txt
            c.paragraphs[0].runs[0].font.name = 'Arial'
            c.paragraphs[0].runs[0].font.size = Pt(12)

# ─── PORTADA ───────────────────────────────────────────────
titulo_principal(doc, 'ESTRATEGIA DE UNIFICACION DE BASE')
titulo_principal(doc, 'MAGISTERIO SONORENSE')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"Una Propuesta para Cada Maestro"')
r.font.name='Arial'; r.font.size=Pt(13); r.font.bold=True; r.italic=True
espacio(doc)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Documento Estrategico para el CORE MS — Mayo 2026')
r2.font.name='Arial'; r2.font.size=Pt(12); r2.italic=True
espacio(doc)
parrafo_bold(doc,'Clasificacion: ','Estrategia interna — CORE MS')
parrafo_bold(doc,'Concepto central: ','Segmentacion territorial con demanda "genetica" por zona + PIP como eje unificador')

doc.add_page_break()

# ─── EL PROBLEMA QUE RESUELVE ──────────────────────────────
seccion(doc, 'EL PROBLEMA QUE RESUELVE ESTA ESTRATEGIA')
espacio(doc)
parrafo(doc,
    'Las grandes organizaciones sindicales (CNTE, SNTE) pierden bases porque lanzan demandas genericas que no conectan con la realidad especifica del maestro de Nogales, ni del maestro de la sierra, ni del que esta a 3 anos de jubilarse en Obregon.')
espacio(doc)
parrafo(doc, 'El MS tiene una ventaja unica: propuestas tecnicas diferenciadas que pueden modularse por zona.')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run('"Cada maestro de Sonora debe encontrar SU demanda en la plataforma del MS."')
r.font.name='Arial'; r.font.size=Pt(12); r.bold=True; r.italic=True
espacio(doc)
parrafo(doc,
    'La PIP es la demanda de todos. Pero el camino para llegar a ella tiene un apellido diferente segun donde vivas y que te duela mas hoy.')

doc.add_page_break()

# ─── ARQUITECTURA ──────────────────────────────────────────
seccion(doc, 'ARQUITECTURA ESTRATEGICA: EL MODELO ARBOL MS')
espacio(doc)
parrafo(doc,
    'La PIP (Pension Intergeneracional Protegida) es la copa del arbol — visible para todos. Cada rama llega a un maestro diferente con su argumento especifico:')
espacio(doc)
tabla(doc,
    ['Zona','Municipios principales','Demanda genetica (la rama)'],
    [
        ['FRONTERA NORTE','Nogales, Agua Prieta, San Luis RC, Sonoyta, Cananea','UMA-Fronteriza'],
        ['CAPITAL','Hermosillo','Rezonificacion economica metropolitana'],
        ['COSTA / PORTUARIA','Guaymas, Empalme','Rezonificacion portuaria + Quinquenios concepto 07'],
        ['SUR AGRICOLA','Ciudad Obregon, Navojoa, Huatabampo','Quinquenios concepto 07 + PIP generacion Decimo Transitorio'],
        ['SIERRA / INDIGENA','Sahuaripa, Yécora, comunidades Yaqui y Mayo','Regularizacion de plazas + bilateralidad USICAMM'],
    ]
)
espacio(doc)
parrafo_bold(doc,'La raiz que une a todos: ',
    'La PIP conecta cada zona con un destino comun: una jubilacion digna, independientemente del municipio donde se haya trabajado.')

doc.add_page_break()

# ─── ZONA 1 ────────────────────────────────────────────────
seccion(doc, 'ZONA 1 — FRONTERA NORTE')
espacio(doc)
subseccion(doc,'Municipios incluidos:')
parrafo(doc,'Nogales · Agua Prieta · San Luis Rio Colorado · Sonoyta · Naco · Santa Cruz · Saric · Altar · Cananea · General Plutarco Elias Calles · Puerto Penasco')
espacio(doc)
subseccion(doc,'Realidad especifica de esta zona:')
bala(doc,'Viven en la Zona Libre de la Frontera Norte (ZLFN) — salario minimo oficial ya diferenciado por el gobierno federal.')
bala(doc,'Costo de vida 30-40% superior al promedio estatal por dinamica binacional con EE.UU.')
bala(doc,'Compiten laboralmente con el sector privado que paga en dolares.')
bala(doc,'Fuga de docentes documentada hacia empleos binacionales.')
espacio(doc)
subseccion(doc,'SU DEMANDA GENETICA: UMA-FRONTERIZA')
parrafo(doc,
    '"Si el gobierno ya reconoce que vivir en la frontera cuesta mas (Zona Libre), por que nuestra pension no lo reconoce? El MS propone que el tope de pension en municipios fronterizos sea de 15 UMAs en el corto plazo y 25 UMAs a mediano plazo — igual que el IMSS."')
espacio(doc)
subseccion(doc,'Mensaje de activacion (WhatsApp):')
p = doc.add_paragraph()
r = p.add_run(
    '"Companero maestro de la frontera: gastas en dolares, te jubilan en pesos de 10 UMAs. '
    'El MS ya metio al Senado la propuesta para que tu pension reconozca el costo de vida donde trabajas. '
    'Eso se llama UMA-Fronteriza. Se parte."')
r.font.name='Arial'; r.font.size=Pt(12); r.italic=True

doc.add_page_break()

# ─── ZONA 2 ────────────────────────────────────────────────
seccion(doc, 'ZONA 2 — CAPITAL Y AREA METROPOLITANA')
espacio(doc)
subseccion(doc,'Municipios incluidos:')
parrafo(doc,'Hermosillo · Bahia de Kino · La Colorada')
espacio(doc)
subseccion(doc,'Realidad especifica de esta zona:')
bala(doc,'Mayor concentracion de docentes federalizados en el estado.')
bala(doc,'Costo de vida metropolitano — vivienda, transporte y servicios al alza constante.')
bala(doc,'Concentracion de docentes con mas anos de servicio, proximos a jubilarse.')
bala(doc,'Mayor nivel educativo y acceso a informacion tecnica — base receptiva a propuestas.')
bala(doc,'Presencia de universidades publicas — zona de influencia politica estrategica.')
espacio(doc)
subseccion(doc,'SU DEMANDA GENETICA: REZONIFICACION ECONOMICA METROPOLITANA')
parrafo(doc,
    '"Hermosillo ya no es una ciudad de salario minimo general. Su costo de vida es el mas alto de Sonora despues de la frontera. El MS exige que los tabuladores de zonas economicas para la SEP reconozcan la realidad de la capital."')
espacio(doc)
subseccion(doc,'Mensaje de activacion (WhatsApp):')
p = doc.add_paragraph()
r = p.add_run(
    '"Maestro de Hermosillo: tu costo de renta, gasolina y servicios subio, pero tu zona economica sigue igual que hace 20 anos. '
    'El MS tiene una propuesta de rezonificacion que pondra a la capital en el tabulador que le corresponde. '
    'Unete al movimiento que si propone."')
r.font.name='Arial'; r.font.size=Pt(12); r.italic=True

doc.add_page_break()

# ─── ZONA 3 ────────────────────────────────────────────────
seccion(doc, 'ZONA 3 — COSTA Y ZONA PORTUARIA')
espacio(doc)
subseccion(doc,'Municipios incluidos:')
parrafo(doc,'Guaymas · Empalme · San Carlos · Cajeme (zona industrial)')
espacio(doc)
subseccion(doc,'Realidad especifica de esta zona:')
bala(doc,'Dinamica economica portuaria e industrial — salarios privados mas altos que el magisterio.')
bala(doc,'Competencia laboral con sector maritimo, maquiladora y pesca.')
bala(doc,'Docentes con antiguedad acumulada pero sueldos desfasados respecto al entorno economico.')
bala(doc,'Acceso limitado a programas de promocion y carrera magisterial.')
espacio(doc)
subseccion(doc,'SU DEMANDA GENETICA: REZONIFICACION PORTUARIA + QUINQUENIOS AL SUELDO BASE')
parrafo(doc,
    '"Trabajar en un municipio costero e industrial en Sonora tiene un costo de vida especifico. El MS exige rezonificacion economica para Guaymas-Empalme y la correcta compactacion de los quinquenios al sueldo base (concepto 07), que hoy te estan pagando mal."')
espacio(doc)
subseccion(doc,'Mensaje de activacion (WhatsApp):')
p = doc.add_paragraph()
r = p.add_run(
    '"Companero docente de Guaymas: llevas anos cotizando mas de lo que te regresa el sistema. '
    'Tus quinquenios estan mal calculados — el MS ya tiene la investigacion legal que demuestra '
    'que el concepto 07 debe compactarse al sueldo base. Eso aumenta tu jubilacion. '
    'Esto es lo que el SNTE no te dice."')
r.font.name='Arial'; r.font.size=Pt(12); r.italic=True

doc.add_page_break()

# ─── ZONA 4 ────────────────────────────────────────────────
seccion(doc, 'ZONA 4 — SUR AGRICOLA (VALLE DEL MAYO Y YAQUI)')
espacio(doc)
subseccion(doc,'Municipios incluidos:')
parrafo(doc,'Ciudad Obregon (Cajeme) · Navojoa · Huatabampo · Alamos · Quiriego · Etchojoa · Bacum · San Ignacio Rio Muerto')
espacio(doc)
subseccion(doc,'Realidad especifica de esta zona:')
bala(doc,'Mayor densidad de maestros federalizados por escuelas rurales y agropecuarias.')
bala(doc,'Alta concentracion de maestros con 20-30 anos de servicio — la generacion mas afectada por el Decimo Transitorio.')
bala(doc,'Muchos con quinquenios no pagados correctamente (concepto 07 desindexado del sueldo base).')
bala(doc,'Ciudad Obregon — sede del SNTE local y presencia de CNTE Sonora (Mtra. Mercedes). Zona de disputa directa.')
espacio(doc)
subseccion(doc,'SU DEMANDA GENETICA: QUINQUENIOS CONCEPTO 07 + PIP')
parrafo(doc,
    '"El maestro del sur de Sonora lleva 25 anos de servicio y al jubilarse descubre que le calcularon mal sus quinquenios durante toda su carrera, y que su pension topa en 10 UMAs aunque haya cotizado mas. El MS tiene la propuesta tecnica para corregir ambos: la compactacion del concepto 07 y la PIP."')
espacio(doc)
subseccion(doc,'Mensaje de activacion (WhatsApp):')
p = doc.add_paragraph()
r = p.add_run(
    '"Maestro del sur: llevas decadas construyendo Sonora. Al momento de jubilarte, el sistema te paga menos de lo que cotizaste. '
    'El MS tiene la investigacion que demuestra que tus quinquenios estan mal liquidados '
    'y que la PIP puede reparar tu pension. No es ideologia — son numeros. Unete."')
r.font.name='Arial'; r.font.size=Pt(12); r.italic=True

doc.add_page_break()

# ─── ZONA 5 ────────────────────────────────────────────────
seccion(doc, 'ZONA 5 — SIERRA E INDIGENA')
espacio(doc)
subseccion(doc,'Municipios incluidos:')
parrafo(doc,'Sahuaripa · Bacanora · Arivechi · Yécora · Rosario · Quiriego · San Pedro de la Cueva · Comunidades Yaqui · Comunidades Mayo · Comunidades Seri · Zona serrana de Alamos')
espacio(doc)
subseccion(doc,'Realidad especifica de esta zona:')
bala(doc,'Maestros de educacion indigena con contratos temporales y plazas sin regularizar.')
bala(doc,'Rezago en plazas de educacion indigena — demanda historicamente incumplida.')
bala(doc,'USICAMM aplicado sin considerar contexto bilingue-bicultural de estas comunidades.')
bala(doc,'Carrera magisterial congelada o inexistente para este segmento.')
bala(doc,'Mayor vulnerabilidad ante lideres locales que explotan su desinformacion.')
espacio(doc)
subseccion(doc,'SU DEMANDA GENETICA: REGULARIZACION DE PLAZAS INDIGENAS + BILATERALIDAD USICAMM')
parrafo(doc,
    '"El maestro indigena de la sierra sonorense trabaja en condiciones que el sistema no reconoce. Sus plazas siguen sin regularizarse. Los concursos de USICAMM no contemplan la especificidad bilingue-bicultural. El MS exige regularizacion de plazas y bilateralidad en los procesos de asignacion para la educacion indigena en Sonora."')
espacio(doc)
subseccion(doc,'Mensaje de activacion (WhatsApp):')
p = doc.add_paragraph()
r = p.add_run(
    '"Companero maestro de la sierra: llevas anos en el aula sin que tu plaza este regularizada. '
    'Los concursos del USICAMM no fueron hechos para tu realidad. '
    'El MS exige que el gobierno federal cumpla con las plazas prometidas y que los procesos '
    'de ingreso y promocion respeten tu contexto. Tu lucha tambien es nuestra propuesta."')
r.font.name='Arial'; r.font.size=Pt(12); r.italic=True

doc.add_page_break()

# ─── SEGMENTOS TRANSVERSALES ───────────────────────────────
seccion(doc, 'SEGMENTOS TRANSVERSALES — CRUZAN TODAS LAS ZONAS')
espacio(doc)
parrafo(doc,
    'Ademas de la segmentacion geografica, existen 3 perfiles demograficos que coexisten en todas las zonas y requieren mensajes especificos:')
espacio(doc)

subseccion(doc,'SEGMENTO A — "La Generacion del Decimo Transitorio"')
parrafo_bold(doc,'Perfil: ','Maestros de 48-62 anos con 20-35 anos de servicio. Seran jubilados bajo el Decimo Transitorio con tope de 10 UMAs.')
parrafo_bold(doc,'Su argumento: ',
    '"La PIP no es abstracta para ti — es la diferencia entre $35,662 y $89,155 de por vida. Ya esta en el Senado. Exigela."')
espacio(doc)

subseccion(doc,'SEGMENTO B — "La Generacion Afore" (maestros jovenes)')
parrafo_bold(doc,'Perfil: ','Maestros de 25-40 anos incorporados despues de 2008. Su ahorro esta en cuenta individual de Afore PENSIONISSSTE.')
parrafo_bold(doc,'Su argumento: ',
    '"Tu pension hoy depende de como le vaya a la bolsa de valores el dia que te jubiles. La PIP propone un fondo colectivo con pension garantizada. Tu tambien tienes algo que ganar con el MS."')
espacio(doc)

subseccion(doc,'SEGMENTO C — "Las Maestras" (mayoria de la base)')
parrafo_bold(doc,'Perfil: ','Las mujeres representan mas del 60% del magisterio federalizado. El sistema las penaliza dos veces: en sueldo activo y en pension futura.')
parrafo_bold(doc,'Su argumento: ',
    '"Cotizaste igual que tus colegas hombres, pero el sistema te jubila con menos. La PIP del MS incluye criterios de equidad de genero en el calculo de pensiones. La justicia laboral tambien es tu derecho."')

doc.add_page_break()

# ─── LA PIP COMO EJE UNIFICADOR ────────────────────────────
seccion(doc, 'LA PIP — EL EJE QUE UNIFICA TODAS LAS ZONAS')
espacio(doc)
parrafo(doc,
    'Independientemente de la zona o segmento, todos los maestros de Sonora comparten un destino comun: la jubilacion bajo un sistema que los castigara si no cambia. La PIP es el lenguaje comun de unificacion:')
espacio(doc)
tabla(doc,
    ['Zona / Segmento','Como llegan a la PIP'],
    [
        ['Frontera','La UMA-Fronteriza es el primer paso — la PIP es la meta final'],
        ['Hermosillo','La rezonificacion mejora tu sueldo hoy — la PIP protege tu pension manana'],
        ['Guaymas','Los quinquenios correctos suben tu base — la PIP sube tu techo'],
        ['Sur (Obregon/Navojoa)','Corregimos el concepto 07 — la PIP cambia el techo de tu jubilacion'],
        ['Sierra / Indigena','Regularizamos tus plazas — la PIP garantiza tu retiro digno'],
        ['Generacion Decimo Transitorio','La PIP puede cambiar lo que te van a pagar de aqui a 3 anos'],
        ['Generacion Afore (jovenes)','La PIP es tu unica garantia frente a las Afores'],
        ['Las maestras','Equidad de genero en el calculo pensionario — incluida en la PIP'],
    ]
)

doc.add_page_break()

# ─── ESTRUCTURA ORGANIZACIONAL ─────────────────────────────
seccion(doc, 'ESTRUCTURA ORGANIZACIONAL — COORDINADORES DE ZONA')
espacio(doc)
parrafo(doc,
    'Para activar esta estrategia, el CORE MS necesita designar un Coordinador de Zona con responsabilidad de activar su zona geografica con los argumentos especificos de su territorio:')
espacio(doc)
tabla(doc,
    ['Zona','Municipios clave','Perfil del Coordinador ideal'],
    [
        ['Zona 1 — Frontera','Nogales, Agua Prieta, San Luis RC','Docente con antiguedad, conocimiento de dinamica fronteriza'],
        ['Zona 2 — Capital','Hermosillo','Docente con habilidades de comunicacion y redes sociales'],
        ['Zona 3 — Costa','Guaymas, Empalme','Docente con conocimiento de quinquenios y tabuladores'],
        ['Zona 4 — Sur','Obregon, Navojoa','Docente con liderazgo de base — zona de disputa con CNTE'],
        ['Zona 5 — Sierra','Sahuaripa, comunidades','Docente bilingue o con conocimiento del sistema de plazas indigenas'],
    ]
)
espacio(doc)
subseccion(doc,'Cada coordinador recibe:')
bala(doc,'Su "carpeta de zona" con argumentos especificos para su region.')
bala(doc,'Kit de mensajes de WhatsApp listos para su zona.')
bala(doc,'Acceso al material tecnico del MS (PIP, UMA-Fronteriza, Quinquenios concepto 07).')
bala(doc,'Linea directa con el CORE para consultas tecnicas.')

doc.add_page_break()

# ─── PLAN DE ACTIVACION ────────────────────────────────────
seccion(doc, 'PLAN DE ACTIVACION EN 3 FASES')
espacio(doc)

subseccion(doc,'FASE 1 — SIEMBRA (Semanas 1-2): "Que cada maestro reciba SU mensaje"')
bala(doc,'Designar coordinadores de zona.')
bala(doc,'Preparar kit de materiales diferenciados por zona (PDF, imagen, mensaje WhatsApp).')
bala(doc,'Identificar grupos de WhatsApp activos por municipio y zona.')
bala(doc,'Primer mensaje en cada grupo: el argumento especifico de esa zona.')
bala(doc,'Crear un hilo de "introduccion MS" diferenciado por zona.')
espacio(doc)

subseccion(doc,'FASE 2 — ARRAIGO (Semanas 3-5): "Que cada maestro conozca la PIP"')
bala(doc,'Segunda ronda de mensajes conectando la demanda local con la PIP.')
bala(doc,'Video corto (2-3 min) del CORE explicando: "Tu zona, tu demanda, tu PIP".')
bala(doc,'Evento virtual por zona (videollamada de 30 min) con coordinador de zona.')
bala(doc,'Recoleccion de firmas de adhesion por zona — dato de fuerza para la Carpeta Ejecutiva.')
espacio(doc)

subseccion(doc,'FASE 3 — MOVILIZACION (Semanas 6+): "Una base unida con una voz"')
bala(doc,'Comunicado unificado firmado por coordinadores de todas las zonas.')
bala(doc,'Presentacion publica de la Carpeta Ejecutiva con respaldo territorial.')
bala(doc,'Gestion de reunion con legisladores federales de Sonora (uno por zona).')
bala(doc,'Evento presencial MS en Hermosillo: "Sonora Unida por Pensiones Dignas".')

doc.add_page_break()

# ─── CONCLUSION ────────────────────────────────────────────
seccion(doc, 'CONCLUSION ESTRATEGICA')
espacio(doc)
parrafo(doc,'La diferencia entre el MS y la CNTE no es solo tecnica — es de precision.')
espacio(doc)
parrafo(doc,
    'La CNTE lanza una red generica y pesca poco. El MS lanza anzuelos especificos y pesca a cada maestro por su necesidad real.')
espacio(doc)
p = doc.add_paragraph()
r = p.add_run(
    '"Cuando cada maestro de Sonora sienta que el MS habla de SU problema, '
    'el MS tendra la base mas solida del estado — sin haber parado una sola escuela."')
r.font.name='Arial'; r.font.size=Pt(12); r.bold=True; r.italic=True
espacio(doc)
subseccion(doc,'El arbol del MS crece desde sus raices:')
bala(doc,'PIP como raiz profunda que nutre a todos.')
bala(doc,'5 zonas como ramas que llegan a cada municipio de Sonora.')
bala(doc,'3 segmentos transversales (generaciones + genero) que completan la cobertura.')
bala(doc,'Coordinadores de zona como el sistema que hace circular la informacion.')
espacio(doc)
doc.add_paragraph('_' * 60)
parrafo(doc,'Estrategia elaborada para uso exclusivo del CORE del Magisterio Sonorense')
parrafo(doc,'Fecha: 03/05/2026 | Version 1.0')

# GUARDAR
out = r'C:\Users\USUARIO\.gemini\antigravity\scratch\investigaciones\Estrategia_Unificacion_Base_MS_Zonas.docx'
doc.save(out)
print(f"LISTO. Guardado en: {out}")
